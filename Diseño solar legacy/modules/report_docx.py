"""
SALVI Solar – Informe Word profesional.
Design System: Salvi Studio v1 — tokens, tipografía y jerarquía de color.

Tokens aplicados:
  --salvi-black:   #1E1E1E   acción primaria, texto, cabeceras
  --salvi-grey:    #6A6A6A   texto secundario, metadatos
  --salvi-cream:   #FCF9F5   fondo superficies amplias
  --salvi-surface: #F7F4EF   cards, filas alternas
  --salvi-line:    #E8E2D8   bordes, separadores, tablas
  --salvi-muted:   #A09A91   texto apagado, notas al pie
  --state-success: #1F7A4D   verde SOLO para estados OK (fiabilidad, CO₂)
  --state-warning: #B7791F   ámbar para riesgos, año 10
  --state-danger:  #B42318   rojo para errores y valores críticos

Fuentes:
  Georgia (serif) → simula "Exposure" — portada, H1 de sección
  Calibri (sans)  → simula "Helvetica Neue" — cuerpo, tablas, H2
"""
from io import BytesIO
from datetime import datetime

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Design tokens ──────────────────────────────────────────────────────────────
T = {
    'black':   '#1E1E1E',
    'grey':    '#6A6A6A',
    'cream':   '#FCF9F5',
    'surface': '#F7F4EF',
    'line':    '#E8E2D8',
    'muted':   '#A09A91',
    'success': '#1F7A4D',
    'warning': '#B7791F',
    'danger':  '#B42318',
    'white':   '#FFFFFF',
}

def _rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

RC = {k: _rgb(v) for k, v in T.items()}

FONT_BRAND = 'Georgia'    # Exposure substitute
FONT_UI    = 'Calibri'    # Helvetica Neue substitute

MONTHS_ES   = ['Ene','Feb','Mar','Abr','May','Jun','Jul','Ago','Sep','Oct','Nov','Dic']
MONTHS_FULL = ['Enero','Febrero','Marzo','Abril','Mayo','Junio',
               'Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']


# ── Matplotlib style (Salvi palette) ───────────────────────────────────────────
plt.rcParams.update({
    'font.family':        'DejaVu Sans',
    'axes.spines.top':    False,
    'axes.spines.right':  False,
    'axes.grid':          True,
    'grid.alpha':         0.20,
    'grid.linestyle':     '--',
    'grid.color':         '#E8E2D8',
    'figure.facecolor':   '#FFFFFF',
    'axes.facecolor':     '#FFFFFF',
    'axes.edgecolor':     '#6A6A6A',
    'axes.labelcolor':    '#1E1E1E',
    'xtick.color':        '#6A6A6A',
    'ytick.color':        '#6A6A6A',
    'text.color':         '#1E1E1E',
})


# ── XML/cell helpers ───────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def _cell_text(cell, text, size=9, bold=False, color_key='black',
               align=WD_ALIGN_PARAGRAPH.LEFT, font=FONT_UI):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text) if text is not None else '–')
    r.font.name = font
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RC.get(color_key, RC['black'])
    return r

def _add_hr(doc, color_key='line'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), T[color_key].lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def _h1(doc, text):
    """Section heading — Georgia (brand/Exposure), black, uppercase."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text.upper())
    r.font.name  = FONT_BRAND
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = RC['black']
    # thin rule after heading
    return p

def _h2(doc, text):
    """Sub-heading — Calibri bold, black."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name  = FONT_UI
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = RC['black']
    return p

def _body(doc, text, size=9.5, italic=False, color_key='black', space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.name   = FONT_UI
    r.font.size   = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = RC.get(color_key, RC['black'])
    return p

def _img(doc, png_bytes, width_inches=6.0):
    doc.add_picture(BytesIO(png_bytes), width=Inches(width_inches))

def _quote(doc, text, attribution=None):
    """Pull-quote card — surface background, Georgia italic bold, brand tone of voice."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, T['surface'].lstrip('#'))
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4) if attribution else Pt(8)
    r = p.add_run(f'“{text}”')
    r.font.name  = FONT_BRAND
    r.font.size  = Pt(11.5)
    r.font.italic = True
    r.font.bold  = True
    r.font.color.rgb = RC['black']
    if attribution:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(attribution)
        r2.font.name = FONT_UI
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RC['grey']
    return tbl

def _table_header(table, cols, font=FONT_UI):
    """Black header row, white text — per Design System primary action = black."""
    row = table.rows[0]
    for cell, text in zip(row.cells, cols):
        _cell_text(cell, text, size=8.5, bold=True, color_key='white',
                   align=WD_ALIGN_PARAGRAPH.CENTER, font=font)
        _set_cell_bg(cell, T['black'])

def _fmt_eur(v):
    if v is None: return '–'
    return f'{int(round(v)):,} €'.replace(',','.')

def _safe(d, *keys, default=None):
    for k in keys:
        if isinstance(d, dict): d = d.get(k)
        else: return default
        if d is None: return default
    return d


# ── Chart generators (Salvi palette) ──────────────────────────────────────────

def _chart_prod_cons(monthly_data, product_name=''):
    prod = [m.get('production_kwh') or m.get('prod_kwh') or 0 for m in monthly_data[:12]]
    cons = [m.get('consumption_kwh') or m.get('cons_kwh') or 0 for m in monthly_data[:12]]
    x, w = np.arange(12), 0.38
    fig, ax = plt.subplots(figsize=(9, 3.2))
    ax.bar(x - w/2, prod, w, label='Producción solar (kWh)', color=T['black'], alpha=0.82)
    ax.bar(x + w/2, cons, w, label='Consumo estimado (kWh)', color=T['grey'],  alpha=0.60)
    ax.set_xticks(x); ax.set_xticklabels(MONTHS_ES, fontsize=9)
    ax.set_ylabel('kWh/mes', fontsize=9)
    ax.legend(fontsize=9, framealpha=0.4)
    if product_name:
        ax.set_title(f'Producción vs Consumo mensual — {product_name}',
                     fontsize=10, fontweight='bold', pad=8)
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig); return buf.getvalue()


def _chart_soc(monthly_data, product_name=''):
    soc1  = [m.get('soc_min_pct') or m.get('soc_new') or 0 for m in monthly_data[:12]]
    soc10 = [m.get('soc_min_y10') or m.get('soc_y10') or 0 for m in monthly_data[:12]]
    x = np.arange(12)
    fig, ax = plt.subplots(figsize=(9, 3.0))
    ax.fill_between(x, soc1, alpha=0.08, color=T['black'])
    ax.plot(x, soc1,  'o-',  color=T['black'],   linewidth=2, markersize=4, label='SOC mín. año 1')
    ax.plot(x, soc10, 's--', color=T['warning'],  linewidth=1.8, markersize=4, label='SOC mín. año 10')
    ax.axhline(20, color=T['danger'], linewidth=0.8, linestyle=':', alpha=0.7)
    ax.text(11.15, 22, 'Límite crítico', fontsize=7.5, color=T['danger'], va='bottom')
    ax.set_xticks(x); ax.set_xticklabels(MONTHS_ES, fontsize=9)
    ax.set_ylim(0, 108); ax.set_ylabel('SOC (%)', fontsize=9)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda y,_: f'{int(y)}%'))
    ax.legend(fontsize=9, framealpha=0.4)
    ax.set_title('Estado de carga de batería (SOC mínimo mensual)',
                 fontsize=10, fontweight='bold', pad=8)
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig); return buf.getvalue()


def _chart_loss_waterfall(lt):
    labels = [
        'Recurso solar teórico',
        'Geometría / tilt',
        'Suciedad (soiling)',
        'Temperatura',
        'Regulador / cableado',
        'Batería (ida+vuelta)',
        'Degradación promedio',
        'Energía útil disponible',
    ]
    losses = [
        0,
        lt.get('geometry_loss_pct', 0),
        lt.get('soiling_loss_pct', 0),
        lt.get('temperature_loss_pct', 0),
        lt.get('controller_loss_pct', 0),
        lt.get('battery_roundtrip_loss_pct', 0),
        lt.get('degradation_loss_pct', 0),
        0,
    ]
    running = 100.0
    values = [100.0]
    for l in losses[1:-1]:
        running = max(0, running - l)
        values.append(running)
    values.append(lt.get('energy_available_pct', running))

    colors = [T['black']] + [T['grey']] * 6 + [T['success']]
    fig, ax = plt.subplots(figsize=(9, 3.8))
    y = np.arange(len(labels))
    bars = ax.barh(y, values, color=colors, alpha=0.82, height=0.52)
    for bar, v, loss, lbl in zip(bars, values, losses, labels):
        if lbl == 'Recurso solar teórico':
            label = f'{v:.1f}%'
        elif lbl == 'Energía útil disponible':
            label = f'{v:.1f}%'
        else:
            label = f'−{loss:.1f}%  →  {v:.1f}%'
        ax.text(max(v + 1.5, 3), bar.get_y() + bar.get_height()/2,
                label, va='center', fontsize=8.5, color=T['black'],
                fontweight='bold' if lbl in ('Recurso solar teórico','Energía útil disponible') else 'normal')
    ax.set_yticks(y); ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlim(0, 135)
    ax.set_xlabel('Energía remanente (%)', fontsize=9)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x,_: f'{int(x)}%'))
    ax.invert_yaxis()
    ax.set_title('Árbol de pérdidas del sistema (acumulativo)', fontsize=10, fontweight='bold', pad=8)
    ax.spines['left'].set_visible(False); ax.tick_params(axis='y', length=0)
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig); return buf.getvalue()


def _chart_tco_comparison(candidates):
    valid = [c for c in candidates if not c.get('error') and c.get('tco_10y_sale')]
    if not valid: return None
    names    = [c.get('product_name') or c.get('product_id','') for c in valid]
    tco_sale = [c['tco_10y_sale'] for c in valid]
    tco_cost = [c.get('tco_10y_cost') or 0 for c in valid]
    margin   = [s - c for s, c in zip(tco_sale, tco_cost)]
    recs     = [c.get('recommended', False) for c in valid]
    x = np.arange(len(valid))
    fig, ax = plt.subplots(figsize=(8, 3.0))
    bar_colors = [T['black'] if r else T['grey'] for r in recs]
    ax.bar(x, tco_cost, label='Coste directo',     color=bar_colors, alpha=0.80)
    ax.bar(x, margin, bottom=tco_cost, label='Margen comercial',
           color=[T['muted']]*len(valid), alpha=0.55)
    for i, (v, r) in enumerate(zip(tco_sale, recs)):
        ax.text(i, v + max(tco_sale)*0.02, _fmt_eur(v),
                ha='center', va='bottom', fontsize=8,
                fontweight='bold' if r else 'normal',
                color=T['black'] if r else T['grey'])
    ax.set_xticks(x); ax.set_xticklabels(names, fontsize=8.5, rotation=10, ha='right')
    ax.set_ylabel('TCO 10 años (€)', fontsize=9)
    ax.legend(fontsize=8.5, framealpha=0.4)
    ax.set_title('Comparativa TCO 10 años', fontsize=10, fontweight='bold', pad=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda y,_: f'{int(y/1000)}k €'))
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig); return buf.getvalue()


def _chart_smartec_dimming(periods, system_power_w, night_hours=12):
    if not periods: return None
    fig, ax = plt.subplots(figsize=(8, 2.6))
    colors_p = [T['black'], T['grey'], T['muted']]
    t = 0
    for i, p in enumerate(periods[:3]):
        dur  = p.get('duration_pct', 0.333) * night_hours
        pres = p.get('presence_ratio', 0.3)
        dp   = p.get('dimming_presence', 0.8)
        dnp  = p.get('dimming_no_presence', 0.2)
        avg  = pres * dp + (1 - pres) * dnp
        ax.barh(0, dur, left=t, height=0.38,
                color=colors_p[i % 3], alpha=0.82,
                label=f'P{i+1} — pres. {dp*100:.0f}% / sin pres. {dnp*100:.0f}%')
        ax.text(t + dur/2, 0, f'{avg*100:.0f}%\n{system_power_w*avg:.0f} W',
                ha='center', va='center', fontsize=8.5,
                color='white', fontweight='bold')
        t += dur
    ax.set_xlim(0, night_hours); ax.set_yticks([])
    ax.set_xlabel('Horas de noche', fontsize=9)
    ax.set_title('Perfil de dimming nocturno Smartec', fontsize=10, fontweight='bold', pad=8)
    ax.legend(fontsize=8, loc='lower right', framealpha=0.4)
    ax.xaxis.set_major_locator(mticker.MaxNLocator(integer=True))
    fig.tight_layout()
    buf = BytesIO(); fig.savefig(buf, format='png', dpi=140, bbox_inches='tight')
    plt.close(fig); return buf.getvalue()


# ── Technology descriptions ─────────────────────────────────────────────────────

TECH_DESC = {
    'sil_horizontal': {
        'title': 'Panel solar horizontal integrado — SIL-H',
        'body': (
            'El panel SIL en configuración horizontal emplea células monocristalinas PERC '
            '(Passivated Emitter and Rear Cell) con ángulo de inclinación optimizado para '
            'la latitud del emplazamiento mediante el algoritmo PVGIS/JRC. '
            'Este formato integrado en la luminaria minimiza la carga de viento al presentar '
            'sección frontal reducida, y maximiza la captación de irradiancia directa durante '
            'las horas centrales del día. La tecnología PERC reduce la recombinación en la '
            'cara posterior de la célula, incrementando la eficiencia entre 0.5 y 1% respecto '
            'a células convencionales, con rendimientos típicos del 21–22%.'
        ),
    },
    'cylinder_250': {
        'title': 'Panel cilíndrico omnidireccional — SIL-C 250',
        'body': (
            'El panel cilíndrico SIL-C representa una innovación radical en el diseño de paneles '
            'solares para alumbrado público. Su geometría cilíndrica con 12 sectores facetados '
            'permite la captación de irradiancia solar desde cualquier orientación azimutal, '
            'eliminando la necesidad de orientar la luminaria hacia el sur. '
            'Las células Back Contact (BC) ubican todos los contactos eléctricos en la cara '
            'posterior de la célula, eliminando las pérdidas por sombreado de rejilla frontal. '
            'Eficiencia típica de células BC M10: 24–25%, con respuesta mejorada a luz difusa.'
        ),
    },
    'cylinder_300': {
        'title': 'Panel cilíndrico omnidireccional — SIL-C 300',
        'body': (
            'Versión de mayor potencia del panel cilíndrico SIL-C, con células Back Contact M10 '
            '(182×182 mm) en configuración 3×9 por sector, alcanzando ~27 W por sector y ~324 Wp '
            'totales en el conjunto de 12 sectores. La tecnología BC elimina el sombreado frontal '
            'y mejora la respuesta a luz difusa. Captación omnidireccional ideal para instalaciones '
            'con orientación variable de la vía.'
        ),
    },
    'cylinder_350': {
        'title': 'Panel cilíndrico omnidireccional — SIL-C 350',
        'body': (
            'Variante SIL-C de mayor densidad, con células Back Contact M10 en configuración '
            '3×10 por sector (30 células/sector × 12 sectores). Con eficiencias BC superiores '
            'al 24.5%, el sistema alcanza aproximadamente 360 Wp totales en formato compacto. '
            'Especialmente eficiente en zonas de alta irradiancia difusa (costa, clima atlántico).'
        ),
    },
    'double_vertical_eo': {
        'title': 'Panel dual Este-Oeste vertical — SIL-EO',
        'body': (
            'La configuración dual Este-Oeste emplea dos paneles verticales orientados '
            'perpendicularmente al eje de la vía, capturando irradiancia en las horas de '
            'mañana (panel Este) y tarde (panel Oeste). Esta distribución temporal de la '
            'captación evita picos de potencia al mediodía solar, reduciendo el estrés térmico '
            'de la batería. Células Back Contact M10 en 3×8 filas por panel (24 cél./panel), '
            '~200 Wp totales. La orientación vertical reduce el depósito de suciedad superficial '
            'respecto a paneles inclinados.'
        ),
    },
    'custom_orientable': {
        'title': 'Panel orientable de alta eficiencia — SIL-OPT',
        'body': (
            'El panel orientable SIL-OPT combina células monocristalinas PERC de alta eficiencia '
            'con sistema de ajuste de inclinación, permitiendo optimizar el ángulo de tilt para '
            'cada latitud. El algoritmo PVGIS determina el ángulo óptimo que maximiza la '
            'producción anual, típicamente entre 25° y 40° según latitud.'
        ),
    },
    'sil_independent': {
        'title': 'Panel solar independiente de orientación libre — SIL-IND',
        'body': (
            'El panel independiente SIL-IND se instala separado del cuerpo de la luminaria, '
            'sobre un soporte propio orientable libremente hacia el sur, sin las restricciones '
            'geométricas del poste. Esta libertad de orientación e inclinación permite maximizar '
            'la captación solar para la latitud exacta del emplazamiento, con el algoritmo PVGIS '
            'determinando el ángulo óptimo. El panel y la capacidad de batería se dimensionan a '
            'medida para cada proyecto, en lugar de partir de un tamaño de catálogo fijo, '
            'garantizando la fiabilidad objetivo con el menor coste total posible.'
        ),
    },
}

def _tech(geometry_type):
    return TECH_DESC.get(geometry_type, {'title': geometry_type, 'body': 'Geometría específica del producto.'})


# ── Narrative generators ────────────────────────────────────────────────────────

def _narrative_recommendation(rec, all_candidates, project, photometry):
    if not rec: return ''
    product_name = rec.get('product_name') or rec.get('product_id','–')
    rel1         = 100 - (rec.get('annual_failure_rate_pct') or 0)
    rel10_raw    = rec.get('annual_failure_rate_pct_y10')
    rel10        = (100 - rel10_raw) if rel10_raw is not None else None
    city         = project.get('city') or 'la ubicación seleccionada'
    lat          = project.get('lat') or 0
    n_cands      = len([c for c in all_candidates if not c.get('error')])
    lclass       = photometry.get('lighting_class','–')
    autonomy     = rec.get('autonomy_days')
    pv_wp        = rec.get('pv_peak_power_wp', 0)
    bat_wh       = rec.get('battery_nominal_wh', 0)
    n_units      = rec.get('n_units', 1)

    monthly = rec.get('monthly_data') or []
    worst_month, worst_soc_val = '', 100
    if monthly:
        socs = [(m.get('soc_min_pct') or m.get('soc_new') or 100, i) for i, m in enumerate(monthly[:12])]
        worst_soc_val, worst_i = min(socs)
        worst_month = MONTHS_FULL[worst_i]

    others = [c for c in all_candidates if not c.get('error') and c.get('product_id') != rec.get('product_id')]
    second = others[0] if others else None

    paras = []
    paras.append(
        f'El análisis de {n_cands} soluciones candidatas bajo las condiciones específicas de '
        f'{city} ({lat:.3f}°N) ha determinado que {product_name} '
        f'presenta el óptimo técnico-económico para este proyecto de alumbrado público clase {lclass}.'
    )
    paras.append(
        f'Con {pv_wp} Wp de panel fotovoltaico y {bat_wh} Wh de capacidad de batería'
        + (f' (configuración de {n_units} unidades)' if n_units and n_units > 1 else '')
        + f', el sistema garantiza una fiabilidad del {rel1:.1f}% en el primer año de operación, '
        f'equivalente a cubrir con plenas condiciones de servicio {rel1/100*365:.0f} noches de cada 365.'
        + (f' En el mes crítico de {worst_month}, el estado de carga mínimo de batería es del '
           f'{worst_soc_val:.0f}%, dentro del margen operacional seguro.' if worst_month else '')
    )
    if rel10 is not None:
        trend = 'manteniéndose dentro del umbral de servicio' if rel10 >= 95 else 'con ligera degradación controlada'
        paras.append(
            f'La proyección al décimo año, incorporando la degradación natural de la batería '
            f'(~2% anual de capacidad), sitúa la fiabilidad en el {rel10:.1f}%, '
            f'{trend}. Este resultado confirma la idoneidad del dimensionado para el ciclo de vida completo del sistema.'
        )
    if second and rec.get('tco_10y_sale') and second.get('tco_10y_sale'):
        rel_second = 100 - (second.get('annual_failure_rate_pct') or 0)
        tco_diff = rec['tco_10y_sale'] - second['tco_10y_sale']
        rel_diff = rel1 - rel_second
        if abs(tco_diff) > 50:
            direction = 'superior' if tco_diff > 0 else 'inferior'
            paras.append(
                f'En comparación con la siguiente alternativa evaluada '
                f'({second.get("product_name") or second.get("product_id","–")}), '
                f'la solución recomendada presenta un TCO {direction} en {abs(tco_diff):.0f} € '
                f'a 10 años, con una variación de {rel_diff:+.1f} puntos porcentuales en fiabilidad. '
                f'El análisis de valor identifica esta relación coste-prestación como la óptima para los parámetros del proyecto.'
            )
    if autonomy:
        paras.append(
            f'La autonomía de diseño de {autonomy} días sin irradiancia solar proporciona resiliencia '
            f'ante episodios de meteorología adversa prolongada, especialmente relevante en instalaciones '
            f'de alumbrado viario donde la continuidad del servicio es crítica.'
        )
    return '\n\n'.join(paras)


def _smartec_savings(periods, sys_pw, night_hours=12):
    """Ahorro energético real Smartec vs. potencia plena constante, calculado a partir del
    % de presencia y la curva de dimming con/sin presencia configurada en el proyecto.
    Es una variable de cada informe, no una cifra de marketing fija — puede llegar hasta
    ~90% en vías de tráfico muy bajo con dimming agresivo en ausencia de presencia."""
    if not periods or not sys_pw:
        return None, None, None
    full_wh = sys_pw * night_hours
    smartec_wh = 0
    for p in periods:
        dur_h = p.get('duration_pct', 0) * night_hours
        pres  = p.get('presence_ratio', 0.3)
        dp    = p.get('dimming_presence', 0.8)
        dnp   = p.get('dimming_no_presence', 0.2)
        smartec_wh += sys_pw * (pres * dp + (1 - pres) * dnp) * dur_h
    savings_pct = (1 - smartec_wh / full_wh) * 100 if full_wh > 0 else 0
    return savings_pct, smartec_wh, full_wh


def _narrative_smartec(periods, photometry, rec):
    if not periods:
        return 'El perfil de control nocturno Smartec no ha sido configurado para esta simulación.'
    sys_pw = photometry.get('system_power_w', 0)
    savings_pct, smartec_wh, full_wh = _smartec_savings(periods, sys_pw)
    bat_wh = (rec.get('battery_nominal_wh') or 0) if rec else 0
    text = (
        f'El sistema Smartec implementa un algoritmo de gestión predictiva de la energía que '
        f'divide la noche en {len(periods)} períodos funcionales con niveles de dimming '
        f'diferenciados en función de la presencia de tráfico y peatones.\n\n'
        f'Con los parámetros configurados, el consumo nocturno estimado es de {smartec_wh:.0f} Wh, '
        f'frente a los {full_wh:.0f} Wh de operación a potencia plena constante. '
        f'Esto representa un ahorro energético del {savings_pct:.1f}% respecto al modo sin dimming — '
        f'un valor propio de este proyecto, calculado a partir del % de presencia y la curva de '
        f'dimming configurados, y no una cifra fija: según el tipo de vía y el nivel de dimming en '
        f'ausencia de presencia, el ahorro puede llegar hasta un 90%.\n\n'
        'La gestión predictiva opera en tres niveles:\n'
        '• Modo presencia: iluminación al nivel de conformidad normativa cuando se detecta '
        'movimiento en la zona de influencia del sensor.\n'
        '• Modo stand-by: reducción al nivel de dimming mínimo en ausencia de presencia, '
        'manteniendo la visibilidad de seguridad requerida por la clase de alumbrado.\n'
        '• Modo protección: activación automática cuando el SOC de la batería cae por debajo '
        'del umbral crítico, garantizando la continuidad del servicio mediante reducción adicional del nivel de luz.'
    )
    if bat_wh:
        text += (
            f'\n\nLa batería de {bat_wh:.0f} Wh dimensionada para esta solución incorpora el perfil '
            f'Smartec como parámetro de cálculo, optimizando la relación panel/batería. '
            f'El ahorro energético del {savings_pct:.1f}% se traduce directamente en menores ciclos '
            f'de carga/descarga = mayor vida útil de la batería.'
        )
    return text


# ── Section builders ───────────────────────────────────────────────────────────

def _build_cover(doc, project, rec, sim_date):
    """Portada: cabecera negra, Georgia, ficha proyecto, 4 KPIs."""

    # ── Black header bar ──
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, T['black'].lstrip('#'))
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    # Brand name — Georgia (Exposure substitute)
    r1 = p.add_run('SALVI SOLAR')
    r1.font.name  = FONT_BRAND
    r1.font.size  = Pt(28)
    r1.font.bold  = True
    r1.font.color.rgb = RC['white']
    p.add_run('\n')
    r2 = p.add_run('Informe de simulación fotovoltaica autónoma para alumbrado público')
    r2.font.name  = FONT_UI
    r2.font.size  = Pt(10.5)
    r2.font.bold  = False
    r2.font.color.rgb = RGBColor(0xA0, 0x9A, 0x91)  # muted on dark

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Project info table ──
    pairs = [
        ('Proyecto',      project.get('name') or '–'),
        ('Ubicación',     f"{project.get('city') or '–'}, {project.get('country') or '–'}"),
        ('Coordenadas',   f"{project.get('lat','–')}, {project.get('lon','–')}"),
        ('Solución rec.', rec.get('product_name') if rec else '–'),
        ('Módulo',        'SALVI Solar'),
        ('Fecha',         sim_date),
    ]
    info = doc.add_table(rows=len(pairs), cols=2)
    info.style = 'Table Grid'
    for i, (row, (k, v)) in enumerate(zip(info.rows, pairs)):
        _cell_text(row.cells[0], k, size=9, bold=True, color_key='grey')
        _cell_text(row.cells[1], v, size=9)
        _set_cell_bg(row.cells[0], T['surface'].lstrip('#'))
        _set_cell_bg(row.cells[1], T['cream'].lstrip('#') if i % 2 else T['white'].lstrip('#'))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _add_hr(doc, 'line')

    # ── 4-column KPI row ──
    if rec:
        rel1      = 100 - (rec.get('annual_failure_rate_pct') or 0)
        rel10_raw = rec.get('annual_failure_rate_pct_y10')
        rel10     = (100 - rel10_raw) if rel10_raw is not None else None
        kpi_labels = ['FIABILIDAD AÑO 1',  'TCO 10 AÑOS',
                      'CO₂ EVITADO 10a',   'AUTONOMÍA SIN SOL']
        kpi_values = [
            f'{rel1:.1f}%',
            _fmt_eur(rec.get('tco_10y_sale')),
            f'{int(round(rec["co2_saved_10y_kg"]))} kg' if rec.get('co2_saved_10y_kg') else '–',
            f'{rec["autonomy_days"]} días' if rec.get('autonomy_days') else '–',
        ]
        # color_key for value: success if positive KPI, else black
        kpi_colors = ['success', 'black', 'success', 'black']
        kpi_tbl = doc.add_table(rows=2, cols=4)
        kpi_tbl.style = 'Table Grid'
        for cell, lbl in zip(kpi_tbl.rows[0].cells, kpi_labels):
            _cell_text(cell, lbl, size=8, bold=True, color_key='grey',
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(cell, T['surface'].lstrip('#'))
        for cell, val, ck in zip(kpi_tbl.rows[1].cells, kpi_values, kpi_colors):
            _cell_text(cell, val, size=16, bold=True, color_key=ck,
                       align=WD_ALIGN_PARAGRAPH.CENTER, font=FONT_BRAND)
            _set_cell_bg(cell, T['white'].lstrip('#'))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    _body(doc,
          'Documento generado automáticamente por SALVI Solar Studio. '
          'Datos de irradiancia: PVGIS/JRC, Comisión Europea. Confidencial.',
          size=8, italic=True, color_key='muted', space_after=0)


def _build_value_proposition(doc, rec, night=None, photometry=None):
    """Sección 01: propuesta de valor SIL + SMARTEC — por qué la gestión dinámica es imprescindible."""
    night      = night or {}
    photometry = photometry or {}
    _h1(doc, '01 · Por qué SIL + SMARTEC')
    _add_hr(doc, 'line')

    _quote(doc,
        'El alumbrado solar no falla porque el sol sea insuficiente. Falla cuando el panel, '
        'la batería, el controlador y la luminaria no trabajan de forma coordinada.')

    _body(doc,
        'Muchas soluciones solares convencionales dimensionan el sistema para el peor mes del '
        'año y aplican una reducción de potencia preprogramada durante toda la noche, con '
        'independencia de la presencia real de tráfico o peatones. Esta rigidez desaprovecha el '
        'excedente energético de los meses favorables y compromete la autonomía en los meses críticos.')
    _body(doc,
        'SIL con SMARTEC no es una luminaria LED con un panel añadido: es un sistema energético '
        'autónomo donde el panel fotovoltaico, la batería, el controlador y la luminaria trabajan '
        'de forma coordinada, con una lógica de gestión que decide cada noche cuánta energía '
        'consumir para garantizar el servicio hasta el amanecer.')

    _h2(doc, 'SMARTEC Adaptive Logic')
    _body(doc,
        'El algoritmo Adaptive Logic analiza en tiempo real la energía almacenada, el consumo '
        'previsto y la previsión de carga del día siguiente, ajustando de forma predictiva el '
        'nivel de luz para asegurar que la luminaria llegue encendida hasta el amanecer. Si '
        'detecta menor capacidad de carga esperada, reduce el consumo nocturno de forma '
        'preventiva, en lugar de esperar a que la batería llegue a un nivel crítico.')

    _h2(doc, 'Iluminación adaptativa y red mallada')
    _body(doc,
        'Sensores de presencia elevan el nivel de luz solo cuando hay actividad real —peatones '
        'o vehículos— y lo reducen al nivel base cuando la vía está vacía. Las luminarias SMARTEC '
        'además se comunican entre sí mediante una red mallada: al detectar movimiento, una '
        'luminaria avisa a las vecinas para que se anticipen, generando una onda de luz '
        'coordinada en lugar de puntos aislados de detección.')

    _h2(doc, 'Protección activa de la batería')
    _body(doc,
        'SMARTEC gestiona de forma dinámica la profundidad de descarga de la batería, evitando '
        'descargas profundas innecesarias noche tras noche. No se trata de poder descargar la '
        'batería al 100%, sino de no tener que hacerlo cada noche: esta gestión activa alarga la '
        'vida útil real de la batería frente a un dimensionado fijo.')

    savings_pct, _, _ = _smartec_savings(night.get('periods', []), photometry.get('system_power_w', 0))
    _h2(doc, 'Ahorro real frente a una reducción preprogramada')
    if savings_pct is not None:
        _body(doc,
            f'Para este proyecto, con el % de presencia y la curva de dimming configurados, el '
            f'ahorro energético de SMARTEC frente a una operación a potencia plena constante es '
            f'del {savings_pct:.1f}% (ver detalle en la sección 04). No es una cifra de marketing '
            f'fija: es una variable propia de cada instalación, que depende del tipo de vía y del '
            f'tráfico real, y que en vías de baja intensidad con dimming agresivo en ausencia de '
            f'presencia puede llegar hasta un 90%.')
    else:
        _body(doc,
            'El ahorro energético de SMARTEC frente a una operación a potencia plena constante no '
            'es una cifra de marketing fija: depende del % de presencia y de la curva de dimming '
            'configurados para cada proyecto, y puede llegar hasta un 90% en vías de baja '
            'intensidad con dimming agresivo en ausencia de presencia.')

    geo = (rec.get('geometry_type') if rec else '') or ''
    if geo in ('cylinder_250', 'cylinder_300', 'cylinder_350', 'double_vertical_eo'):
        _h2(doc, 'Paneles Back Contact: resiliencia frente a sombras y suciedad')
        _body(doc,
            'En un panel solar convencional, una sombra parcial o un depósito de suciedad '
            'localizado (polvo, hojas, excrementos de aves) puede provocar una pérdida de '
            'potencia desproporcionada respecto a la superficie afectada, del orden del 70 al '
            '85%. La tecnología Back Contact (BC) empleada en este panel ubica los contactos '
            'eléctricos en la cara posterior de la célula, reduciendo esa pérdida residual hasta '
            'aproximadamente un 8-15%.',
            size=9, color_key='grey')

    _body(doc,
        'El resultado es que la energía disponible deja de estar limitada por el peor mes del '
        'año: el sistema aprovecha el excedente de los meses favorables y protege la autonomía '
        'en los meses críticos, en lugar de aplicar un perfil fijo pensado para el escenario más '
        'desfavorable.', italic=True, color_key='grey')


def _build_municipal_benefits(doc):
    """Sección 05: beneficios segmentados por perfil de interlocutor municipal."""
    doc.add_page_break()
    _h1(doc, '06 · Beneficios para el municipio')
    _add_hr(doc, 'line')

    blocks = [
        ('Responsable político', [
            'Mayor seguridad percibida y mejor servicio al ciudadano.',
            'Reducción de incidencias y menor coste operativo.',
            'Imagen de ciudad innovadora y sostenible.',
            'Menor dependencia de obra civil y de la red eléctrica.',
            'Menor huella energética y ambiental.',
        ]),
        ('Ingeniero municipal', [
            'Datos reales de funcionamiento y diagnóstico remoto.',
            'Alarmas y mantenimiento predictivo.',
            'Control de SOC, temperatura, potencia fotovoltaica y potencia LED.',
            'Ajuste dinámico de niveles de iluminación.',
            'Fotometrías adaptadas a la norma EN 13201.',
            'Plataforma abierta (DALI-2, d4i, Zhaga Book 18) y escalable.',
        ]),
        ('Instalador / mantenedor', [
            'Menos visitas correctivas.',
            'Identificación rápida de averías.',
            'Reprogramación remota de parámetros.',
            'Información técnica clara para intervenir solo cuando es necesario.',
        ]),
    ]
    for title, items in blocks:
        _h2(doc, title)
        for item in items:
            _body(doc, f'•  {item}', size=9.5, space_after=2)


def _build_executive_summary(doc, rec, all_candidates, project, photometry):
    """Sección 02: resumen ejecutivo con gráficos."""
    _h1(doc, '02 · Resumen ejecutivo')
    _add_hr(doc, 'line')

    if not rec:
        _body(doc, 'No se ha podido identificar una solución recomendada.')
        return

    monthly   = rec.get('monthly_data') or []
    rel1      = 100 - (rec.get('annual_failure_rate_pct') or 0)
    rel10_raw = rec.get('annual_failure_rate_pct_y10')
    rel10     = (100 - rel10_raw) if rel10_raw is not None else None
    city      = project.get('city') or 'la ubicación'
    n_cands   = len([c for c in all_candidates if not c.get('error')])
    lclass    = photometry.get('lighting_class','–')

    _body(doc,
        f'Análisis de {n_cands} soluciones candidatas para alumbrado público clase {lclass} '
        f'en {city}. Solución seleccionada: '
        f'{rec.get("product_name") or rec.get("product_id","–")} '
        f'— fiabilidad {rel1:.1f}% (año 1)'
        + (f', {rel10:.1f}% (año 10).' if rel10 is not None else '.')
    )

    if monthly:
        _h2(doc, 'Producción solar vs Consumo estimado')
        _img(doc, _chart_prod_cons(monthly, rec.get('product_name','')), 6.2)
        _h2(doc, 'Estado de carga de batería')
        _img(doc, _chart_soc(monthly, rec.get('product_name','')), 6.2)

    _h2(doc, 'Indicadores clave de la solución recomendada')
    kpi_rows = [
        ('Panel fotovoltaico',   f'{rec.get("pv_peak_power_wp","–")} Wp'),
        ('Batería',              f'{rec.get("battery_nominal_wh","–")} Wh'),
        ('Peso estimado',        f'{rec.get("weight_kg","–")} kg'),
        ('Fiabilidad año 1',     f'{rel1:.1f}%'),
        ('Fiabilidad año 10',    f'{rel10:.1f}%' if rel10 else '–'),
        ('TCO 10 años (venta)',  _fmt_eur(rec.get('tco_10y_sale'))),
        ('TCO 10 años (coste)',  _fmt_eur(rec.get('tco_10y_cost'))),
        ('CO₂ evitado 10 años',  f'{int(round(rec["co2_saved_10y_kg"]))} kg' if rec.get('co2_saved_10y_kg') else '–'),
        ('Autonomía sin sol',    f'{rec.get("autonomy_days","–")} días'),
    ]
    tbl = doc.add_table(rows=len(kpi_rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (row, (k, v)) in enumerate(zip(tbl.rows, kpi_rows)):
        _cell_text(row.cells[0], k, size=9, bold=True, color_key='grey')
        _cell_text(row.cells[1], v, size=9)
        _set_cell_bg(row.cells[0], T['surface'].lstrip('#'))
        _set_cell_bg(row.cells[1], T['cream'].lstrip('#') if i % 2 else T['white'].lstrip('#'))


def _build_recommended_detail(doc, rec, all_candidates, project, photometry, nightProfile):
    """Sección 02: justificación y árbol de pérdidas."""
    if not rec: return
    doc.add_page_break()
    product_name = rec.get('product_name') or rec.get('product_id','–')
    geo  = rec.get('geometry_type','')
    tech = _tech(geo)

    _h1(doc, f'03 · Solución recomendada — {product_name}')
    _add_hr(doc, 'line')

    # AI narrative
    narrative = _narrative_recommendation(rec, all_candidates, project, photometry)
    for para in narrative.split('\n\n'):
        if para.strip():
            _body(doc, para.strip())

    # Technology
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _h2(doc, f'Tecnología: {tech["title"]}')
    _body(doc, tech['body'])

    # Loss tree
    lt = rec.get('loss_tree')
    if lt:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        _h2(doc, 'Árbol de pérdidas del sistema')
        _body(doc,
            'Las barras muestran la energía solar remanente acumulada tras cada etapa de pérdida, '
            'desde el recurso solar teórico hasta la energía útil disponible para la luminaria.',
            size=9, color_key='grey')
        _img(doc, _chart_loss_waterfall(lt), 6.2)

        losses_detail = [
            ('Geometría / tilt',       lt.get('geometry_loss_pct', 0),
             'Ángulo de inclinación subóptimo respecto a irradiancia máxima teórica.'),
            ('Suciedad (soiling)',      lt.get('soiling_loss_pct', 0),
             'Depósito de polvo y partículas en la superficie del panel.'),
            ('Temperatura',            lt.get('temperature_loss_pct', 0),
             'Reducción de eficiencia por temperatura de célula > 25°C (NOCT).'),
            ('Regulador / cableado',   lt.get('controller_loss_pct', 0),
             'Pérdidas eléctricas en regulador MPPT, cableado y conectores.'),
            ('Batería (ida+vuelta)',    lt.get('battery_roundtrip_loss_pct', 0),
             'Pérdida de rendimiento en ciclo carga/descarga LiFePO4.'),
            ('Degradación promedio',   lt.get('degradation_loss_pct', 0),
             'Promedio degradación panel (~0.5%/año) y batería (~2%/año).'),
        ]
        running = 100.0
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        _table_header(tbl, ['Factor de pérdida', 'Pérdida', 'Restante', 'Descripción'])
        for i, (lbl, val, desc) in enumerate(losses_detail):
            running = max(0, running - val)
            row = tbl.add_row()
            _cell_text(row.cells[0], lbl, size=8.5)
            _cell_text(row.cells[1], f'−{val:.1f}%', size=8.5, color_key='warning',
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[2], f'{running:.1f}%', size=8.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[3], desc, size=8, color_key='grey')
            for cell in row.cells:
                _set_cell_bg(cell, T['surface'].lstrip('#') if i % 2 else T['white'].lstrip('#'))
        row = tbl.add_row()
        _cell_text(row.cells[0], '▶ Energía útil disponible', size=8.5, bold=True)
        _cell_text(row.cells[1], '', size=8.5)
        avail = lt.get('energy_available_pct', running)
        _cell_text(row.cells[2], f'{avail:.1f}%', size=9, bold=True, color_key='success',
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row.cells[3], 'Energía solar efectiva entregada al sistema.', size=8, color_key='grey')
        for cell in row.cells:
            _set_cell_bg(cell, T['cream'].lstrip('#'))


def _build_shading_correction(doc, rec):
    """Sección 04: corrección de sombras locales (Shadowmap), solo si se aplicó."""
    shading = rec.get('shading') if rec else None
    if not shading or shading.get('status') == 'not_supported_geometry':
        return
    doc.add_page_break()
    _h1(doc, '04 · Corrección de sombras locales')
    _add_hr(doc, 'line')

    _body(doc,
        'La irradiancia climática base se ha calculado mediante PVGIS/JRC. Cuando ha existido '
        'información 3D local disponible, se ha aplicado una corrección de sombra mediante '
        'Shadowmap para ajustar la componente directa de irradiancia, considerando obstáculos '
        'locales como edificios, terreno y vegetación. La componente difusa no se ha anulado '
        'automáticamente, ya que incluso en sombra directa existe captación por irradiancia difusa.')

    applied = bool(shading.get('shadow_correction_applied'))
    if not applied:
        warn = (shading.get('warnings') or ['No se aplicó corrección de sombra local.'])[0]
        _body(doc, f'⚠ {warn}', color_key='warning')
        return

    height_label = 'centro del panel' if shading.get('height_mode') == 'panel_center_height' else 'proxy a nivel de suelo'
    confidence = shading.get('confidence')

    _h2(doc, 'Parámetros de la corrección')
    rows = [
        ('Fuente de irradiancia base', 'PVGIS/JRC'),
        ('Fuente de sombra local',     shading.get('provider', 'Shadowmap')),
        ('Modo de cálculo',            shading.get('mode', '–')),
        ('Confianza',                  f'{confidence*100:.0f}%' if confidence is not None else '–'),
        ('Altura considerada',         height_label),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    for i, (row, (k, v)) in enumerate(zip(tbl.rows, rows)):
        _cell_text(row.cells[0], k, size=9, bold=True, color_key='grey')
        _cell_text(row.cells[1], v, size=9)
        _set_cell_bg(row.cells[0], T['surface'].lstrip('#'))
        _set_cell_bg(row.cells[1], T['cream'].lstrip('#') if i % 2 else T['white'].lstrip('#'))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _h2(doc, 'Pérdida por sombra local')
    loss_rows = [
        ('Pérdida directa anual',  shading.get('annual_direct_shadow_loss_pct')),
        ('Pérdida total anual',    shading.get('annual_total_shadow_loss_pct')),
        ('Pérdida en mes crítico', shading.get('critical_month_shadow_loss_pct')),
    ]
    tbl2 = doc.add_table(rows=1, cols=len(loss_rows))
    tbl2.style = 'Table Grid'
    _table_header(tbl2, [k for k, _ in loss_rows])
    row = tbl2.add_row()
    for cell, (_, v) in zip(row.cells, loss_rows):
        _cell_text(cell, f'{v:.1f}%' if v is not None else '–', size=10, bold=True,
                    color_key='warning', align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(cell, T['cream'].lstrip('#'))

    comparison = rec.get('shading_comparison')
    if comparison:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        _h2(doc, 'Comparativa PVGIS puro vs. PVGIS + Shadowmap')
        base = comparison.get('base_case', {})
        corr = comparison.get('corrected_case', {})
        tbl3 = doc.add_table(rows=1, cols=3)
        tbl3.style = 'Table Grid'
        _table_header(tbl3, ['Escenario', 'Producción anual', 'Fiabilidad año 1'])
        for i, c in enumerate((base, corr)):
            row = tbl3.add_row()
            rel = 100 - (c.get('annual_failure_rate_pct') or 0)
            for cell, val in zip(row.cells, [
                c.get('label', '–'),
                f"{c.get('annual_production_kwh', 0):.1f} kWh",
                f'{rel:.1f}%',
            ]):
                _cell_text(cell, val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_bg(cell, T['surface'].lstrip('#') if i % 2 else T['white'].lstrip('#'))

    warnings = shading.get('warnings') or []
    if warnings:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        for w in warnings:
            _body(doc, f'⚠ {w}', size=9, color_key='warning')


def _build_smartec(doc, nightProfile, photometry, rec):
    """Sección 05: gestión predictiva Smartec."""
    doc.add_page_break()
    _h1(doc, '05 · Gestión predictiva Smartec')
    _add_hr(doc, 'line')

    _quote(doc,
        'Con SMARTEC, la energía se administra como un recurso valioso. No se desperdicia: '
        'se utiliza cuando realmente aporta valor.')

    periods = nightProfile.get('periods', [])
    narrative = _narrative_smartec(periods, photometry, rec)
    for para in narrative.split('\n\n'):
        if para.strip():
            _body(doc, para.strip())

    sys_pw = photometry.get('system_power_w', 0)
    if periods:
        png_dim = _chart_smartec_dimming(periods, sys_pw)
        if png_dim:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            _img(doc, png_dim, 6.0)

        _h2(doc, 'Detalle de períodos nocturnos')
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = 'Table Grid'
        _table_header(tbl, ['Período', 'Duración', 'Presencia', 'Dim. pres.', 'Dim. sin pres.', 'Consumo est.'])
        for i, p in enumerate(periods):
            dur_h   = p.get('duration_pct', 0) * 12
            pres    = p.get('presence_ratio', 0)
            dp      = p.get('dimming_presence', 1.0)
            dnp     = p.get('dimming_no_presence', 1.0)
            avg_dim = pres * dp + (1 - pres) * dnp
            wh      = sys_pw * avg_dim * dur_h
            row = tbl.add_row()
            for cell, val in zip(row.cells, [
                f'Período {i+1}', f'{p.get("duration_pct",0)*100:.0f}%  (~{dur_h:.1f}h)',
                f'{pres*100:.0f}%', f'{dp*100:.0f}%', f'{dnp*100:.0f}%', f'{wh:.0f} Wh',
            ]):
                _cell_text(cell, val, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_bg(cell, T['surface'].lstrip('#') if i % 2 else T['white'].lstrip('#'))

    mo  = nightProfile.get('margin_on_min', -15)
    mof = nightProfile.get('margin_off_min', 15)
    _h2(doc, 'Horarios de encendido / apagado')
    _body(doc,
        f'Encendido: {mo:+d} min respecto al ocaso astronómico local.  '
        f'Apagado: +{mof} min respecto al alba astronómica local.',
        size=9.5)


def _build_comparative(doc, all_candidates, photometry):
    """Sección 06: comparativa completa de todas las soluciones."""
    doc.add_page_break()
    _h1(doc, '07 · Comparativa de soluciones')
    _add_hr(doc, 'line')

    valid  = [c for c in all_candidates if not c.get('error')]
    errors = [c for c in all_candidates if c.get('error')]
    _body(doc,
        f'{len(all_candidates)} soluciones evaluadas — {len(valid)} completadas'
        + (f', {len(errors)} con error.' if errors else '.'))

    if valid:
        cols = ['Producto', 'Panel\nWp', 'Bat.\nWh', 'Peso\nkg',
                'Fiab.\naño 1', 'Fiab.\naño 10', 'TCO 10a\n(venta)', 'TCO 10a\n(coste)', 'CO₂\n10a kg', 'Escenario']
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = 'Table Grid'
        _table_header(tbl, cols)
        for i, c in enumerate(valid):
            rel1      = 100 - (c.get('annual_failure_rate_pct') or 0)
            rel10_raw = c.get('annual_failure_rate_pct_y10')
            rel10     = (100 - rel10_raw) if rel10_raw is not None else None
            tags      = c.get('scenario_tags') or (['recommended'] if c.get('recommended') else [])
            tag_str   = ', '.join({
                'recommended': '★ Recomendada',
                'low_capex':   '↓ Bajo CAPEX',
                'max_reliability': '▲ Máx. Fiab.',
                'hybrid':      '⚡ Híbrida',
            }.get(t, t) for t in tags)
            row = tbl.add_row()
            for cell, val in zip(row.cells, [
                c.get('product_name') or c.get('product_id','–'),
                c.get('pv_peak_power_wp','–'), c.get('battery_nominal_wh','–'), c.get('weight_kg','–'),
                f'{rel1:.1f}%', f'{rel10:.1f}%' if rel10 is not None else '–',
                _fmt_eur(c.get('tco_10y_sale')), _fmt_eur(c.get('tco_10y_cost')),
                f'{round(c["co2_saved_10y_kg"])} kg' if c.get('co2_saved_10y_kg') else '–',
                tag_str or '–',
            ]):
                _cell_text(cell, val, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_bg(cell, T['cream'].lstrip('#') if c.get('recommended') else
                             (T['surface'].lstrip('#') if i % 2 == 1 else T['white'].lstrip('#')))

    tco_png = _chart_tco_comparison(all_candidates)
    if tco_png:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        _img(doc, tco_png, 6.2)

    # Monthly table for recommended
    rec = next((c for c in valid if c.get('recommended')), valid[0] if valid else None)
    if rec and rec.get('monthly_data'):
        _h2(doc, f'Datos mensuales — {rec.get("product_name") or rec.get("product_id","–")}')
        _body(doc, 'Producción = kWh/día media del mes. Consumo = kWh/noche media. SOC mínimo = peor noche del mes.',
              size=8.5, color_key='grey')
        monthly = rec['monthly_data']
        tbl2 = doc.add_table(rows=1, cols=7)
        tbl2.style = 'Table Grid'
        _table_header(tbl2, ['Mes', 'Prod. kWh/día', 'Cons. kWh/noche', 'Balance', 'SOC mín. a1', 'SOC mín. a10', 'Noches críticas'])
        for i, m in enumerate(monthly[:12]):
            prod = m.get('production_kwh') or 0
            cons = m.get('consumption_kwh') or 0
            bal  = prod - cons
            soc1 = m.get('soc_min_pct') or 0
            fails = m.get('failures') or 0
            row  = tbl2.add_row()
            soc_ck = 'success' if soc1 >= 30 else ('warning' if soc1 >= 15 else 'danger')
            for j, (cell, val, ck) in enumerate(zip(row.cells, [
                MONTHS_ES[i] if i < 12 else str(i+1),
                f'{prod:.3f}', f'{cons:.3f}',
                f'+{bal:.3f}' if bal >= 0 else f'{bal:.3f}',
                f'{soc1:.0f}%',
                f'{(m.get("soc_min_y10") or 0):.0f}%',
                str(fails) if fails else '–',
            ], [
                'black', 'black', 'black',
                'success' if bal >= 0 else 'danger',
                soc_ck, 'muted',
                'danger' if fails else 'muted',
            ])):
                _cell_text(cell, val, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                           color_key=ck, bold=(j in (3, 4)))
                _set_cell_bg(cell, T['surface'].lstrip('#') if i % 2 else T['white'].lstrip('#'))


def _build_economic(doc, all_candidates, env):
    """Sección 07: análisis económico y ambiental."""
    doc.add_page_break()
    _h1(doc, '08 · Análisis económico y ambiental')
    _add_hr(doc, 'line')

    elec_cost  = env.get('electricity_cost', 0.20)
    co2_factor = env.get('country_co2_factor', 0.25)

    _h2(doc, 'Coste total de propiedad (TCO) a 10 años')
    _body(doc,
        f'El TCO incluye el coste de adquisición, los costes de mantenimiento preventivo y '
        f'correctivo estimados, y el coste de oportunidad energética frente a la conexión a '
        f'red eléctrica a {elec_cost:.2f} €/kWh. La amortización considera la vida útil del '
        f'panel fotovoltaico (>25 años) y de la batería LiFePO4 (8–12 años con gestión Smartec).')

    valid = [c for c in all_candidates if not c.get('error') and c.get('tco_10y_sale')]
    if valid:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = 'Table Grid'
        _table_header(tbl, ['Producto', 'TCO venta', 'TCO coste', 'Margen', 'CO₂ 10a'])
        for i, c in enumerate(valid):
            tco_s = c.get('tco_10y_sale') or 0
            tco_c = c.get('tco_10y_cost') or 0
            margin_pct = (tco_s - tco_c) / tco_s * 100 if tco_s else 0
            row = tbl.add_row()
            for cell, val in zip(row.cells, [
                c.get('product_name') or c.get('product_id','–'),
                _fmt_eur(tco_s), _fmt_eur(tco_c),
                f'{margin_pct:.1f}%',
                f'{round(c["co2_saved_10y_kg"])} kg' if c.get('co2_saved_10y_kg') else '–',
            ]):
                _cell_text(cell, val, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_bg(cell, T['cream'].lstrip('#') if c.get('recommended') else
                             (T['surface'].lstrip('#') if i % 2 else T['white'].lstrip('#')))

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    _h2(doc, 'Impacto medioambiental')
    rec = next((c for c in all_candidates if c.get('recommended') and not c.get('error')), None)
    co2_total = rec.get('co2_saved_10y_kg') if rec else None
    _body(doc,
        f'La sustitución de puntos de luz convencionales conectados a red por sistemas '
        f'autónomos SALVI Solar evita'
        + (f' aproximadamente {co2_total:.0f} kg de CO₂ equivalente en 10 años '
           f'(factor de red: {co2_factor:.2f} kg CO₂/kWh). ' if co2_total else ' emisiones de CO₂ a la red eléctrica. ')
        + 'La generación fotovoltaica in-situ elimina las pérdidas de transporte y distribución '
        'de la red eléctrica (~8–12%), y contribuye a la resiliencia energética local ante '
        'cortes de suministro.')


def _build_methodology(doc, project):
    """Sección 08: metodología y fuentes de datos."""
    doc.add_page_break()
    _h1(doc, '09 · Metodología y fuentes de datos')
    _add_hr(doc, 'line')

    _h2(doc, 'Fuente de datos de irradiancia solar')
    _body(doc,
        'Los datos de irradiancia solar proceden de PVGIS (Photovoltaic Geographical '
        'Information System), desarrollado por el Joint Research Centre (JRC) de la Comisión '
        'Europea. Se utiliza el año meteorológico típico (TMY) de la base de datos PVGIS-SARAH3 '
        'para Europa, África y Asia occidental, o ERA5 para otras regiones. '
        'La resolución espacial es de aproximadamente 5 km × 5 km.')

    _h2(doc, 'Modelo de simulación energética')
    _body(doc,
        'El cálculo de producción energética emplea el modelo de irradiancia horaria de '
        'PVGIS a 1 kWp de potencia de referencia, escalando linealmente por la potencia '
        'nominal del sistema. Las pérdidas del árbol de pérdidas se aplican secuencialmente '
        'sobre la producción horaria bruta. La simulación de la batería realiza un balance '
        'energético horario para las 8.760 horas del año, determinando el SOC en cada instante '
        'con el perfil nocturno Smartec como carga variable.')

    _h2(doc, 'Modelo de degradación')
    _body(doc,
        'La simulación del año 10 aplica una degradación del panel fotovoltaico del 0.5% anual '
        'acumulado (×0.95 en año 10) y una degradación de la capacidad de la batería LiFePO4 '
        'del 2% anual acumulado (×0.82 en año 10). Valores conservadores respecto a las '
        'garantías de fabricante típicas del mercado.')

    _h2(doc, 'Limitaciones')
    _body(doc,
        'Este informe es una estimación basada en datos históricos meteorológicos y modelos '
        'estadísticos. Los resultados reales pueden variar por condiciones locales específicas '
        '(obstrucciones, microclima, orientación exacta). Se recomienda una visita de campo '
        'para validar los parámetros de instalación antes de la oferta definitiva.',
        size=9.5, italic=True, color_key='muted')


# ── Main entry point ───────────────────────────────────────────────────────────

def generate_report(data: dict) -> bytes:
    project    = data.get('project', {})
    photometry = data.get('photometry', {})
    night      = data.get('nightProfile', {})
    env        = data.get('env', {})
    sim        = data.get('simulation', {})
    candidates = sim.get('candidates', [])

    rec = next((c for c in candidates if c.get('recommended') and not c.get('error')), None)
    if not rec and candidates:
        rec = next((c for c in candidates if not c.get('error')), None)

    sim_date = datetime.now().strftime('%d/%m/%Y')

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Set default style font to Calibri
    doc.styles['Normal'].font.name = FONT_UI

    _build_cover(doc, project, rec, sim_date)
    doc.add_page_break()
    _build_value_proposition(doc, rec, night, photometry)
    _build_executive_summary(doc, rec, candidates, project, photometry)
    _build_recommended_detail(doc, rec, candidates, project, photometry, night)
    _build_shading_correction(doc, rec)
    _build_smartec(doc, night, photometry, rec)
    _build_municipal_benefits(doc)
    _build_comparative(doc, candidates, photometry)
    _build_economic(doc, candidates, env)
    _build_methodology(doc, project)

    _add_hr(doc, 'line')
    _quote(doc,
        'La verdadera pregunta es: ¿quién gestiona la energía cada noche para garantizar que '
        'la luz llegue hasta el amanecer? La respuesta de Salvi Lighting es SIL con SMARTEC.',
        attribution='Light inspired by you.')
    _body(doc,
        f'SALVI Solar Studio  ·  {sim_date}  ·  Datos solares: PVGIS/JRC – Comisión Europea  ·  Documento confidencial.',
        size=8, italic=True, color_key='muted', space_after=0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
