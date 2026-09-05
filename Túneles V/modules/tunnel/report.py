"""
SALVI Tunnel Engine — Generador de Informe Técnico Word
Produce un .docx con el resultado completo del cálculo CIE 88:2004.
Requiere: python-docx
"""

import io
import datetime
import json
import math
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from urllib.parse import urlencode, urlparse
from urllib.request import Request, urlopen
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ══════════════════════════════════════════════════════════════════
# CONSTANTES DE ESTILO
# ══════════════════════════════════════════════════════════════════

BLUE_DARK   = "1A3A6B"   # Cabeceras principales
BLUE_MED    = "1A56B0"   # Cabeceras de tabla
BLUE_LIGHT  = "D5E8F5"   # Filas alternas de tabla
BLUE_BRAND  = "0066CC"   # Textos destacados
GRAY_LIGHT  = "F5F5F5"   # Fondo alternativo gris
WHITE       = "FFFFFF"
BLACK       = "000000"

# Nombres internos de las familias OpenType instaladas en el puesto SALVI.
BODY_FONT   = "Exposure[+10]"
TITLE_FONT  = "Exposure[-50]"

_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_APHEX_ASSET_DIR = _PROJECT_ROOT / "assets" / "productos" / "aphex"
_APHEX_CATALOG = {
    "APHEX_S": {
        "product": "Aphex S Lira", "pdf": "APHEX S-ES.pdf", "dimensions_mm": (472, 400, 90),
        "mounting": "4 a 25 m", "catalog_power_w": 223, "catalog_flux_lm": 39061,
        "weight": "12 kg", "protection": "IP66 / IK08 / IK09 / IK10", "leds": "8 a 50 LED",
    },
    "APHEX_M": {
        "product": "Aphex M Lira", "pdf": "APHEX M-ES.pdf", "dimensions_mm": (656, 412, 90),
        "mounting": "4 a 25 m", "catalog_power_w": 446, "catalog_flux_lm": 78123,
        "weight": "16 kg", "protection": "IP66 / IK08 / IK09", "leds": "32 o 100 LED",
    },
    "APHEX_L": {
        "product": "Aphex L Lira", "pdf": "APHEX L-ES.pdf", "dimensions_mm": (656, 558, 90),
        "mounting": "4 a 25 m", "catalog_power_w": 670, "catalog_flux_lm": 117184,
        "weight": "20 kg", "protection": "IP66 / IK08 / IK09", "leds": "48 o 150 LED",
    },
}

# ══════════════════════════════════════════════════════════════════
# UTILIDADES DOCX
# ══════════════════════════════════════════════════════════════════

def _set_word_font(font, family: str) -> None:
    """Fija las cuatro familias OOXML para evitar sustituciones de Word."""
    font.name = family
    rpr = font._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attribute}"), family)


def _configure_word_typography(doc: Document) -> None:
    """Aplica Exposure +10 al cuerpo y Exposure -50 a encabezados."""
    _set_word_font(doc.styles["Normal"].font, BODY_FONT)
    doc.styles["Normal"].font.size = Pt(10)
    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        try:
            _set_word_font(doc.styles[style_name].font, TITLE_FONT)
        except KeyError:
            continue

def _set_cell_bg(cell, hex_color: str) -> None:
    """Establece el color de fondo de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    # Remove existing shd if present
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold: bool = False, color: str = None,
               size_pt: float = 9.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Escribe texto en una celda con formato."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_separator(doc: Document, color: str = BLUE_MED) -> None:
    """Línea horizontal separadora."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Inserta un vínculo externo compatible con Word sin depender de estilos."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE_BRAND)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _section_heading(doc: Document, title: str, level: int = 1) -> None:
    """Encabezado de sección con estilo SALVI."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    _set_word_font(run.font, TITLE_FONT)
    run.font.size = Pt(12 if level == 1 else 10.5)
    run.font.color.rgb = RGBColor.from_string(BLUE_DARK if level == 1 else BLUE_MED)
    _add_separator(doc, BLUE_DARK if level == 1 else BLUE_MED)


def _kv_table(doc: Document, rows: list, col_widths=(5.5, 8.5)) -> None:
    """Tabla de dos columnas clave→valor."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    for i, (key, val) in enumerate(rows):
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        c0, c1 = table.rows[i].cells
        _set_cell_bg(c0, bg)
        _set_cell_bg(c1, bg)
        _cell_text(c0, key, bold=True, size_pt=9)
        _cell_text(c1, str(val), size_pt=9)
        # Column widths
        c0.width = Cm(col_widths[0])
        c1.width = Cm(col_widths[1])


def _header_row(row, headers: list, bg: str = BLUE_MED) -> None:
    """Fila de cabecera de tabla (fondo azul, texto blanco)."""
    trPr = row._tr.get_or_add_trPr()
    repeat = OxmlElement('w:tblHeader')
    repeat.set(qn('w:val'), 'true')
    trPr.append(repeat)
    for cell, h in zip(row.cells, headers):
        _set_cell_bg(cell, bg)
        _cell_text(cell, h, bold=True, color=WHITE,
                   size_pt=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        for run in cell.paragraphs[0].runs:
            _set_word_font(run.font, TITLE_FONT)


def _set_page_margins(doc: Document,
                      top=2.0, bottom=2.0, left=2.5, right=2.0) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════

def _project_route(params: dict) -> list[tuple[float, float]]:
    """Eje OSM persistido en el formulario, o el portal como último recurso."""
    params = params or {}
    route = []
    raw_route = params.get("osm_tunnel_geometry") or []
    if isinstance(raw_route, list):
        for point in raw_route:
            if not isinstance(point, dict):
                continue
            try:
                lat = float(point.get("lat"))
                lng = float(point.get("lng", point.get("lon")))
                if -85.0 < lat < 85.0 and -180.0 <= lng <= 180.0:
                    route.append((lat, lng))
            except (TypeError, ValueError):
                continue
    if route:
        return route
    try:
        lat, lng = float(params.get("lat")), float(params.get("lng"))
        return [(lat, lng)] if -85.0 < lat < 85.0 and -180.0 <= lng <= 180.0 else []
    except (TypeError, ValueError):
        return []


def _satellite_location_map(params: dict) -> Optional[io.BytesIO]:
    """Crea un recorte satelital reproducible con el eje OSM del túnel.

    La imagen no se guarda en disco: se descarga sólo al generar el informe y
    queda embebida en el DOCX. Si el proyecto no tiene coordenadas o no hay
    conexión, la portada sigue siendo válida y no se inventa una localización.
    """
    params = params or {}
    cached_image = params.get("_report_satellite_map")
    if isinstance(cached_image, bytes):
        return io.BytesIO(cached_image)
    route = _project_route(params)
    if not route:
        return None

    try:
        from PIL import Image, ImageDraw, ImageFont

        latitudes = [point[0] for point in route]
        longitudes = [point[1] for point in route]
        center_lat = (min(latitudes) + max(latitudes)) / 2
        center_lng = (min(longitudes) + max(longitudes)) / 2
        lat_span_m = max(latitudes) - min(latitudes)
        lat_span_m *= 111_320
        lng_span_m = (max(longitudes) - min(longitudes)) * 111_320 * math.cos(math.radians(center_lat))
        extent_m = max(lat_span_m, lng_span_m, float(params.get("length_m", 0) or 0))
        zoom = 16 if extent_m < 450 else (15 if extent_m < 1_200 else 14)
        tiles = 2 ** zoom

        def world_px(latitude: float, longitude: float) -> tuple[float, float]:
            x = (longitude + 180.0) / 360.0 * tiles * 256
            latitude = max(min(latitude, 85.05112878), -85.05112878)
            y = (1 - math.asinh(math.tan(math.radians(latitude))) / math.pi) / 2 * tiles * 256
            return x, y

        center_x, center_y = world_px(center_lat, center_lng)
        # Con 3×3 teselas, 512 px permite recortar siempre alrededor del
        # centro sin salir del mosaico (y mantiene una escala nítida en A4).
        image_size = 512
        top_left_x, top_left_y = center_x - image_size / 2, center_y - image_size / 2
        tile_x0, tile_y0 = math.floor(top_left_x / 256), math.floor(top_left_y / 256)
        canvas = Image.new("RGB", (768, 768), "#D9E2F0")

        def get_tile(tile_x: int, tile_y: int):
            safe_x = tile_x % tiles
            safe_y = min(max(tile_y, 0), tiles - 1)
            url = (
                "https://server.arcgisonline.com/ArcGIS/rest/services/"
                f"World_Imagery/MapServer/tile/{zoom}/{safe_y}/{safe_x}"
            )
            request = Request(url, headers={"User-Agent": "SALVI-Tunnel-Engine/1.2"})
            with urlopen(request, timeout=4) as response:
                return tile_x, tile_y, Image.open(io.BytesIO(response.read())).convert("RGB")

        coordinates = [(tile_x0 + dx, tile_y0 + dy) for dy in range(3) for dx in range(3)]
        with ThreadPoolExecutor(max_workers=9) as executor:
            futures = [executor.submit(get_tile, tile_x, tile_y) for tile_x, tile_y in coordinates]
            for future in as_completed(futures):
                tile_x, tile_y, tile = future.result()
                canvas.paste(tile, ((tile_x - tile_x0) * 256, (tile_y - tile_y0) * 256))

        offset_x, offset_y = top_left_x - tile_x0 * 256, top_left_y - tile_y0 * 256
        image = canvas.crop((int(offset_x), int(offset_y), int(offset_x) + image_size, int(offset_y) + image_size))
        draw = ImageDraw.Draw(image)
        route_pixels = [(x - top_left_x, y - top_left_y) for x, y in (world_px(*point) for point in route)]
        if len(route_pixels) > 1:
            draw.line(route_pixels, fill="#172033", width=9, joint="curve")
            draw.line(route_pixels, fill="#FBBF24", width=5, joint="curve")
        markers = ((route_pixels[0], "#22C55E"), (route_pixels[-1], "#EF4444")) if len(route_pixels) > 1 else ((route_pixels[0], "#FBBF24"),)
        for (x, y), color in markers:
            draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill="white", outline="#172033", width=2)
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=color)
        # Panorámica horizontal para no sobrecargar la portada: conserva el
        # centro del eje y deja espacio para las fotos de las dos bocas.
        crop_top, crop_height = 141, 230
        image = image.crop((0, crop_top, image_size, crop_top + crop_height))
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, crop_height - 26, image_size, crop_height), fill=(20, 31, 49))
        map_note = "Imagen satelital © Esri · eje de túnel OSM · A entrada / B salida" if len(route) > 1 else "Imagen satelital © Esri · localización de portal"
        draw.text((9, crop_height - 20), map_note, fill="white")
        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        image_bytes = output.getvalue()
        # La misma imagen se utiliza en portada y en el plano de situación.
        # Se evita así una segunda descarga y se garantiza que ambas vistas
        # representan exactamente el mismo eje y los mismos portales.
        params["_report_satellite_map"] = image_bytes
        return io.BytesIO(image_bytes)
    except Exception:
        return None


def _report_video_url(params: dict) -> Optional[str]:
    """Acepta sólo URLs http(s) para no crear vínculos inseguros en el Word."""
    url = str((params or {}).get("report_video_url", "") or "").strip()
    try:
        parsed = urlparse(url)
    except ValueError:
        return None
    return url if parsed.scheme in {"http", "https"} and parsed.netloc else None


def _video_window_preview(params: dict, title: str) -> io.BytesIO:
    """Miniatura autónoma para el enlace de vídeo; utiliza el mapa si existe."""
    from PIL import Image, ImageDraw, ImageFont, ImageOps

    size = (960, 540)
    location_map = _satellite_location_map(params)
    if location_map:
        image = Image.open(location_map).convert("RGB")
        resampling = getattr(Image, "Resampling", Image).LANCZOS
        image = ImageOps.fit(image, size, method=resampling, centering=(0.5, 0.5))
    else:
        image = Image.new("RGB", size, "#172033")
        draw_background = ImageDraw.Draw(image)
        for line in range(0, size[0] + size[1], 45):
            draw_background.line((line, 0, 0, line), fill="#1F3D67", width=2)

    overlay = Image.new("RGBA", size, (8, 20, 38, 142))
    image = Image.alpha_composite(image.convert("RGBA"), overlay)
    draw = ImageDraw.Draw(image)
    # Arial está disponible en la instalación Windows de la aplicación. En
    # otros entornos se conserva una miniatura válida, convirtiendo sólo el
    # texto del bitmap a ASCII cuando la fuente por defecto no soporta tildes.
    font_paths = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    font_path = next((path for path in font_paths if path.exists()), None)
    if font_path:
        font_brand = ImageFont.truetype(str(font_path), 18)
        font_subtitle = ImageFont.truetype(str(font_path), 15)
        font_title = ImageFont.truetype(str(font_path), 18)
        image_text = lambda value: value
    else:
        font_brand = font_subtitle = font_title = ImageFont.load_default()
        image_text = lambda value: unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode()
    draw.rounded_rectangle((332, 132, 628, 408), radius=22,
                           fill="#1A56B0", outline="#FFFFFF", width=4)
    draw.polygon([(438, 202), (438, 338), (548, 270)], fill="#FFFFFF")
    draw.text((36, 35), "SALVI TUNNEL ENGINE", fill="#FFFFFF", font=font_brand)
    draw.text((36, 66), image_text("VÍDEO DE RECORRIDO Y VERIFICACIÓN VISUAL"),
              fill="#D5E8F5", font=font_subtitle)
    draw.rounded_rectangle((36, 462, 924, 510), radius=8, fill="#172033")
    draw.text((55, 476), image_text(title[:90]), fill="#FFFFFF", font=font_title)
    output = io.BytesIO()
    image.convert("RGB").save(output, format="PNG", optimize=True)
    output.seek(0)
    return output


def _build_report_video(doc: Document, params: dict) -> None:
    """Ventana de vídeo compatible: miniatura, vínculo y QR para móvil/PDF."""
    video_url = _report_video_url(params)
    if not video_url:
        return
    title = str(params.get("report_video_title") or "Recorrido virtual del túnel").strip()
    title = title or "Recorrido virtual del túnel"
    _section_heading(doc, "Vídeo de recorrido y verificación visual", level=2)
    intro = doc.add_paragraph()
    intro.add_run(
        "La ventana incorpora el acceso al vídeo de apoyo alojado en la plataforma del proyecto. "
        "El QR mantiene el acceso desde la versión PDF o impresa; el vídeo no sustituye "
        "las verificaciones ni las mediciones de recepción."
    ).font.size = Pt(8.5)
    try:
        import qrcode

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        preview_cell, qr_cell = table.rows[0].cells
        _set_cell_bg(preview_cell, "F5F8FC")
        _set_cell_bg(qr_cell, "F5F8FC")
        preview_cell.width = Cm(11.7)
        qr_cell.width = Cm(3.7)
        preview = _video_window_preview(params, title)
        preview_paragraph = preview_cell.paragraphs[0]
        preview_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preview_paragraph.add_run().add_picture(preview, width=Cm(10.9))
        preview_caption = preview_cell.add_paragraph()
        preview_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preview_caption.add_run("▶  ").bold = True
        _add_hyperlink(preview_caption, video_url, "Abrir vídeo de recorrido")

        qr = qrcode.QRCode(version=None, error_correction=qrcode.constants.ERROR_CORRECT_M,
                           box_size=7, border=2)
        qr.add_data(video_url)
        qr.make(fit=True)
        qr_image = io.BytesIO()
        qr.make_image(fill_color="#172033", back_color="white").convert("RGB").save(qr_image, format="PNG")
        qr_image.seek(0)
        qr_paragraph = qr_cell.paragraphs[0]
        qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_paragraph.add_run().add_picture(qr_image, width=Cm(3.0))
        qr_label = qr_cell.add_paragraph()
        qr_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_label.add_run("Escanear para abrir").italic = True
        qr_label.runs[0].font.size = Pt(7.5)
    except Exception:
        # La ventana conserva vínculo textual si el entorno de generación no
        # dispone de la librería QR o de una imagen de mapa.
        fallback = doc.add_paragraph()
        fallback.add_run("▶ ").bold = True
        fallback.add_run(title + ": ").bold = True
        _add_hyperlink(fallback, video_url, "Abrir vídeo")

    link = doc.add_paragraph()
    link.paragraph_format.space_before = Pt(3)
    link.add_run("Enlace de vídeo: ").bold = True
    _add_hyperlink(link, video_url, video_url)


def _portal_street_images(params: dict) -> list[tuple[str, io.BytesIO, str]]:
    """Obtiene vistas viales próximas a las bocas desde KartaView, sin API key."""
    route = _project_route(params)
    if not route:
        return []
    portals = [("Boca A", route[0]), ("Boca B", route[-1])] if len(route) > 1 else [("Boca", route[0])]

    def bearing(origin: tuple[float, float], destination: tuple[float, float]) -> float:
        lat1, lon1, lat2, lon2 = map(math.radians, (*origin, *destination))
        y = math.sin(lon2 - lon1) * math.cos(lat2)
        x = math.cos(lat1) * math.sin(lat2) - math.sin(lat1) * math.cos(lat2) * math.cos(lon2 - lon1)
        return (math.degrees(math.atan2(y, x)) + 360) % 360

    def distance_m(origin: tuple[float, float], destination: tuple[float, float]) -> float:
        dy = (origin[0] - destination[0]) * 111_320
        dx = (origin[1] - destination[1]) * 111_320 * math.cos(math.radians(origin[0]))
        return math.hypot(dx, dy)

    images = []
    for index, (label, portal) in enumerate(portals):
        target_heading = bearing(portal, route[1] if index == 0 else route[-2]) if len(route) > 1 else None
        try:
            query = urlencode({
                "lat": f"{portal[0]:.6f}", "lng": f"{portal[1]:.6f}",
                "zoomLevel": 16, "join": "sequence", "orderBy": "id", "orderDirection": "desc",
            })
            request = Request(f"https://api.openstreetcam.org/2.0/photo/?{query}", headers={"User-Agent": "SALVI-Tunnel-Engine/1.2"})
            with urlopen(request, timeout=5) as response:
                candidates = json.loads(response.read()).get("result", {}).get("data", [])

            def score(candidate: dict) -> float:
                try:
                    point = (float(candidate.get("lat")), float(candidate.get("lng")))
                    proximity = distance_m(portal, point)
                except (TypeError, ValueError):
                    return float("inf")
                if target_heading is None:
                    return proximity
                try:
                    heading = float(candidate.get("heading", candidate.get("headers"))) % 360
                    heading_gap = abs((heading - target_heading + 180) % 360 - 180)
                except (TypeError, ValueError):
                    heading_gap = 90
                return proximity + heading_gap * 0.45

            candidate = min(candidates, key=score, default=None)
            if not candidate or score(candidate) > 250:
                continue
            image_url = candidate.get("imageProcUrl") or candidate.get("fileurlProc")
            if not image_url:
                continue
            with urlopen(Request(image_url, headers={"User-Agent": "SALVI-Tunnel-Engine/1.2"}), timeout=8) as response:
                image_bytes = response.read()
            from PIL import Image
            image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            image.thumbnail((700, 390), Image.Resampling.LANCZOS)
            output = io.BytesIO()
            image.save(output, format="JPEG", quality=88, optimize=True)
            output.seek(0)
            captured = str(candidate.get("dateAdded", "")).split(" ")[0] or "fecha no disponible"
            images.append((label, output, f"Imagen vial © KartaView · captura {captured}"))
        except Exception:
            continue
    return images


def _build_cover(doc: Document, result: dict, params: dict) -> None:
    summary        = result.get('summary', {})
    classification = result.get('classification', {})

    # Cabecera azul
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("SALVI TUNNEL ENGINE")
    run.bold = True
    _set_word_font(run.font, TITLE_FONT)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(BLUE_DARK)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("Informe Técnico de Iluminación de Túneles")
    _set_word_font(run2.font, TITLE_FONT)
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor.from_string(BLUE_MED)

    _add_separator(doc, BLUE_DARK)

    # Subtítulo norma
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    r3 = p3.add_run("Norma: CIE 88:2004 — Guide for the Lighting of Road Tunnels and Underpasses")
    _set_word_font(r3.font, BODY_FONT)
    r3.italic = True
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()  # Espacio

    # Tabla de proyecto
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    fields = [
        ("Proyecto",              summary.get('project', params.get('project_name', '—'))),
        ("Tubo",                  summary.get('tube_id', '—')),
        ("Longitud total",        f"{summary.get('length_m', '—')} m"),
        ("Velocidad de diseño",   f"{summary.get('speed_kmh', '—')} km/h"),
        ("Tráfico de diseño",     f"{params.get('traffic_veh_h', '—')} veh/h"),
        ("Clasificación óptica",  classification.get('optical', '—').replace('_', ' ')),
        ("Iluminación diurna",    classification.get('daylighting', '—')),
        ("Fecha del informe",     datetime.date.today().strftime("%d/%m/%Y")),
    ]
    for i, (k, v) in enumerate(fields):
        bg = BLUE_LIGHT if i % 2 == 0 else WHITE
        c0, c1 = table.rows[i].cells
        _set_cell_bg(c0, bg)
        _set_cell_bg(c1, bg)
        _cell_text(c0, k, bold=True, size_pt=10)
        _cell_text(c1, str(v), size_pt=10)
        c0.width = Cm(6)
        c1.width = Cm(8)

    satellite_map = _satellite_location_map(params)
    if satellite_map:
        doc.add_paragraph()
        map_paragraph = doc.add_paragraph()
        map_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        map_paragraph.add_run().add_picture(satellite_map, width=Cm(13.8))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run("Situación del túnel, eje de cálculo y portales A/B")
        caption_run.italic = True
        caption_run.font.size = Pt(8)
        caption_run.font.color.rgb = RGBColor.from_string("555555")

    street_images = _portal_street_images(params)
    if street_images:
        doc.add_paragraph()
        photos = doc.add_table(rows=1, cols=len(street_images))
        photos.autofit = False
        for cell, (label, image, source) in zip(photos.rows[0].cells, street_images):
            cell.text = ""
            title = cell.paragraphs[0]
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading = title.add_run(label)
            heading.bold = True
            heading.font.size = Pt(8.5)
            image_paragraph = cell.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(
                image, width=Cm(6.45 if len(street_images) > 1 else 13.2)
            )
            source_paragraph = cell.add_paragraph()
            source_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source_run = source_paragraph.add_run(source)
            source_run.italic = True
            source_run.font.size = Pt(6.5)
            source_run.font.color.rgb = RGBColor.from_string("666666")

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    rn = p_note.add_run(
        "Redactado con SALVI Tunnel Engine (STE) — Motor de cálculo conforme a CIE 88:2004."
    )
    rn.font.size = Pt(8)
    rn.font.color.rgb = RGBColor.from_string("888888")
    rn.italic = True


# ══════════════════════════════════════════════════════════════════
# SECCIONES DEL INFORME
# ══════════════════════════════════════════════════════════════════

def _build_classification(doc: Document, result: dict,
                          title: str = "1. Clasificación del Túnel (CIE 88:2004 §6)") -> None:
    cls   = result.get('classification', {})
    speed = result.get('speed', {})
    _section_heading(doc, title)

    _kv_table(doc, [
        ("Clasificación geométrica",    cls.get('geometric', '—').replace('_', ' ')),
        ("Clasificación óptica",        cls.get('optical', '—').replace('_', ' ')),
        ("Necesidad de ilum. diurna",   cls.get('daylighting', '—')),
        ("Distancia de parada (SD)",    f"{speed.get('SD_m', '—')} m"),
        ("Distancia de reacción (dr)",  f"{speed.get('d_reaction_m', '—')} m"),
        ("Distancia de frenado (df)",   f"{speed.get('d_braking_m', '—')} m"),
    ])

    if cls.get('justification'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("Justificación: ")
        run.bold = True
        run.font.size = Pt(9)
        p.add_run(cls['justification']).font.size = Pt(9)


def _build_luminances(doc: Document, result: dict,
                      title: str = "2. Luminancias de Diseño (CIE 88:2004 §12–16)") -> None:
    s   = result.get('summary', {})
    l20 = result.get('l20', {})
    lth = result.get('lth', {})

    _section_heading(doc, title)

    _kv_table(doc, [
        ("L₂₀ — Luminancia campo 20°",     f"{s.get('L20', '—')} cd/m²"),
        ("Método de cálculo L₂₀",          l20.get('method', '—')),
        ("Confianza del valor L₂₀",         l20.get('confidence', '—')),
        ("Lth — Luminancia de umbral",      f"{s.get('Lth', '—')} cd/m²"),
        ("Factor k (Lth = k · L₂₀)",       s.get('k_factor', '—')),
        ("Coef. revelación contraste qc",   s.get('qc', '—')),
        ("Lseq — Luminancia equivalente",   f"{s.get('Lseq', '—')} cd/m²"),
        ("Lin — Luminancia interior",       f"{s.get('Lin', '—')} cd/m²"),
        (
            "L_noche normal",
            f"{s.get('L_night_normal', s.get('Lin', '—'))} cd/m²",
        ),
        (
            "L_noche reducida",
            f"{s.get('L_night_reduced', s.get('L_night', '—'))} cd/m²",
        ),
    ])

    if l20.get('note'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(f"Nota L₂₀: {l20['note']}")
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("666666")


def _build_zones(doc: Document, result: dict,
                 title: str = "3. Zonas Normativas CIE 88:2004") -> None:
    zones_raw = result.get('zones', {})
    _section_heading(doc, title)

    # Extraer lista de zonas del dict
    zones_list = []
    if isinstance(zones_raw, dict):
        # Puede venir como {threshold: {...}, transition: {...}, ...}
        # o como {zones: [...]}
        if 'zones' in zones_raw:
            zones_list = zones_raw['zones']
        else:
            for key, val in zones_raw.items():
                if isinstance(val, dict) and 's_start' in val:
                    zones_list.append(val)
    elif isinstance(zones_raw, list):
        zones_list = zones_raw

    if not zones_list:
        doc.add_paragraph("(Datos de zonas no disponibles)")
        return

    headers = ["Zona", "s inicio (m)", "s fin (m)", "Longitud (m)",
               "L inicio (cd/m²)", "L fin (cd/m²)", "L mín. req. (cd/m²)"]
    col_widths_cm = [3.0, 2.2, 2.2, 2.3, 2.8, 2.8, 3.2]

    table = doc.add_table(rows=1 + len(zones_list), cols=len(headers))
    table.style = 'Table Grid'
    _header_row(table.rows[0], headers)
    for j, w in enumerate(col_widths_cm):
        table.rows[0].cells[j].width = Cm(w)

    ZONE_BG = {
        'threshold':  'FFF3CD',
        'transition': 'CFF4FC',
        'interior':   'D1E7DD',
        'exit':       'E2D9F3',
        'access':     'F8D7DA',
        'parting':    'F8D7DA',
    }

    for i, z in enumerate(zones_list):
        row = table.rows[i + 1]
        zone_type = str(z.get('zone_type', z.get('type', '—'))).lower()
        bg = ZONE_BG.get(zone_type, WHITE)
        values = [
            _zone_label_es(zone_type),
            f"{z.get('s_start', 0):.1f}",
            f"{z.get('s_end', 0):.1f}",
            f"{z.get('length', z.get('s_end', 0) - z.get('s_start', 0)):.1f}",
            f"{z.get('L_start', z.get('L_min_required', 0)):.2f}",
            f"{z.get('L_end', z.get('L_min_required', 0)):.2f}",
            f"{z.get('L_min_required', 0):.2f}",
        ]
        for j, (cell, val) in enumerate(zip(row.cells, values)):
            _set_cell_bg(cell, bg)
            bold = (j == 3 or j == 6)
            _cell_text(cell, val, bold=bold, size_pt=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)
            cell.width = Cm(col_widths_cm[j])


def _build_control(doc: Document, result: dict,
                   title: str = "4. Plan de Control (CIE 88:2004 §87–111)") -> None:
    control = result.get('control')
    if not control:
        return

    _section_heading(doc, title)

    groups = control.get('groups', [])
    scenes = control.get('scenes', [])
    protocol = control.get('protocol', 'DALI')

    if not groups or not scenes:
        doc.add_paragraph("(Plan de control no disponible)")
        return

    # Info del plan
    _kv_table(doc, [
        ("Protocolo de control", protocol),
        ("Nº de grupos",         control.get('n_groups', len(groups))),
        ("Nº de escenas",        control.get('n_scenes', len(scenes))),
    ], col_widths=(5.5, 4.0))

    _section_heading(doc, "Escenas de referencia y regulación continua", level=2)
    continuous = doc.add_paragraph()
    continuous.paragraph_format.space_after = Pt(5)
    r0 = continuous.add_run("Regulación adaptativa continua. ")
    r0.bold = True
    r0.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    r1 = continuous.add_run(
        "Las seis escenas de la tabla son anclas de operación, programación y "
        "puesta en marcha; no son las únicas posiciones que puede adoptar la instalación. "
        "El luminancímetro de acceso mide la luminancia exterior L20 y el controlador "
        "interpela la curva de regulación de cada grupo entre esas anclas, aplicando la "
        "consigna DALI correspondiente a cualquier condición intermedia dentro de los "
        "límites de diseño."
    )
    r1.font.size = Pt(9)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(5)
    note_run = note.add_run(
        "Las escenas de noche normal y noche reducida se seleccionan conforme a la "
        "estrategia de explotación programada; la escena de emergencia permanece "
        "definida por el proyecto eléctrico y de seguridad específico."
    )
    note_run.italic = True
    note_run.font.size = Pt(8.2)
    note_run.font.color.rgb = RGBColor.from_string("666666")

    _section_heading(doc, "Tabla de Regulación (grupos × escenas)", level=2)

    # Columnas: Grupo | L diseño | Escena1 | Escena2 | ...
    n_cols = 2 + len(scenes)
    headers = ["Grupo", "L diseño\n(cd/m²)"] + [
        f"{s['name']}\nL₂₀={s['L20']:.0f}" for s in scenes
    ]

    table = doc.add_table(rows=1 + len(groups), cols=n_cols)
    table.style = 'Table Grid'
    _header_row(table.rows[0], headers)

    # Anchos de columna
    col_w = [5.0, 2.0] + [2.0] * len(scenes)
    for j, w in enumerate(col_w):
        table.rows[0].cells[j].width = Cm(w)

    # Colores de escena
    SCENE_BG = {
        'sunny':    'FFFDE7',
        'normal':   'F5F5F5',
        'overcast': 'ECEFF1',
        'dusk':     'FCE4EC',
        'night':    'E8EAF6',
    }

    ZONE_BG = {
        'threshold':  'FFF3CD',
        'transition': 'CFF4FC',
        'interior':   'D1E7DD',
        'exit':       'E2D9F3',
    }

    for i, g in enumerate(groups):
        row = table.rows[i + 1]
        zone_type = g.get('zone_type', '').lower()
        zone_bg = ZONE_BG.get(zone_type, WHITE)

        # Nombre del grupo
        _set_cell_bg(row.cells[0], zone_bg)
        _cell_text(row.cells[0], g['name'], bold=True, size_pt=8.5)
        row.cells[0].width = Cm(col_w[0])

        # L diseño
        _set_cell_bg(row.cells[1], zone_bg)
        _cell_text(row.cells[1], str(g.get('L_design', '—')),
                   bold=True, color=BLUE_BRAND, size_pt=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        row.cells[1].width = Cm(col_w[1])

        # Niveles por escena
        for j, scene in enumerate(scenes):
            sid  = str(scene['scene_id'])
            pct  = g.get('dimming_levels', {}).get(sid, '—')
            dali = g.get('dali_levels', {}).get(sid, '')
            scene_bg = SCENE_BG.get(scene.get('scene_type', ''), WHITE)

            cell = row.cells[2 + j]
            _set_cell_bg(cell, scene_bg)
            cell.width = Cm(col_w[2 + j])

            # Pct en bold + DALI level pequeño debajo
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"{pct}%")
            r1.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = RGBColor.from_string(BLUE_DARK)
            if protocol == 'DALI' and dali:
                r2 = p.add_run(f"\nDALI {dali}")
                r2.font.size = Pt(7.5)
                r2.font.color.rgb = RGBColor.from_string("666666")

    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(4)
    rn = p_note.add_run(
        "CTH=Umbral · CTRn=Transición · CIN=Interior · CEX=Salida. "
        "Niveles DALI según IEC 62386 (curva logarítmica, rango 0–254)."
    )
    rn.italic = True
    rn.font.size = Pt(8)
    rn.font.color.rgb = RGBColor.from_string("666666")
    emergency = doc.add_paragraph()
    emergency.paragraph_format.space_before = Pt(2)
    er = emergency.add_run(
        "Emergencia: la escena de evacuación y el alumbrado de seguridad se verifican "
        "en el proyecto eléctrico específico; no se infieren de los niveles de regulación ordinaria."
    )
    er.font.size = Pt(8)
    er.font.color.rgb = RGBColor.from_string("666666")

    # Las curvas que utiliza el controlador se incluyen tras la tabla de
    # consignas para que el revisor pueda pasar de los puntos de anclaje a la
    # regulación continua efectiva sin consultar la aplicación.
    curves = control.get("regulation_curves") or []
    curve_by_group = {
        str(curve.get("group_id")): curve
        for curve in curves
        if isinstance(curve, dict) and curve.get("points")
    }
    if curve_by_group:
        _section_heading(doc, "Curvas de atenuación diurna por grupo DALI", level=2)
        p_curve = doc.add_paragraph()
        p_curve.add_run(
            "Cada curva representa la atenuación realmente programable frente a L20. "
            "Las líneas verticales identifican las cuatro escenas diurnas de referencia; "
            "las dos escenas nocturnas son consignas de explotación independientes y se "
            "recogen en la tabla anterior."
        ).font.size = Pt(8.5)
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt

            portal_groups = {"A": []}
            if any(str(group.get("portal") or "").upper() == "B" for group in groups):
                portal_groups["B"] = []
            for group in groups:
                group_id = str(group.get("group_id"))
                if group_id not in curve_by_group:
                    continue
                if str(group.get("layer") or "").lower() == "permanent":
                    for portal in portal_groups:
                        portal_groups[portal].append(group)
                    continue
                portal = str(group.get("portal") or "A").upper()
                portal_groups.setdefault(portal if portal in {"A", "B"} else "A", []).append(group)
            portals = [portal for portal in ("A", "B") if portal_groups.get(portal)]
            if not portals:
                portals = ["A"]

            figure, axes = plt.subplots(
                1, len(portals), figsize=(8.1, 3.45), sharey=True, squeeze=False
            )
            axes = axes[0]
            day_scenes = [
                scene for scene in scenes
                if float(scene.get("L20", 0) or 0) > 0
            ]
            scene_colors = ["#D97706", "#64748B", "#0891B2", "#BE185D"]
            layer_colors = {
                "permanent": "#1A56B0", "reinforcement": "#D97706",
                "adaptation": "#9333EA", "legacy": "#475569",
            }
            accent_colors = ["#D97706", "#9333EA", "#0F766E", "#DC2626", "#0891B2", "#A16207"]
            group_colors = {}
            accent_index = 0
            for group in groups:
                group_id = str(group.get("group_id"))
                if str(group.get("layer") or "").lower() == "permanent":
                    group_colors[group_id] = layer_colors["permanent"]
                else:
                    group_colors[group_id] = accent_colors[accent_index % len(accent_colors)]
                    accent_index += 1

            for axis, portal in zip(axes, portals):
                groups_here = portal_groups[portal]
                for group in groups_here:
                    curve = curve_by_group[str(group.get("group_id"))]
                    points = curve.get("points") or []
                    x_values = [float(point.get("L20", 0) or 0) for point in points]
                    y_values = [float(point.get("dimming_pct", 0) or 0) for point in points]
                    layer = str(group.get("layer", "legacy") or "legacy")
                    group_name = str(group.get("name", "Grupo"))
                    if group_name.startswith("BASE"):
                        label = "BASE permanente"
                    else:
                        label = group_name.replace("REF-", "").replace(" — ", " ")
                    axis.plot(
                        x_values, y_values, linewidth=2.0 if layer == "permanent" else 1.55,
                        color=group_colors.get(str(group.get("group_id")), layer_colors.get(layer, "#475569")),
                        marker="o", markersize=2.2,
                        label=label,
                    )

                for index, scene in enumerate(day_scenes):
                    l20_key = "L20_b" if portal == "B" else "L20"
                    l20_value = float(scene.get(l20_key, scene.get("L20", 0)) or 0)
                    if l20_value <= 0:
                        continue
                    color = scene_colors[index % len(scene_colors)]
                    axis.axvline(l20_value, color=color, linewidth=0.8,
                                linestyle=(0, (3, 2)), alpha=0.78)
                    axis.text(l20_value, 103.5, str(scene.get("name", "Escena")),
                              fontsize=6.4, rotation=90, ha="center", va="bottom", color=color)

                axis.set_title(f"Boca {portal}", fontsize=9.5, fontweight="bold", color="#1A3A6B")
                axis.set_xlabel("L20 medida (cd/m²)", fontsize=8)
                axis.set_xlim(left=0)
                axis.set_ylim(0, 112)
                axis.set_yticks([0, 25, 50, 75, 100])
                axis.grid(True, color="#D9E2F0", linewidth=0.6)
                axis.tick_params(labelsize=7)
                axis.legend(loc="lower right", fontsize=6.6, frameon=True, framealpha=0.92)
            axes[0].set_ylabel("Atenuación / regulación (%)", fontsize=8)
            figure.tight_layout(pad=1.0)
            _save_figure_to_document(doc, figure, width_cm=16.0)

            _section_heading(doc, "Consignas diurnas y nocturnas de referencia", level=2)
            p_scenes = doc.add_paragraph()
            p_scenes.add_run(
                "Las líneas muestran exclusivamente la progresión de las escenas diurnas. "
                "Las marcas cuadradas de noche reducida y noche normal son consignas de "
                "explotación independientes: no representan una rampa temporal entre ambas."
            ).font.size = Pt(8.5)
            figure, axes = plt.subplots(
                1, len(portals), figsize=(8.1, 3.25), sharey=True, squeeze=False
            )
            axes = axes[0]
            scene_labels = []
            for scene in scenes:
                name = str(scene.get("name", "Escena"))
                scene_labels.append(
                    name.replace("Noche reducida", "N. reducida").replace("Noche normal", "N. normal")
                )
            positions_x = list(range(len(scenes)))
            daytime_count = sum(
                1 for scene in scenes
                if float(scene.get("L20", 0) or 0) > 0
            )
            daytime_positions = positions_x[:daytime_count]
            nighttime_positions = positions_x[daytime_count:]
            for axis, portal in zip(axes, portals):
                for group in portal_groups[portal]:
                    layer = str(group.get("layer", "legacy") or "legacy")
                    group_name = str(group.get("name", "Grupo"))
                    if group_name.startswith("BASE"):
                        label = "BASE permanente"
                    else:
                        label = group_name.replace("REF-", "").replace(" — ", " ")
                    values = []
                    levels = group.get("dimming_levels") or {}
                    for scene in scenes:
                        raw_value = levels.get(str(scene.get("scene_id")), 0)
                        try:
                            values.append(float(raw_value))
                        except (TypeError, ValueError):
                            values.append(0.0)
                    # Noche reducida y noche normal son modos de operación
                    # independientes. Conectarlos con la curva diurna crea un
                    # falso valle/pico (especialmente en la BASE permanente)
                    # que parece una rampa de regulación inexistente.
                    axis.plot(
                        daytime_positions, values[:daytime_count], linewidth=1.8 if layer == "permanent" else 1.45,
                        color=group_colors.get(str(group.get("group_id")), layer_colors.get(layer, "#475569")),
                        marker="o", markersize=3.1,
                        label=label,
                    )
                    if nighttime_positions:
                        axis.scatter(
                            nighttime_positions, values[daytime_count:],
                            color=group_colors.get(str(group.get("group_id")), layer_colors.get(layer, "#475569")),
                            marker="s", s=25, zorder=3,
                        )
                if nighttime_positions:
                    axis.axvline(daytime_count - 0.5, color="#94A3B8", linewidth=0.8,
                                linestyle=(0, (3, 2)))
                    axis.text(
                        (daytime_count + len(scenes) - 1) / 2, 103,
                        "Modos nocturnos independientes", ha="center", va="bottom",
                        fontsize=6.2, color="#475569",
                    )
                axis.set_title(f"Boca {portal}", fontsize=9.5, fontweight="bold", color="#1A3A6B")
                axis.set_xticks(positions_x, scene_labels, rotation=22, ha="right", fontsize=6.8)
                axis.set_ylim(0, 105)
                axis.set_yticks([0, 25, 50, 75, 100])
                axis.grid(axis="y", color="#D9E2F0", linewidth=0.6)
                axis.tick_params(axis="y", labelsize=7)
                axis.legend(loc="best", fontsize=6.4, frameon=True, framealpha=0.92)
            axes[0].set_ylabel("Atenuación / regulación (%)", fontsize=8)
            figure.tight_layout(pad=1.0)
            _save_figure_to_document(doc, figure, width_cm=16.0)
        except Exception as exc:
            chart_note = doc.add_paragraph()
            chart_note.add_run(f"No se pudieron generar las curvas de atenuación: {exc}").font.size = Pt(8)
            chart_note.runs[0].font.color.rgb = RGBColor.from_string("B26A00")


def _as_luminaire_dict(luminaire: Any) -> dict:
    """Normaliza el resultado de diseño para que el informe no dependa de su clase."""
    if not luminaire:
        return {}
    if isinstance(luminaire, dict):
        return luminaire
    to_dict = getattr(luminaire, "to_dict", None)
    if callable(to_dict):
        return to_dict()
    return {}


def _humanize(value: Any) -> str:
    """Hace legibles valores de enumeraciones y claves internas."""
    if value in (None, ""):
        return "—"
    return str(value).replace("_", " ").replace("-", " ").strip().title()


def _zone_label_es(zone_type: Any, fallback: Any = None) -> str:
    """Etiqueta corta en español para planos y tablas de informe."""
    labels = {
        "threshold": "Umbral A", "threshold_b": "Umbral B",
        "transition": "Transición A", "transition_b": "Transición B",
        "interior": "Interior", "interior_base": "Base interior",
        "exit": "Salida", "access": "Acceso", "parting": "Partida",
    }
    value = str(zone_type or "").lower()
    return labels.get(value, str(fallback or _humanize(value)))


def _report_status(result: dict, photometric: dict = None) -> tuple[str, str]:
    """Estado único de lectura rápida; nunca declara conformidad sin CIE 140."""
    profile_ok = bool((result.get("validation") or {}).get("valid"))
    ph = photometric or {}
    if not ph.get("available"):
        return (
            "PENDIENTE DE VERIFICACIÓN FOTOMÉTRICA",
            "El perfil CIE 88 es válido, pero falta la comprobación CIE 140 con la fotometría LDT.",
        ) if profile_ok else (
            "REVISAR",
            "El perfil normativo contiene advertencias o incumplimientos que deben resolverse.",
        )
    if profile_ok and ph.get("overall_compliant"):
        return (
            "CONFORME",
            "Se cumplen el perfil CIE 88 y los criterios CIE 140 de todas las zonas verificadas.",
        )
    return (
        "REVISAR",
        "Existe al menos un criterio de perfil o de calidad fotométrica que no cumple.",
    )


def _build_executive_summary(doc: Document, result: dict,
                             photometric: dict = None,
                             luminaire: Any = None) -> None:
    """Primera página técnica: permite validar el caso sin recorrer los anexos."""
    summary = result.get("summary", {})
    lth = result.get("lth", {})
    lum = _as_luminaire_dict(luminaire)
    status, explanation = _report_status(result, photometric)
    status_color = "1A7A3C" if status == "CONFORME" else (
        "B26A00" if status.startswith("PENDIENTE") else "C0392B"
    )

    _section_heading(doc, "1. Resumen de Validación")
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    _set_cell_bg(banner.cell(0, 0), "E8F5E9" if status == "CONFORME" else "FFF4E5")
    _cell_text(
        banner.cell(0, 0),
        f"{status} — {explanation}",
        bold=True,
        color=status_color,
        size_pt=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    rows = [
        ("Tubo / longitud", f"{summary.get('tube_id', '—')}  /  {summary.get('length_m', '—')} m"),
        ("Condición de diseño", f"{summary.get('speed_kmh', '—')} km/h · SD {summary.get('SD_m', '—')} m"),
        ("Portal A", f"L20 {summary.get('L20', '—')} cd/m² · Lth {summary.get('Lth', '—')} cd/m²"),
        ("Portal B", f"L20 {lth.get('L20_b', '—')} cd/m² · Lth {lth.get('Lth_b', '—')} cd/m²"),
        (
            "Interior / noche",
            f"Lin {summary.get('Lin', '—')} cd/m² · "
            f"noche normal "
            f"{summary.get('L_night_normal', summary.get('Lin', '—'))} cd/m² · "
            f"noche reducida "
            f"{summary.get('L_night_reduced', summary.get('L_night', '—'))} cd/m²",
        ),
    ]
    totals = lum.get("totals", {})
    if totals:
        rows.append((
            "Instalación calculada",
            f"{totals.get('n_luminaires', '—')} luminarias · {totals.get('power_kw', '—')} kW instalados",
        ))
    _kv_table(doc, rows)

    warnings = list(result.get("warnings") or [])
    warnings.extend(lum.get("warnings") or [])
    if warnings:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        run = p.add_run("Condicionantes: ")
        run.bold = True
        run.font.size = Pt(8.5)
        detail = p.add_run(" · ".join(str(item) for item in warnings[:3]))
        detail.font.size = Pt(8.5)
        detail.font.color.rgb = RGBColor.from_string("9A4B00")


def _build_design_inputs(doc: Document, result: dict, params: dict,
                         luminaire: Any = None) -> None:
    """Datos de partida que un revisor debe localizar antes de validar resultados."""
    summary = result.get("summary", {})
    speed = result.get("speed", {})
    lth = result.get("lth", {})
    lum = _as_luminaire_dict(luminaire)

    _section_heading(doc, "2. Datos de Partida y Alcance")
    _section_heading(doc, "Geometría, tráfico y explotación", level=2)
    _kv_table(doc, [
        ("Longitud del tubo", f"{summary.get('length_m', params.get('length_m', '—'))} m"),
        ("Calzada", f"{params.get('width_m', params.get('road_width_m', '—'))} m · {params.get('num_lanes', '—')} carriles de {params.get('lane_width_m', '—')} m"),
        ("Sección / altura", f"{_humanize(params.get('tunnel_shape'))} · {params.get('height_m', '—')} m"),
        ("Circulación", _humanize(params.get('traffic_direction'))),
        ("Velocidad / distancia de parada", f"{summary.get('speed_kmh', '—')} km/h · {speed.get('SD_m', summary.get('SD_m', '—'))} m"),
        ("Pendiente / fricción", f"{params.get('gradient_pct', 0)} % · μ={speed.get('friction_coefficient', '—')}"),
        ("Tráfico de proyecto", f"{params.get('traffic_veh_h', '—')} veh/h"),
        ("Pavimento", _humanize((lum.get('road_surface') or {}).get('label') or params.get('road_surface'))),
    ])

    _section_heading(doc, "Entorno y parámetros luminotécnicos", level=2)
    maintenance_factor = (lum.get('luminaire') or {}).get(
        'maintenance_factor', params.get('maintenance_factor', '—')
    )
    try:
        maintenance_factor = f"{float(maintenance_factor):.2f} (aplicado)"
    except (TypeError, ValueError):
        maintenance_factor = str(maintenance_factor)
    rows = [
        ("Método L20 / entorno", f"{_humanize(params.get('l20_method'))} · {_humanize(params.get('environment_type'))}"),
        ("Orientación / cielo", f"{params.get('portal_orientation', '—')} · {_humanize(params.get('sky_condition'))}"),
        ("Portal A: L20 / Lth", f"{summary.get('L20', '—')} / {summary.get('Lth', '—')} cd/m²"),
        ("Portal B: L20 / Lth", f"{lth.get('L20_b', '—')} / {lth.get('Lth_b', '—')} cd/m²"),
        ("Factor de mantenimiento", maintenance_factor),
        ("Reflectancias", f"Paredes ρ={params.get('rho_wall', params.get('wall_reflectance', '—'))} · Techo ρ={params.get('rho_ceiling', params.get('ceiling_reflectance', '—'))}"),
    ]
    _kv_table(doc, rows)


def _build_methodology(doc: Document, result: dict, photometric: dict = None,
                       params: dict = None, luminaire: Any = None) -> None:
    """Explica el método y las decisiones reales de optimización del proyecto."""
    params = params or {}
    lum = _as_luminaire_dict(luminaire)
    lum_spec = lum.get("luminaire") or {}
    designed_zones = lum.get("zones") or []

    def _parameter(*keys, default="—"):
        """Busca un parámetro de diseño sin depender de la versión del payload."""
        for source in (lum, lum_spec, params):
            for key in keys:
                value = source.get(key) if isinstance(source, dict) else None
                if value not in (None, ""):
                    return value
        return default

    goal = str(_parameter("optimization_goal", default="min_luminaires")).lower()
    if goal not in {"min_luminaires", "min_power"}:
        goal = "min_luminaires"
    goal_text = (
        "mínimo número de luminarias: se adopta el mayor vano instalable que cumple todos los límites"
        if goal == "min_luminaires" else
        "mínima potencia lineal: se compara la malla completa y se escoge la menor potencia por metro compatible"
    )
    fixed_spacing = _parameter("d_fixed", "fixed_spacing_m", default=None)
    fixed_layout = fixed_spacing not in (None, "", 0, "0")

    _section_heading(doc, "3. Criterios y Metodología")
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(6)
    run = intro.add_run(
        "El cálculo no se limita a comprobar una solución dibujada previamente. SALVI "
        "Tunnel Engine construye una solución instalable a partir de los datos de proyecto, "
        "la optimiza bajo restricciones luminotécnicas y de producto, y la verifica de nuevo "
        "con la geometría, la fotometría LDT y las escenas de explotación que se entregan. "
        "Así se mantiene la trazabilidad desde la adaptación visual del conductor hasta el "
        "replanteo, la programación DALI y la medición de recepción."
    )
    run.font.size = Pt(9)

    _section_heading(doc, "Marco normativo aplicado", level=2)
    _kv_table(doc, [
        (
            "CIE 88:2004 — diseño del túnel",
            "Define el planteamiento de alumbrado: clasificación y condiciones de "
            "tráfico, velocidad y distancia de parada, evaluación L20 por portal, "
            "luminancia de umbral, zonas y perfil longitudinal de diseño."
        ),
        (
            "CIE 140:2019 — verificación",
            "Comprueba la implantación realmente proyectada con ficheros LDT y mallas "
            "de cálculo: Lavg, uniformidad general U0, uniformidad longitudinal Ul e "
            "incremento de umbral TI."
        ),
        (
            "CIE 144 — pavimento",
            "Aporta las tablas R empleadas para relacionar iluminancia y luminancia en "
            "la calzada; la clase de pavimento elegida es, por tanto, un dato de diseño "
            "que debe confirmarse en obra."
        ),
        (
            "IEC 62386 — control DALI",
            "Da el marco de interoperabilidad para la programación de grupos y escenas. "
            "La tabla DALI de este informe traduce el diseño luminotécnico a consignas "
            "operables y comprobables durante la puesta en marcha."
        ),
    ], col_widths=(5.0, 9.0))

    _section_heading(doc, "3.1. Objetivo de diseño y orden de las decisiones", level=2)
    p = doc.add_paragraph()
    p.add_run("Objetivo principal. ").bold = True
    p.add_run(
        "Obtener un perfil de luminancias seguro y continuo en todo el recorrido, "
        "con una instalación que pueda fabricarse, montarse, regularse y mantenerse. "
        "La eficiencia es un criterio de selección entre soluciones conformes; nunca se "
        "acepta una reducción de potencia o de unidades que comprometa los requisitos de calidad."
    )
    p.runs[-1].font.size = Pt(9)

    _kv_table(doc, [
        ("Estrategia de optimización seleccionada", goal_text),
        ("Punto de partida", "Zona interior como ancla de la retícula; desde ella se resuelven transición y umbral hacia cada portal."),
        ("Variables de diseño", "Interdistancia, óptica LDT, disposición transversal, altura/posición, orientación (tilt), flujo, modelo, driver y corriente."),
        ("Restricciones no negociables", "Perfil CIE 88, Lavg, U0, Ul y TI CIE 140, capacidad de flujo/corriente, geometría real y límites de instalación."),
        ("Modo de implantación", (
            f"Retrofit o interdistancia fija: {fixed_spacing} m; se optimizan óptica, tilt y corriente sobre las posiciones existentes."
            if fixed_layout else
            "Implantación nueva: el motor busca una interdistancia instalable dentro de la malla definida por el proyecto."
        )),
    ], col_widths=(5.2, 8.8))

    p = doc.add_paragraph()
    p.add_run("Principio de prudencia. ").bold = True
    p.add_run(
        "El resultado se considera una propuesta de cálculo, no una sustitución del criterio del proyectista. "
        "La sección, reflectancias, pavimento, fotometría, limitaciones de obra y estrategia de control se "
        "mantienen como hipótesis explícitas y deben confirmarse antes de la aceptación."
    )
    p.runs[-1].font.size = Pt(8.7)

    # Página 2: el proceso de cálculo se muestra como una secuencia auditable,
    # separando el diseño de las comprobaciones de aceptación.
    _add_page_break(doc)
    _section_heading(doc, "3.2. Proceso de optimización luminotécnica", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La optimización se realiza sobre combinaciones físicas de luminaria y no mediante factores de utilización genéricos. "
        "Para cada alternativa se calcula el campo de luminancias derivado de la LDT, la geometría de montaje, la tabla R del pavimento y el factor de mantenimiento."
    ).font.size = Pt(9)

    process_headers = ["Fase", "Qué se varía o resuelve", "Criterio de aceptación", "Resultado trazable"]
    process_rows = [
        (
            "1. Necesidad visual", "L20 y condiciones de tráfico para cada portal; Lth, zonas y perfil longitudinal.",
            "CIE 88:2004 y datos de proyecto.",
            "Perfil objetivo por posición y sentido de circulación.",
        ),
        (
            "2. Ancla interior", "Malla de interdistancias, ópticas F151/F2MD/F2M2 y reglajes de tilt.",
            "U0 y Ul de la retícula CIE 140; flujo alcanzable dentro del driver.",
            "Disposición, vano, óptica y reglaje base más eficientes entre los conformes.",
        ),
        (
            "3. Transición y umbral", "Geometría de cada tramo, flujo requerido por posición y selección de modelo/driver/corriente.",
            "Luminancia exigida, uniformidades y capacidad de producto simultáneamente.",
            "Tramos graduados desde el interior hacia el portal, sin saltos no justificados.",
        ),
        (
            "4. Ajuste conjunto", "Aportación cruzada de luminarias vecinas mediante la matriz de influencia L = A·φ.",
            "Mínimo exceso máximo y, a igualdad, mínimo flujo total; límites de flujo y monotonía por tramo.",
            "Consignas de flujo/corriente coherentes para toda la implantación.",
        ),
        (
            "5. Verificación final", "Instalación completa, observadores y campos bidimensionales de cálculo.",
            "Lavg, U0, Ul, TI y perfil CIE 88; paredes/radiosidad cuando están activadas.",
            "Resultados que se declaran en capítulos 4 y anexos, no estimaciones de anteproyecto.",
        ),
    ]
    table = doc.add_table(rows=1 + len(process_rows), cols=len(process_headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], process_headers)
    process_widths = [2.3, 4.4, 4.2, 4.1]
    for row_index, values in enumerate(process_rows, start=1):
        row = table.rows[row_index]
        bg = GRAY_LIGHT if row_index % 2 else WHITE
        for column, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = Cm(process_widths[column])
            _set_cell_bg(cell, bg)
            _cell_text(cell, value, bold=column == 0, size_pt=7.6,
                       align=WD_ALIGN_PARAGRAPH.LEFT)

    _section_heading(doc, "Cómo se selecciona una solución", level=2)
    selection_points = [
        "La zona interior fija la retícula de referencia porque es la condición permanente y repetitiva de mayor longitud. Para cada vano se prueban las ópticas y los ángulos de tilt definidos por el proyecto.",
        "La uniformidad se comprueba antes de elegir la potencia: aumentar el flujo puede elevar Lavg, pero no corrige una distribución fotométrica o un vano que produzca U0 o Ul insuficientes.",
        "Sólo las combinaciones que cumplen simultáneamente calidad y flujo alcanzable pasan a la selección de modelo, driver y corriente. La variante se elige con el punto de operación necesario, no con una potencia nominal arbitraria.",
        "En transición y umbral se conserva la coherencia con la retícula interior. El flujo se ajusta conjuntamente para considerar la contribución de luminarias de tramos contiguos, evitando sobredimensionar cada punto como si estuviera aislado.",
        "Si se fija una interdistancia por condicionante de obra o retrofit, el cálculo no disimula esa restricción: conserva las posiciones y declara la solución o la advertencia resultante."
    ]
    for item in selection_points:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_after = Pt(2)
        bullet.add_run(item).font.size = Pt(8.8)

    # Esquema gráfico: de las hipótesis a la aprobación final. Ayuda a un
    # revisor experto a localizar la decisión que gobierna cada resultado.
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.patches import FancyBboxPatch

        figure, axis = plt.subplots(figsize=(8.2, 2.5))
        axis.set_xlim(0, 10)
        axis.set_ylim(0, 2.5)
        axis.axis("off")
        steps = [
            (0.12, "Datos", "geometría · L20\ntráfico"),
            (1.67, "CIE 88", "Lth · zonas\nperfil"),
            (3.27, "Retícula", "d · óptica\ntilt"),
            (4.98, "Flujo", "L = A·φ\ndriver"),
            (6.68, "CIE 140", "Lavg · U0\nUl · TI"),
            (8.42, "Obra", "DALI\nmedición"),
        ]
        colors = ["#EAF1FB", "#FFF3CD", "#E0F2FE", "#F3E8FF", "#E8F5E9", "#FDECEC"]
        for index, ((x, title, subtitle), color) in enumerate(zip(steps, colors)):
            width = 1.18 if index not in (2, 3) else 1.35
            patch = FancyBboxPatch(
                (x, 0.98), width, 0.78,
                boxstyle="round,pad=0.06,rounding_size=0.08",
                linewidth=1.0, edgecolor="#1A56B0", facecolor=color,
            )
            axis.add_patch(patch)
            axis.text(x + width / 2, 1.49, title, ha="center", va="center",
                      fontsize=8.4, fontweight="bold", color="#1A3A6B")
            axis.text(x + width / 2, 1.20, subtitle, ha="center", va="center",
                      fontsize=6.4, color="#344054", wrap=True)
            if index < len(steps) - 1:
                next_x = steps[index + 1][0]
                axis.annotate("", xy=(next_x - 0.10, 1.37), xytext=(x + width + 0.06, 1.37),
                              arrowprops=dict(arrowstyle="->", color="#64748B", lw=1.2))
        axis.text(5.0, 0.55,
                  "Filtro obligatorio en cada alternativa: perfil CIE 88 + Lavg/U0/Ul/TI CIE 140 + límite de producto.",
                  ha="center", fontsize=7.1, color="#475569", style="italic")
        axis.text(5.0, 0.17,
                  "Los portales A y B se evalúan independientemente; el resultado puede ser asimétrico.",
                  ha="center", fontsize=7.1, color="#475569", style="italic")
        figure.tight_layout(pad=0.2)
        _save_figure_to_document(doc, figure, width_cm=16.2)
    except Exception:
        pass

    # Página 3: permite que la dirección facultativa relacione el algoritmo
    # con los valores efectivamente seleccionados y con sus límites de validez.
    _add_page_break(doc)
    _section_heading(doc, "3.3. Parámetros que gobiernan esta solución", level=2)
    design_rows = [
        ("Disposición y sección", f"{_humanize(_parameter('arrangement'))} · H={_parameter('mounting_height_m', 'height_m')} m · calzada={_parameter('road_width_m', 'width_m')} m"),
        ("Óptica / temperatura de color", f"{_parameter('optic')} · {_parameter('cct')}"),
        ("Corriente y límite de driver", f"mínimo configurado {_parameter('I_min_pct', default='—')} · máximo {_parameter('I_max_mA')} mA"),
        ("Malla de interdistancias", f"paso {_parameter('spacing_quantum_m', default='—')} m · {'vano fijo ' + str(fixed_spacing) + ' m' if fixed_layout else 'vano elegido por el optimizador'}"),
        ("Pavimento / mantenimiento", f"tabla {_parameter('rtable', default='—')} · MF={_parameter('maintenance_factor', default='—')}"),
        ("Posición transversal", f"distancia a pared {_parameter('wall_offset_m', default='—')} m · offset de eje {_parameter('axis_offset_m', default='—')} m"),
    ]
    _kv_table(doc, design_rows, col_widths=(5.2, 8.8))

    if designed_zones:
        _section_heading(doc, "Decisiones obtenidas por tramo", level=2)
        note = doc.add_paragraph()
        note.add_run(
            "Esta síntesis muestra la consecuencia de la optimización. El detalle completo de emplazamientos se entrega en el plano, el Excel y, cuando procede, el anexo de replanteo."
        ).font.size = Pt(8.5)
        headers = ["Tramo", "L objetivo", "Óptica / modelo", "d (m)", "Tilt", "mA / W", "Decisión"]
        table = doc.add_table(rows=1 + len(designed_zones), cols=len(headers))
        table.style = "Table Grid"
        _header_row(table.rows[0], headers)
        widths = [2.25, 1.65, 3.25, 1.15, 1.15, 1.65, 3.0]
        for row_index, zone in enumerate(designed_zones, start=1):
            row = table.rows[row_index]
            bg = GRAY_LIGHT if row_index % 2 else WHITE
            values = [
                _zone_label_es(zone.get("zone_type"), zone.get("zone_name")),
                f"{zone.get('L_required', zone.get('L_req', '—'))} cd/m²",
                f"{zone.get('optic', lum.get('optic', '—'))} / {zone.get('model', '—')}",
                str(zone.get("d_used", "—")),
                f"{zone.get('tilt_deg', 0)}°",
                f"{zone.get('current_mA', '—')} / {zone.get('power_w', '—')}",
                "Conforme a filtros" if zone.get("feasible", True) else "Revisar advertencia",
            ]
            for column, (cell, value) in enumerate(zip(row.cells, values)):
                cell.width = Cm(widths[column])
                _set_cell_bg(cell, bg)
                _cell_text(cell, value, bold=column in (0, 6), size_pt=7.3,
                           color="1A7A3C" if column == 6 and zone.get("feasible", True) else ("C0392B" if column == 6 else None),
                           align=WD_ALIGN_PARAGRAPH.CENTER if column in (1, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT)

    _section_heading(doc, "3.4. Comprobaciones independientes y límites de validez", level=2)
    validation_paragraphs = [
        "Los dos portales no se fuerzan a compartir una misma solución. Cada uno parte de su propia valoración L20, orientación y escena de entrada; por ello la potencia, longitud de transición y regulación pueden ser distintas. Esta asimetría es una consecuencia del entorno, no una incoherencia del cálculo.",
        "La optimización selecciona una configuración candidata; la conformidad se decide en una segunda comprobación con la instalación completa y los observadores CIE 140. Esta separación evita declarar conforme una estimación local de una sola luminaria o de una celda idealizada.",
        "Las luminarias se evalúan con la LDT asociada a la óptica documentada. Un cambio de modelo, driver, flujo, corriente, fotometría, altura, posición, inclinación, pavimento o reflectancias invalida la equivalencia y exige recalcular o verificar formalmente la desviación.",
        "El factor de mantenimiento ya forma parte del cálculo; no es un margen que pueda aplicarse de nuevo en obra. El mantenimiento efectivo y las mediciones de recepción deben cotejarse con la estrategia declarada en el capítulo 8."
    ]
    if (photometric or {}).get("radiosity", {}).get("enabled"):
        validation_paragraphs.append(
            "En esta ejecución está activada la radiosidad difusa: la comprobación incluye la contribución indirecta de paredes y techo con las reflectancias declaradas. La recepción debe confirmar que dichas superficies mantienen un estado compatible con las hipótesis de proyecto."
        )
    for item in validation_paragraphs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item).font.size = Pt(8.8)

    _section_heading(doc, "3.5. Evidencias para revisión y puesta en marcha", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "El revisor puede reconstruir la decisión sin depender de una explicación verbal: el capítulo 4 aporta el perfil y los márgenes; el capítulo 5 fija la implantación; el capítulo 6 identifica producto y LDT; el capítulo 7 convierte las consignas en control DALI; y los anexos conservan los campos de cálculo, el libro Excel y los documentos de obra."
    ).font.size = Pt(9)


def _build_longitudinal_profile(doc: Document, result: dict,
                                photometric: dict = None) -> None:
    """Incluye L requerida y Lavg CIE 140; no reconstruye Lavg con estimaciones."""
    _section_heading(doc, "4. Resultado Luminotécnico — Perfil Longitudinal")
    chart = result.get("chart", {})
    required = chart.get("data", []) if isinstance(chart, dict) else []
    # Usar el perfil serializado por el motor como fuente normativa canónica
    # evita que una versión resumida del chart mantenga Lth constante durante
    # todo el umbral; la rampa empieza en el 50 % según CIE 88.
    design_profile = result.get("profile", [])
    if isinstance(design_profile, list) and len(design_profile) >= 2:
        required = [
            {"s": point.get("s"), "L": point.get("L"),
             "zone": point.get("zone")}
            for point in design_profile
            if isinstance(point, dict)
            and point.get("s") is not None
            and point.get("L") is not None
        ]
    profile = (photometric or {}).get("real_profile") or {}
    real_points = profile.get("points", []) if isinstance(profile, dict) else []

    if not required:
        doc.add_paragraph("No hay puntos disponibles para representar el perfil longitudinal requerido.")
        return
    if not profile.get("available") or profile.get("metric") != "CIE140_Lavg" or not real_points:
        p = doc.add_paragraph()
        run = p.add_run(
            "La curva de Lavg CIE 140 no está disponible. No se representa ninguna curva estimada; recalcule las luminarias con fotometría LDT para completar la validación."
        )
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("B26A00")
        return

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        required_x = [float(point.get("s", 0)) for point in required]
        required_y = [float(point.get("L", 0)) for point in required]
        actual_x = [float(point.get("s", 0)) for point in real_points]
        actual_y = [float(point.get("L", 0)) for point in real_points]

        fig, ax = plt.subplots(figsize=(8.2, 3.3), dpi=160)
        ax.plot(required_x, required_y, color="#1A56B0", linewidth=2.2,
                label="Lreq CIE 88 (por zona)")
        ax.plot(actual_x, actual_y, color="#C0392B", linewidth=1.8,
                marker="o", markersize=2.4, label="Lavg calculada CIE 140")

        # Bandas y nombres de las zonas normativas, usando los mismos límites
        # que el motor CIE 88 (incluidos los portales B cuando existan).
        zone_colors = {
            "access": "#E5E7EB", "threshold": "#FEF3C7",
            "threshold_b": "#FEF3C7", "transition": "#E0F2FE",
            "transition_b": "#BAE6FD", "interior": "#D1FAE5",
            "exit": "#EDE9FE", "parting": "#F3F4F6",
        }
        zone_labels = {
            "access": "Acceso", "threshold": "Umbral A",
            "threshold_b": "Umbral B", "transition": "Transición A",
            "transition_b": "Transición B", "interior": "Interior",
            "exit": "Salida", "parting": "Post-salida",
        }
        zone_items = []
        for key, zone in (result.get("zones") or {}).items():
            if not isinstance(zone, dict):
                continue
            try:
                start = float(zone.get("s_start", 0) or 0)
                end = float(zone.get("s_end", start) or start)
            except (TypeError, ValueError):
                continue
            if end <= start:
                continue
            zkey = str(zone.get("zone_type", key) or key).lower()
            color = zone_colors.get(zkey, "#E5E7EB")
            ax.axvspan(start, end, color=color, alpha=0.42, linewidth=0,
                       zorder=0)
            if zkey in {"threshold", "threshold_b"}:
                # CIE 88 mantiene Lth en la primera mitad del umbral y
                # comienza el descenso en su punto medio.
                ax.axvline(start + (end - start) * 0.5,
                           color="#D97706", linewidth=0.75,
                           linestyle=(0, (2, 2)), zorder=1)
            zone_items.append((start, end, zkey, color))

        # Referencias de boca/interior, como en la gráfica longitudinal de la SPA.
        summary = result.get("summary") or {}
        try:
            lth_ref = float(summary.get("Lth"))
        except (TypeError, ValueError):
            lth_ref = 0.0
        try:
            lin_ref = float(summary.get("Lin"))
        except (TypeError, ValueError):
            lin_ref = 0.0
        if lth_ref > 0:
            ax.axhline(lth_ref, color="#D97706", linewidth=0.9,
                       linestyle=(0, (4, 3)), label=f"Lth = {lth_ref:.1f}")
        if lin_ref > 0:
            ax.axhline(lin_ref, color="#2E8B57", linewidth=0.9,
                       linestyle=(0, (4, 3)), label=f"Lin = {lin_ref:.1f}")

        # La escala logarítmica coincide con la vista longitudinal de la SPA y
        # hace visible el descenso desde Lth hasta Lin.
        positive_values = [value for value in required_y + actual_y if value > 0]
        if positive_values:
            ax.set_yscale("log")
            ax.set_ylim(max(min(positive_values) * 0.75, 0.01),
                        max(positive_values) * 1.35)
        for start, end, zkey, _color in zone_items:
            if end - start >= max(20.0, (max(required_x or [end]) * 0.035)):
                top = ax.get_ylim()[1]
                ax.text((start + end) / 2.0, top / 1.12,
                        zone_labels.get(zkey, zkey.replace("_", " ").title()),
                        ha="center", va="top", fontsize=6.8,
                        fontweight="bold", color="#475569", clip_on=True)
        ax.set_xlabel("Posición longitudinal (m)", fontsize=8)
        ax.set_ylabel("Luminancia (cd/m²)", fontsize=8)
        ax.grid(True, color="#D9E2F0", linewidth=0.6)
        ax.tick_params(labelsize=7)
        ax.legend(loc="upper right", frameon=False, fontsize=7)
        fig.tight_layout(pad=0.8)
        image = io.BytesIO()
        fig.savefig(image, format="png", transparent=False)
        plt.close(fig)
        image.seek(0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(image, width=Cm(16.0))
    except Exception as exc:
        p = doc.add_paragraph()
        run = p.add_run(f"No se pudo generar la gráfica longitudinal: {exc}")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("C0392B")
        return

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    run = note.add_run(
        "La curva roja procede exclusivamente de Lavg calculada en campos bidimensionales CIE 140 entre luminarias consecutivas."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")


def _save_figure_to_document(doc: Document, figure, width_cm: float = 16.0) -> None:
    """Inserta una figura Matplotlib en el Word sin crear archivos temporales."""
    image = io.BytesIO()
    figure.savefig(image, format="png", dpi=180, facecolor="white")
    import matplotlib.pyplot as plt
    plt.close(figure)
    image.seek(0)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image, width=Cm(width_cm))


def _luminaire_y_positions(arrangement: str, width_m: float,
                           wall_offset_m: float) -> list[float]:
    """Posiciones transversales, coherentes con el motor de cálculo."""
    width = max(float(width_m or 0.0), 0.1)
    offset = min(max(0.05, float(wall_offset_m or 0.30)), width / 2.0 - 0.05)
    arrangement = str(arrangement or "central_single")
    if arrangement in {
        "bilateral_sym", "bilateral_stag", "bilateral", "staggered",
        "central_double",
    }:
        return [offset, width - offset]
    if arrangement in {"central_offset", "lateral_left"}:
        return [offset]
    if arrangement in {"lateral_right", "unilateral"}:
        return [width - offset]
    return [width / 2.0]


def _build_installation_visuals(doc: Document, result: dict, params: dict,
                                luminaire: Any) -> None:
    """Planta de situación/implantación y sección acotada de montaje."""
    lum = _as_luminaire_dict(luminaire)
    lum_zones = lum.get("zones", [])
    if not lum_zones:
        return

    _section_heading(doc, "Planta de situación e implantación", level=2)
    satellite_map = _satellite_location_map(params)
    if satellite_map:
        map_paragraph = doc.add_paragraph()
        map_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        map_paragraph.add_run().add_picture(satellite_map, width=Cm(11.8))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run(
            "Situación en planta: ortofoto, eje de cálculo y portales A/B. Imagen satelital © Esri."
        )
        caption_run.italic = True
        caption_run.font.size = Pt(8)
        caption_run.font.color.rgb = RGBColor.from_string("555555")
    else:
        note = doc.add_paragraph(
            "No se incluye la situación satelital porque el proyecto no dispone de eje OSM ni coordenadas válidas."
        )
        note.runs[0].font.size = Pt(8.5)
        note.runs[0].font.color.rgb = RGBColor.from_string("666666")

    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        tube_length = float(result.get("summary", {}).get(
            "length_m", lum.get("tube_length_m", 0)
        ) or 0)
        zone_colors = {
            "threshold": "#F6C85F", "threshold_b": "#F6C85F",
            "transition": "#6FB1E8", "transition_b": "#6FB1E8",
            "interior": "#8BCB9B", "interior_base": "#8BCB9B",
            "exit": "#B9A4D8", "access": "#E8A0A0", "parting": "#E8A0A0",
        }
        layer_y = {"permanent": 0.68, "adaptation": 0.50,
                   "reinforcement": 0.32, "legacy": 0.50}
        layer_colors = {"permanent": "#1A56B0", "adaptation": "#A855F7",
                        "reinforcement": "#D97706", "legacy": "#374151"}
        layer_labels = {"permanent": "Base permanente", "adaptation": "Adaptación",
                        "reinforcement": "Refuerzo", "legacy": "Instalación"}

        _section_heading(doc, "Planta longitudinal de luminarias", level=2)
        figure, axis = plt.subplots(figsize=(9.0, 3.6))
        for zone in (result.get("zones") or {}).values():
            if not isinstance(zone, dict):
                continue
            start = float(zone.get("s_start", 0) or 0)
            end = float(zone.get("s_end", start) or start)
            z_type = str(zone.get("zone_type", zone.get("type", ""))).lower()
            color = zone_colors.get(z_type, "#D9E2F0")
            axis.axvspan(start, end, ymin=0.02, ymax=0.98, color=color,
                         alpha=0.40, linewidth=0)
            if end > start + max(25.0, tube_length * 0.04):
                axis.text(
                    (start + end) / 2, 0.92,
                    _zone_label_es(z_type, zone.get("zone_name")),
                    ha="center", va="top", fontsize=7, fontweight="bold",
                )

        used_layers = set()
        for zone in lum_zones:
            layer = str(zone.get("control_layer", "legacy") or "legacy")
            y = layer_y.get(layer, 0.50)
            points = zone.get("setpoints", []) or []
            positions = [float(point.get("s", 0) or 0) for point in points]
            if not positions:
                start = float(zone.get("s_start", 0) or 0)
                end = float(zone.get("s_end", start) or start)
                count = int(zone.get("n_positions", zone.get("n_luminaires", 0)) or 0)
                if count > 0 and end > start:
                    positions = [start + (index + 0.5) * (end - start) / count
                                 for index in range(count)]
            if positions:
                label = layer_labels.get(layer, _humanize(layer)) if layer not in used_layers else None
                axis.scatter(positions, [y] * len(positions), s=16,
                             color=layer_colors.get(layer, "#374151"),
                             label=label, zorder=3)
                used_layers.add(layer)

        axis.set_xlim(0, max(tube_length, 1.0))
        axis.set_ylim(0.12, 1.02)
        axis.set_yticks([0.32, 0.50, 0.68], ["Refuerzo", "Adaptación", "Base permanente"])
        axis.set_xlabel("Posición longitudinal (m)", fontsize=8)
        axis.set_title("Zonas CIE 88 y posiciones físicas de luminarias", fontsize=10,
                       fontweight="bold")
        axis.grid(axis="x", color="#D9E2F0", linewidth=0.7)
        axis.tick_params(labelsize=7)
        if used_layers:
            axis.legend(loc="lower center", bbox_to_anchor=(0.5, -0.35),
                        ncol=max(1, len(used_layers)), frameon=False, fontsize=7)
        figure.tight_layout(pad=1.0)
        _save_figure_to_document(doc, figure)

        width = float(params.get("width_m", lum.get("road_width_m", 7.0)) or 7.0)
        height = float(params.get("height_m", 5.0) or 5.0)
        mount_height = float((lum.get("luminaire") or {}).get(
            "mounting_height_m", params.get("mounting_height_m", 4.5)
        ) or 4.5)
        offset = float((lum.get("luminaire") or {}).get(
            "wall_offset_m", params.get("wall_offset_m", 0.30)
        ) or 0.30)
        positions = _luminaire_y_positions(lum.get("arrangement"), width, offset)

        _section_heading(doc, "Sección transversal y cotas de montaje", level=2)
        # Esta figura replica la geometría de TunnelSectionPreview de la
        # aplicación. No se sustituye la bóveda por un rectángulo: la forma,
        # los hastiales, los carriles y las filas de luminarias coinciden con
        # la sección definida por el proyectista.
        shape = str(params.get("tunnel_shape", "horseshoe") or "horseshoe").lower()
        if shape not in {"horseshoe", "circular", "rectangular"}:
            shape = "horseshoe"
        wall_height = min(max(float(params.get("H_pared_m", 3.0) or 3.0), 0.0), height)
        if shape == "rectangular":
            outline = [(0, 0), (0, height), (width, height), (width, 0)]
            shape_label = "Sección rectangular"
        elif shape == "circular":
            radius = ((width / 2.0) ** 2 + height ** 2) / (2.0 * height)
            centre_z = height - radius
            angle_limit = math.acos(max(-1.0, min(1.0, -centre_z / radius)))
            arch = [
                (width / 2.0 + radius * math.sin(-angle_limit + index * 2.0 * angle_limit / 36.0),
                 centre_z + radius * math.cos(-angle_limit + index * 2.0 * angle_limit / 36.0))
                for index in range(37)
            ]
            outline = [(0, 0), *arch, (width, 0)]
            shape_label = f"Sección circular · R = {radius:.2f} m"
        else:
            arch = [
                (width / 2.0 + (width / 2.0) * math.cos(math.pi - index * math.pi / 36.0),
                 wall_height + (height - wall_height) * math.sin(math.pi - index * math.pi / 36.0))
                for index in range(37)
            ]
            outline = [(0, 0), (0, wall_height), *arch, (width, wall_height), (width, 0)]
            shape_label = "Sección en herradura / bóveda"

        figure, axis = plt.subplots(figsize=(7.5, 3.55))
        axis.set_facecolor("#0F172A")
        figure.patch.set_facecolor("white")
        outline_x = [point[0] for point in outline]
        outline_z = [point[1] for point in outline]
        axis.fill(outline_x, outline_z, facecolor="#1E293B", edgecolor="#94A3B8",
                  linewidth=1.8, zorder=1)
        axis.add_patch(plt.Rectangle((0, 0), width, 0.11, color="#374151", zorder=3))

        num_lanes = max(1, int(params.get("num_lanes", 2) or 2))
        lane_width = float(params.get("lane_width_m", 3.5) or 3.5)
        shoulder_left = max(0.0, float(params.get("shoulder_left_m", 0) or 0))
        carriageway_end = min(width, shoulder_left + num_lanes * lane_width)
        for lane_index in range(num_lanes + 1):
            x_lane = shoulder_left + lane_index * lane_width
            if 0.0 < x_lane < width and (lane_index in (0, num_lanes) or lane_index < num_lanes):
                axis.plot([x_lane, x_lane], [0.01, 0.11], color="#F8FAFC", linewidth=0.85,
                          linestyle="--" if lane_index not in (0, num_lanes) else "-", zorder=4)
        if carriageway_end < width - 0.01:
            axis.plot([carriageway_end, carriageway_end], [0.01, 0.11], color="#F8FAFC",
                      linewidth=0.85, zorder=4)
        if shape == "horseshoe" and 0.0 < wall_height < height:
            axis.plot([0, width], [wall_height, wall_height], color="#64748B", linewidth=0.7,
                      linestyle=(0, (4, 3)), alpha=0.85, zorder=2)

        for pos in positions:
            inside = 0 <= pos <= width and 0 <= mount_height <= height
            if inside and shape == "horseshoe" and mount_height > wall_height:
                a, b = width / 2.0, height - wall_height
                if b > 0:
                    inside = ((pos - width / 2.0) ** 2 / a ** 2
                              + (mount_height - wall_height) ** 2 / b ** 2) <= 1.02
            lum_color = "#FBBF24" if inside else "#F87171"
            axis.scatter([pos], [mount_height], color=lum_color, marker="s", s=70,
                         edgecolors="#FFFFFF", linewidths=1.1, zorder=6)
            axis.plot([pos, pos], [0.11, mount_height], color=lum_color, linewidth=0.85,
                      linestyle=(0, (3, 2)), alpha=0.7, zorder=2)
            axis.text(pos + 0.10, (mount_height + 0.11) / 2, f"h={mount_height:.2f} m",
                      color="#CBD5E1", fontsize=6.8, ha="left", va="center", zorder=7)
        # Cotas de geometría: ancho, altura libre y altura de montaje.
        dim_style = dict(arrowstyle="<->", color="#1A56B0", lw=1.0)
        internal_dim_style = dict(arrowstyle="<->", color="#60A5FA", lw=1.0)
        axis.annotate("", xy=(0, -0.36), xytext=(width, -0.36), arrowprops=dim_style)
        axis.text(width / 2, -0.55, f"Ancho de calzada: {width:.2f} m",
                  ha="center", va="top", fontsize=8, color="#1A3A6B")
        axis.annotate("", xy=(-0.42, 0), xytext=(-0.42, height), arrowprops=dim_style)
        axis.text(-0.54, height / 2, f"Altura libre\n{height:.2f} m",
                  ha="right", va="center", fontsize=8, color="#1A3A6B")
        axis.annotate("", xy=(width + 0.42, 0), xytext=(width + 0.42, mount_height),
                      arrowprops=dim_style)
        axis.text(width + 0.56, mount_height / 2, f"Altura de montaje\n{mount_height:.2f} m",
                  ha="left", va="center", fontsize=8, color="#1A3A6B")

        # La posición transversal se muestra desde la pared más próxima; en
        # disposición doble se acota cada fila para que la figura sea válida
        # como referencia de replanteo, no sólo como esquema ilustrativo.
        for index, pos in enumerate(positions):
            label = f"x{index + 1} = {pos:.2f} m"
            if index == 0:
                axis.annotate("", xy=(0, 0.46), xytext=(pos, 0.46), arrowprops=internal_dim_style)
                axis.text(pos / 2, 0.60, f"{label} desde pared izq.",
                          ha="center", va="bottom", fontsize=7.2, color="#DBEAFE")
            else:
                distance_right = width - pos
                axis.annotate("", xy=(pos, 0.76), xytext=(width, 0.76), arrowprops=internal_dim_style)
                axis.text((pos + width) / 2, 0.90,
                          f"x{index + 1} = {distance_right:.2f} m desde pared der.",
                          ha="center", va="bottom", fontsize=7.2, color="#DBEAFE")
        if shape == "horseshoe" and 0.0 < wall_height < height:
            axis.text(width + 0.12, wall_height / 2, f"Hastial\n{wall_height:.2f} m",
                      ha="left", va="center", fontsize=7.2, color="#1A3A6B")
        axis.set_xlim(-1.15, width + 2.15)
        axis.set_ylim(-0.85, max(height + 0.55, mount_height + 0.55))
        axis.set_aspect("equal", adjustable="box")
        axis.axis("off")
        axis.set_title(
            f"{shape_label} — {_humanize(lum.get('arrangement'))} · luminarias en amarillo",
            fontsize=10, fontweight="bold", color="#1A3A6B",
        )
        figure.tight_layout(pad=0.8)
        _save_figure_to_document(doc, figure, width_cm=14.5)
    except Exception as exc:
        note = doc.add_paragraph()
        run = note.add_run(f"No se pudieron generar los planos de implantación: {exc}")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("C0392B")


def _build_luminaire_catalog(doc: Document, luminaire: Any) -> None:
    """Fichas de las combinaciones modelo/driver/óptica efectivamente instaladas."""
    lum = _as_luminaire_dict(luminaire)
    zones = lum.get("zones", []) or []
    if not zones:
        return

    combinations = {}
    for zone in zones:
        model = str(zone.get("model", "") or "")
        optic = str(zone.get("optic", lum.get("optic", "")) or "").upper()
        if not model:
            continue
        key = (model, optic)
        combinations.setdefault(key, []).append(zone)
    if not combinations:
        return

    _section_heading(doc, "6. Luminarias Empleadas: Fichas, Geometría y Fotometría")
    intro = doc.add_paragraph(
        "Se documentan exclusivamente las combinaciones de modelo, driver y óptica instaladas en este tubo. "
        "Las cotas y la imagen proceden de las fichas técnicas APHEX; las curvas polares se obtienen de la LDT usada en la verificación CIE 140."
    )
    intro.runs[0].font.size = Pt(9)

    for item_index, ((model, optic), selected_zones) in enumerate(combinations.items(), start=1):
        if item_index > 1:
            _add_page_break(doc)
        family = next((name for name in _APHEX_CATALOG if model.startswith(name)), "")
        catalog = _APHEX_CATALOG.get(family, {})
        _section_heading(doc, f"6.{item_index}. {model} · Óptica {optic or '—'}", level=2)

        variant = None
        try:
            from modules.tunnel.led_engine import VARIANTS_BY_ID
            variant = VARIANTS_BY_ID.get(model)
        except Exception:
            pass

        product_image = None
        try:
            import fitz
            from PIL import Image
            pdf_path = _APHEX_ASSET_DIR / str(catalog.get("pdf", ""))
            if pdf_path.exists():
                pdf = fitz.open(pdf_path)
                page = pdf[0]
                clip = fitz.Rect(0, 0, page.rect.width, page.rect.height * 0.79)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.7, 1.7), clip=clip, alpha=False)
                product_image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                pdf.close()
        except Exception:
            product_image = None

        photometry = None
        # Usar exactamente la misma fuente LDT que el motor CIE 140.  Mantener
        # una segunda tabla en el informe permite que las curvas documentadas
        # se desincronicen de las empleadas para dimensionar el proyecto.
        from modules.tunnel.optimizer import _LDT_DIR, _OPTIC_LDT
        ldt_filename = _OPTIC_LDT.get(optic)
        if ldt_filename:
            try:
                from photometric_engine.salvi_photometry.ldt_parser import load_ldt
                ldt_path = _LDT_DIR / ldt_filename
                if ldt_path.exists():
                    photometry = load_ldt(ldt_path)
            except Exception:
                photometry = None

        # Figura compacta: fotografía oficial, plano acotado y curva polar
        # de la misma LDT que alimenta la comprobación CIE 140.
        try:
            import numpy as np
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            from matplotlib.patches import Rectangle

            # Se concede deliberadamente más superficie a la foto: es la
            # referencia visual de compra e instalación de esta ficha.
            fig = plt.figure(figsize=(8.35, 3.45), dpi=180)
            grid = fig.add_gridspec(1, 3, width_ratios=[1.52, 0.86, 1.08], wspace=0.34)
            photo_axis = fig.add_subplot(grid[0, 0])
            if product_image is not None:
                photo_axis.imshow(product_image)
            else:
                photo_axis.add_patch(Rectangle((0.1, 0.32), 0.8, 0.30, facecolor="#D5E8F5", edgecolor="#1A56B0", linewidth=1.5))
                photo_axis.text(0.5, 0.48, catalog.get("product", family or model), ha="center", va="center", fontsize=8, fontweight="bold")
            photo_axis.set_title("Imagen de producto", fontsize=8.5, fontweight="bold", pad=5)
            photo_axis.axis("off")

            geometry_axis = fig.add_subplot(grid[0, 1])
            length, width, thickness = catalog.get("dimensions_mm", (0, 0, 0))
            if length and width:
                ratio = width / length
                geometry_axis.add_patch(Rectangle((-0.46, -ratio / 2), 0.92, ratio, facecolor="#EAF1FB", edgecolor="#1A56B0", linewidth=1.5))
                geometry_axis.annotate("", xy=(-0.46, ratio / 2 + 0.10), xytext=(0.46, ratio / 2 + 0.10), arrowprops=dict(arrowstyle="<->", color="#475569", lw=1))
                geometry_axis.text(0, ratio / 2 + 0.14, f"{length} mm", ha="center", fontsize=7.6, color="#1A3A6B")
                geometry_axis.annotate("", xy=(-0.60, -ratio / 2), xytext=(-0.60, ratio / 2), arrowprops=dict(arrowstyle="<->", color="#475569", lw=1))
                geometry_axis.text(-0.67, 0, f"{width} mm", ha="center", va="center", fontsize=7.4, color="#1A3A6B", rotation=90)
                geometry_axis.add_patch(Rectangle((-0.46, -ratio / 2 - 0.31), 0.92, 0.05, facecolor="#CBD5E1", edgecolor="#475569", linewidth=0.8))
                geometry_axis.annotate("", xy=(0.57, -ratio / 2 - 0.31), xytext=(0.57, -ratio / 2 - 0.26), arrowprops=dict(arrowstyle="<->", color="#475569", lw=1))
                geometry_axis.text(0.64, -ratio / 2 - 0.285, f"{thickness} mm", ha="left", va="center", fontsize=7.2, color="#1A3A6B")
                geometry_axis.set_xlim(-0.82, 0.96)
                geometry_axis.set_ylim(-ratio / 2 - 0.43, ratio / 2 + 0.27)
            else:
                geometry_axis.text(0.5, 0.5, "Cotas no disponibles", ha="center", va="center", fontsize=8)
            geometry_axis.set_title("Geometría de catálogo (sin lira)", fontsize=8.5, fontweight="bold", pad=5)
            geometry_axis.set_aspect("equal", adjustable="box")
            geometry_axis.axis("off")

            polar_axis = fig.add_subplot(grid[0, 2], projection="polar")
            if photometry is not None:
                gamma = np.linspace(0, min(180.0, float(photometry.g_angles[-1])), 181)
                i_c0 = np.array([photometry.normalized_intensity(0, angle) for angle in gamma])
                i_c180 = np.array([photometry.normalized_intensity(180, angle) for angle in gamma])
                i_c90 = np.array([photometry.normalized_intensity(90, angle) for angle in gamma])
                i_c270 = np.array([photometry.normalized_intensity(270, angle) for angle in gamma])
                normalizer = max(
                    float(i_c0.max()), float(i_c180.max()),
                    float(i_c90.max()), float(i_c270.max()), 1e-6,
                )
                theta = np.deg2rad(gamma)
                # Un plano C completo no es el espejo de una semicurva: para
                # ópticas asimétricas (como F151) cada lado procede de su
                # plano opuesto, C0/C180 o C90/C270 respectivamente.
                for positive, negative, color, label in (
                    (i_c0, i_c180, "#1A56B0", "C0–C180"),
                    (i_c90, i_c270, "#D97706", "C90–C270"),
                ):
                    positive = positive / normalizer * 100
                    negative = negative / normalizer * 100
                    polar_axis.plot(
                        np.r_[-theta[:0:-1], theta],
                        np.r_[negative[:0:-1], positive],
                        color=color, lw=1.5, label=label,
                    )
                polar_axis.set_rmax(100)
                polar_axis.set_rticks([25, 50, 75, 100])
                # El LDT se representa en la orientación de montaje: el eje
                # fotométrico 0° que originalmente se dibuja hacia la derecha
                # queda girado 90° en sentido horario, apuntando a calzada.
                # No se invierte el sentido angular: C0/C180 y C90/C270 ya
                # conservan los dos semiplanos reales del fichero LDT.
                polar_axis.set_theta_zero_location("S")
                polar_axis.set_theta_direction(1)
                polar_axis.set_rlabel_position(135)
                polar_axis.set_thetagrids([0, 45, 90, 135, 180, 225, 315])
                polar_axis.tick_params(labelsize=5.5)
                polar_axis.legend(loc="lower center", bbox_to_anchor=(0.5, -0.17), ncol=1, fontsize=5.8, frameon=False)
            else:
                polar_axis.text(0.5, 0.5, "LDT no disponible", transform=polar_axis.transAxes, ha="center", va="center", fontsize=8)
            polar_axis.set_title("Curvas polares normalizadas\nC0–C180 · C90–C270 · 0° hacia calzada", fontsize=8.5, fontweight="bold", pad=8)
            fig.tight_layout(pad=0.5)
            _save_figure_to_document(doc, fig, width_cm=16.0)
        except Exception as exc:
            p = doc.add_paragraph(f"No se pudo componer la ficha visual de la luminaria: {exc}")
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor.from_string("C0392B")

        installed_count = sum(int(zone.get("n_luminaires", 0) or 0) for zone in selected_zones)
        currents = [float(zone.get("current_mA", 0) or 0) for zone in selected_zones]
        powers = [float(zone.get("power_w", 0) or 0) for zone in selected_zones]
        fluxes = [float(zone.get("flux_lm", 0) or 0) for zone in selected_zones]
        zones_text = ", ".join(str(zone.get("zone_name", "—")) for zone in selected_zones)
        driver_text = "—"
        if variant is not None:
            driver_text = (
                f"{variant.driver_count} × {variant.driver_manufacturer} {variant.driver_model} "
                f"({variant.driver_rated_power_total_w:.0f} W nominales; límite de luminaria {variant.luminaire_max_input_power_w:.0f} W)"
            )
        _kv_table(doc, [
            ("Producto / configuración", f"{catalog.get('product', family or model)} · {model} · {optic or 'óptica no indicada'}"),
            ("Unidades y tramos", f"{installed_count} luminarias · {zones_text}"),
            ("Driver", driver_text),
            ("Punto de operación calculado", f"{min(currents, default=0):.1f}–{max(currents, default=0):.1f} mA · {min(powers, default=0):.1f}–{max(powers, default=0):.1f} W · {min(fluxes, default=0):,.0f}–{max(fluxes, default=0):,.0f} lm"),
            ("Características de familia", f"{catalog.get('protection', '—')} · {catalog.get('weight', '—')} · {catalog.get('leds', '—')} · instalación {catalog.get('mounting', '—')}"),
            ("Construcción y control", "Cuerpo de aluminio, vidrio templado y lentes PMMA. Fijación con lira orientable cada 5°; compatible con DALI/Zhaga según ficha de producto."),
            ("Fotometría de cálculo", f"{ldt_filename or '—'} · flujo LDT {getattr(photometry, '_flux_file_lm', 0):,.0f} lm · curva normalizada a 1 klm." if photometry is not None else "LDT de referencia no disponible."),
            ("Fuente y alcance", f"Ficha oficial {catalog.get('pdf', 'APHEX')} (ed. 2025-01) y LDT de referencia del motor. Verificar la revisión de ficha/driver antes de la compra."),
        ], col_widths=(5.4, 8.6))


def _operating_scenes(luminaire: Any) -> list[tuple[str, dict]]:
    """Devuelve una sola vez cada escena que el operador reconoce en obra."""
    lum = _as_luminaire_dict(luminaire)
    scenarios = lum.get("scenarios", {}) or {}
    aliases = [
        ("sunny", ("sunny",)),
        ("normal", ("normal",)),
        ("overcast", ("overcast",)),
        ("dusk", ("dusk",)),
        ("night_normal", ("night_normal", "night")),
        ("night_reduced", ("night_reduced",)),
    ]
    selected = []
    used = set()
    for canonical, candidates in aliases:
        key = next((candidate for candidate in candidates
                    if isinstance(scenarios.get(candidate), dict)
                    and scenarios[candidate].get("power_kw") is not None), None)
        if key and key not in used:
            selected.append((canonical, scenarios[key]))
            used.add(key)
    return selected


def _build_operation_chart(doc: Document, luminaire: Any) -> None:
    """Gráfica de potencia por escena para operación y eficiencia energética."""
    scenarios = [value for _key, value in _operating_scenes(luminaire)]
    if not scenarios:
        return
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        labels = [str(item.get("name", "Escena")) for item in scenarios]
        powers = [float(item.get("power_kw", 0) or 0) for item in scenarios]
        figure, axis = plt.subplots(figsize=(7.5, 2.9))
        bars = axis.bar(labels, powers, color="#1A56B0")
        for bar, value in zip(bars, powers):
            axis.text(bar.get_x() + bar.get_width() / 2, bar.get_height(),
                      f"{value:.2f}", ha="center", va="bottom", fontsize=8)
        axis.set_ylabel("Potencia (kW)", fontsize=8)
        axis.set_title("Potencia eléctrica por escena de operación", fontsize=10,
                       fontweight="bold")
        axis.grid(axis="y", color="#D9E2F0", linewidth=0.7)
        axis.tick_params(axis="x", labelrotation=18, labelsize=7)
        axis.tick_params(axis="y", labelsize=7)
        figure.tight_layout(pad=0.8)
        _save_figure_to_document(doc, figure, width_cm=15.0)
    except Exception:
        return


def _field_luminaires_for_report(luminaire: Any, field_start: float,
                                 field_end: float) -> list[dict]:
    """Expande las unidades físicas que delimitan un campo CIE 140.

    La representación usa las mismas coordenadas Y (desde pared izquierda)
    que el Excel y el motor fotométrico. Así el revisor ve las luminarias que
    fijan el intervalo, no símbolos genéricos superpuestos al mapa.
    """
    lum = _as_luminaire_dict(luminaire)
    zones = lum.get("zones", []) or []
    if not zones:
        return []
    spec = lum.get("luminaire") or {}
    road_width = float(lum.get("road_width_m", 0) or 0)
    if road_width <= 0:
        return []
    arrangement = str(lum.get("arrangement", "central_single") or "central_single")
    wall_offset = float(spec.get("wall_offset_m", lum.get("wall_offset_m", 0.30)) or 0.30)
    offset = min(max(0.05, wall_offset), max(0.05, road_width / 2.0 - 0.05))

    def rows_for_station(station: int) -> list[float]:
        if arrangement in {"central_double", "bilateral_sym", "bilateral"}:
            return [offset, road_width - offset]
        if arrangement in {"bilateral_stag", "staggered"}:
            return [offset if station % 2 else road_width - offset]
        if arrangement in {"central_offset", "lateral_left"}:
            return [offset]
        if arrangement in {"lateral_right", "unilateral"}:
            return [road_width - offset]
        return [road_width / 2.0]

    selected = []
    tolerance = 0.08
    for zone in zones:
        for sequence, point in enumerate(zone.get("setpoints", []) or [], start=1):
            try:
                x = float(point.get("s", 0) or 0)
            except (TypeError, ValueError):
                continue
            if x < field_start - tolerance or x > field_end + tolerance:
                continue
            station = int(point.get("idx", sequence) or sequence)
            for y in rows_for_station(station):
                selected.append({
                    "x": x,
                    "y": float(y),
                    "model": str(point.get("model", zone.get("model", "—")) or "—"),
                    "optic": str(point.get("optic", zone.get("optic", lum.get("optic", "—"))) or "—"),
                    "power_w": point.get("power_w", zone.get("power_w", "—")),
                    "current_mA": point.get("current_mA", zone.get("current_mA", "—")),
                    "tilt_deg": point.get("tilt_deg", zone.get("tilt_deg", 0)),
                })
    selected.sort(key=lambda item: (item["x"], item["y"]))
    return selected


def _build_cie140_field_annex(doc: Document, photometric: dict = None,
                              luminaire: Any = None) -> None:
    """Anexo breve por zona: sólo los campos gobernantes, nunca una matriz completa."""
    profile = (photometric or {}).get("real_profile") or {}
    fields = profile.get("fields", []) if profile.get("available") else []
    if not fields:
        return

    grouped = {}
    for field in fields:
        name = str(field.get("zone_name") or field.get("zone_type") or "Zona")
        grouped.setdefault(name, []).append(field)

    for index, (zone_name, zone_fields) in enumerate(grouped.items(), start=1):
        _add_page_break(doc)
        _section_heading(doc, f"ANEXO B.{index}. Caso crítico CIE 140 — {zone_name}")

        def field_range(item: dict) -> str:
            return f"{float(item.get('field_start', 0) or 0):.1f}–{float(item.get('field_end', 0) or 0):.1f}"

        def observer(item: dict) -> str:
            direction = 'A→B' if int(item.get('observer_direction', 1) or 1) >= 0 else 'B→A'
            return f"C{item.get('observer_lane_number', '—')} / {direction}"

        # El informe principal ya declara el resultado por zona. Aquí se
        # conservan únicamente los cuatro campos que gobiernan la aceptación:
        # es suficiente para auditoría y evita reproducir una matriz por cada
        # observador. El detalle punto a punto queda disponible en Excel.
        critical = [
            ("Luminancia media mínima", min(zone_fields, key=lambda item: float(item.get('L', 0) or 0)), "Lavg", "cd/m²"),
            ("Uniformidad general mínima", min(zone_fields, key=lambda item: float(item.get('U0', 0) or 0)), "U0", ""),
            ("Uniformidad longitudinal mínima", min(zone_fields, key=lambda item: float(item.get('Ul', 0) or 0)), "Ul", ""),
            ("Incremento de umbral máximo", max(zone_fields, key=lambda item: float(item.get('TI', 0) or 0)), "TI", "%"),
        ]
        table = doc.add_table(rows=1 + len(critical), cols=5)
        table.style = "Table Grid"
        _header_row(table.rows[0], ["Criterio gobernante", "Campo (m)", "Valor", "Observador", "Referencia"])
        for row_index, (criterion, item, key, unit) in enumerate(critical, start=1):
            bg = GRAY_LIGHT if row_index % 2 else WHITE
            value = float(item.get(key, 0) or 0)
            values = [criterion, field_range(item), f"{value:.3f}" if key != "TI" else f"{value:.1f}", observer(item), unit or "—"]
            for cell, text in zip(table.rows[row_index].cells, values):
                _set_cell_bg(cell, bg)
                _cell_text(cell, text, bold=False, size_pt=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Para la única imagen de esta zona se utiliza el campo con peor U0,
        # que permite revisar visualmente el patrón más sensible sin convertir
        # el anexo en una sucesión de mapas repetitivos.
        governing = critical[1][1]
        grid = governing.get("grid_points", []) or []
        if not grid:
            continue
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            from matplotlib.colors import LinearSegmentedColormap

            xs = [float(point.get("x", 0) or 0) for point in grid]
            ys = [float(point.get("y", 0) or 0) for point in grid]
            values = [float(point.get("L", 0) or 0) for point in grid]
            field_start = float(governing.get("field_start", 0) or 0)
            field_end = float(governing.get("field_end", 0) or 0)
            field_luminaires = _field_luminaires_for_report(
                luminaire, field_start, field_end,
            )
            # Arena → ocre → granate: misma lectura de severidad que la paleta
            # anterior, con una saturación visual mayor para distinguir niveles
            # próximos sin recurrir a colores estridentes.
            cmap = LinearSegmentedColormap.from_list(
                "salvi_isocandela",
                ["#E8E0C6", "#D7C29D", "#C48E78", "#A65F68", "#783C58"],
                N=256,
            )
            figure, axis = plt.subplots(figsize=(8.2, 4.25))
            contour = axis.tricontourf(xs, ys, values, levels=12, cmap=cmap)
            figure.colorbar(contour, ax=axis, label="Luminancia (cd/m²)")
            axis.axvline(field_start, color="#1A3A6B", linewidth=0.9,
                         linestyle=(0, (3, 2)), alpha=0.80)
            axis.axvline(field_end, color="#1A3A6B", linewidth=0.9,
                         linestyle=(0, (3, 2)), alpha=0.80)
            if field_luminaires:
                lum_y = [item["y"] for item in field_luminaires]
                for lum_index, item in enumerate(field_luminaires, start=1):
                    axis.scatter(
                        [item["x"]], [item["y"]], s=72, marker="s",
                        facecolor="#FBBF24", edgecolor="#172033", linewidth=1.1,
                        zorder=7,
                    )
                    y_offset = 9 if item["y"] <= sum(lum_y) / len(lum_y) else -13
                    axis.annotate(
                        f"L{lum_index}", xy=(item["x"], item["y"]),
                        xytext=(5, y_offset), textcoords="offset points",
                        fontsize=6.8, fontweight="bold", color="#172033",
                        bbox=dict(boxstyle="round,pad=0.15", fc="#FFFFFF", ec="#172033", lw=0.55),
                        zorder=8,
                    )
                y_min = min(min(ys), min(lum_y)) - 0.32
                y_max = max(max(ys), max(lum_y)) + 0.32
                axis.set_ylim(y_min, y_max)
                detail_lines = []
                for lum_index, item in enumerate(field_luminaires, start=1):
                    def displayed(value, unit=""):
                        try:
                            return f"{float(value):.0f}{unit}"
                        except (TypeError, ValueError):
                            return f"{value}{unit}" if value not in (None, "") else "—"
                    detail_lines.append(
                        f"L{lum_index}: PK {item['x']:.2f} m · Y {item['y']:.2f} m · "
                        f"{item['model']} · óptica {item['optic']} · "
                        f"{displayed(item['power_w'], ' W')} · {displayed(item['current_mA'], ' mA')} · "
                        f"tilt {displayed(item['tilt_deg'], '°')}"
                    )
                figure.text(
                    0.08, 0.025, "Luminarias delimitadoras del campo:  " + "   |   ".join(detail_lines),
                    fontsize=6.5, color="#172033", wrap=True,
                    bbox=dict(boxstyle="round,pad=0.35", fc="#F5F8FC", ec="#B8C7D9", lw=0.7),
                )
                figure.subplots_adjust(bottom=0.20)
            else:
                axis.text(
                    0.01, 0.02, "No hay coordenadas de luminarias disponibles para superponer en este campo.",
                    transform=axis.transAxes, fontsize=7.3, color="#475569",
                    bbox=dict(boxstyle="round,pad=0.25", fc="#F5F8FC", ec="#B8C7D9", lw=0.7),
                )
            axis.set_xlabel("Posición longitudinal (m)", fontsize=8)
            axis.set_ylabel("Posición transversal (m)", fontsize=8)
            axis.set_title(
                f"Campo crítico {field_start:.1f}–{field_end:.1f} m · luminarias delimitadoras",
                fontsize=10, fontweight="bold",
            )
            axis.tick_params(labelsize=7)
            figure.tight_layout(pad=0.8, rect=(0, 0.09 if field_luminaires else 0, 1, 1))
            _save_figure_to_document(doc, figure, width_cm=15.5)
        except Exception:
            continue


def _build_compliance_summary(doc: Document, result: dict,
                              photometric: dict = None) -> None:
    """Matriz compacta por zona para la lectura principal del informe."""
    _section_heading(doc, "Cumplimiento por zona", level=2)
    ph = photometric or {}
    zones = ph.get("zones", {}) if ph.get("available") else {}
    if not zones:
        doc.add_paragraph(
            "No disponible: la verificación CIE 140 con LDT es necesaria para declarar U0, Ul y TI."
        )
        return

    headers = ["Zona", "Lavg / Lreq", "U0", "Ul", "TI (%)", "Estado"]
    table = doc.add_table(rows=1 + len(zones), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], headers)
    for i, (name, data) in enumerate(zones.items()):
        row = table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        checks = data.get("checks", {})
        state = "CUMPLE" if data.get("compliant") else "REVISAR"
        values = [
            name,
            f"{data.get('L_avg', '—')} / {data.get('L_req', '—')}",
            str(data.get("U0", "—")),
            str(data.get("Ul", "—")),
            str(data.get("TI", "—")),
            state,
        ]
        for col, value in enumerate(values):
            ok = col == 5 and state == "CUMPLE"
            _cell_text(
                row.cells[col], value, bold=col in (0, 5), size_pt=8.5,
                align=WD_ALIGN_PARAGRAPH.LEFT if col == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                color="1A7A3C" if ok else ("C0392B" if col == 5 else None),
            )


def _build_installation_summary(doc: Document, luminaire: Any) -> None:
    """Diseño instalable agrupado por tramo; el detalle unitario se mueve al anexo."""
    _section_heading(doc, "5. Diseño de la Instalación")
    lum = _as_luminaire_dict(luminaire)
    zones = lum.get("zones", [])
    if not zones:
        doc.add_paragraph(
            "No hay diseño de luminarias asociado. Calcule las luminarias APHEX para incluir la especificación de instalación."
        )
        return

    totals = lum.get("totals", {})
    lum_spec = lum.get("luminaire") or {}
    _kv_table(doc, [
        ("Disposición", _humanize(lum.get("arrangement"))),
        ("Luminaria de referencia", lum_spec.get("name") or "APHEX configurada por zona"),
        ("Óptica / CCT / I máx.", f"{lum.get('optic', '—')} · {lum.get('cct', '—')} · {lum.get('I_max_mA', '—')} mA"),
        ("Total instalado", f"{totals.get('n_luminaires', '—')} luminarias · {totals.get('power_kw', '—')} kW"),
        ("Densidad media", f"{totals.get('power_density_wm2', '—')} W/m²"),
    ])

    doc.add_paragraph()
    headers = ["Zona", "Tramo (m)", "Modelo / óptica", "mA", "W", "d (m)", "N", "Tilt", "Capa"]
    table = doc.add_table(rows=1 + len(zones), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], headers)
    for i, zone in enumerate(zones):
        row = table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        values = [
            zone.get("zone_name", zone.get("zone_type", "—")),
            f"{zone.get('s_start', '—')}–{zone.get('s_end', '—')}",
            f"{zone.get('model', '—')} / {zone.get('optic', lum.get('optic', '—'))}",
            str(zone.get("current_mA", "—")),
            str(zone.get("power_w", "—")),
            str(zone.get("d_used", "—")),
            str(zone.get("n_luminaires", "—")),
            f"{zone.get('tilt_deg', 0)}°",
            _humanize(zone.get("control_layer")),
        ]
        for col, value in enumerate(values):
            _cell_text(
                row.cells[col], value, bold=col == 0, size_pt=7.5,
                align=WD_ALIGN_PARAGRAPH.LEFT if col in (0, 2, 8) else WD_ALIGN_PARAGRAPH.CENTER,
            )
    note = doc.add_paragraph()
    note.add_run(
        "La planta y esta tabla por tipo de tramo constituyen la especificación ordinaria de montaje. "
        "El listado unitario se incorpora sólo si existen reglajes singulares; el detalle completo queda disponible en la exportación Excel."
    ).italic = True
    note.runs[0].font.size = Pt(8.2)


def _build_operation_summary(doc: Document, luminaire: Any) -> None:
    """Potencia demandada por escena, separada de la tabla de consignas DALI."""
    _section_heading(doc, "Potencia por escena de operación", level=2)
    rows = _operating_scenes(luminaire)
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    _header_row(table.rows[0], ["Escena", "Luminarias activas", "Potencia (kW)", "Flujo (lm)"])
    for i, (key, value) in enumerate(rows):
        row = table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        for col, text in enumerate([
            value.get("name", _humanize(key)),
            str(value.get("active_luminaires", "—")),
            str(value.get("power_kw", "—")),
            str(value.get("flux_lm", "—")),
        ]):
            _cell_text(row.cells[col], text, bold=col == 0, size_pt=8.5,
                       align=WD_ALIGN_PARAGRAPH.LEFT if col == 0 else WD_ALIGN_PARAGRAPH.CENTER)


def _build_energy_and_maintenance(doc: Document, params: dict,
                                  luminaire: Any,
                                  title: str = "8. Potencia, Energía y Mantenimiento") -> None:
    """Cierra la parte operativa con una estimación energética trazable."""
    lum = _as_luminaire_dict(luminaire)
    totals = lum.get("totals", {}) or {}
    scenarios = lum.get("scenarios", {}) or {}
    installed_kw = float(totals.get("power_kw", 0) or 0)
    if not installed_kw and not scenarios:
        return

    _section_heading(doc, title)
    annual_hours = max(0.0, float(params.get("annual_operation_hours", 8760) or 8760))
    tariff = float(params.get("energy_tariff_eur_kwh", 0.15) or 0.15)
    # Hipótesis de explotación, no requisito CIE: un túnel 24 h reparte
    # 8.760 h entre 4.300 h nocturnas y las horas diurnas restantes.
    night_hours = min(
        annual_hours,
        max(0.0, float(params.get("night_operation_hours", 4300) or 4300)),
    )
    day_hours = annual_hours - night_hours
    night_reduced_share = min(
        100.0,
        max(0.0, float(params.get("night_reduced_share_pct", 30.0) or 0.0)),
    ) / 100.0
    day_reference_weights = {
        "sunny": 1200.0, "normal": 1800.0,
        "overcast": 900.0, "dusk": 600.0,
    }
    day_weight_total = sum(day_reference_weights.values())
    hours_by_scene = {
        key: day_hours * weight / day_weight_total
        for key, weight in day_reference_weights.items()
    }
    scene_aliases = {
        "sunny": ("sunny",), "normal": ("normal",),
        "overcast": ("overcast",), "dusk": ("dusk",),
    }
    has_night_normal = isinstance(scenarios.get("night_normal"), dict)
    has_night_reduced = isinstance(scenarios.get("night_reduced"), dict)
    if has_night_normal and has_night_reduced:
        hours_by_scene["night_normal"] = night_hours * (1.0 - night_reduced_share)
        hours_by_scene["night_reduced"] = night_hours * night_reduced_share
        scene_aliases.update({
            "night_normal": ("night_normal",),
            "night_reduced": ("night_reduced",),
        })
    elif isinstance(scenarios.get("night"), dict):
        # La escena histórica ``night`` recibe todas las horas nocturnas.
        hours_by_scene["night"] = night_hours
        scene_aliases["night"] = ("night",)
    elif has_night_normal:
        hours_by_scene["night_normal"] = night_hours
        scene_aliases["night_normal"] = ("night_normal",)
    elif has_night_reduced:
        hours_by_scene["night_reduced"] = night_hours
        scene_aliases["night_reduced"] = ("night_reduced",)
    rows = []
    consumed_source_keys = set()
    for canonical_key, hours in hours_by_scene.items():
        source_key = next(
            (candidate for candidate in scene_aliases[canonical_key]
             if isinstance(scenarios.get(candidate), dict)
             and scenarios[candidate].get("power_kw") is not None),
            None,
        )
        if source_key is None:
            continue
        # Si el motor sólo devuelve una escena de noche, se le asigna el
        # total de horas nocturnas y no se duplica en la tabla.
        if source_key in consumed_source_keys:
            for item_index, item in enumerate(rows):
                if item[4] == source_key:
                    updated_hours = item[1] + hours
                    rows[item_index] = (
                        item[0], updated_hours, item[2], item[2] * updated_hours,
                        source_key,
                    )
                    break
            continue
        scenario = scenarios[source_key]
        consumed_source_keys.add(source_key)
        power = float(scenario.get("power_kw", 0) or 0)
        name = scenario.get("name", _humanize(canonical_key))
        if canonical_key == "night" and name == "Noche":
            name = "Noche (escena única)"
        rows.append((
            name, hours, power, power * hours,
            source_key,
        ))

    if rows:
        total_kwh = sum(item[3] for item in rows)
        represented_hours = sum(item[1] for item in rows)
        table = doc.add_table(rows=1 + len(rows) + 1, cols=4)
        table.style = "Table Grid"
        _header_row(table.rows[0], ["Escena", "Horas/año", "Potencia (kW)", "Consumo (kWh/año)"])
        for index, (name, hours, power, energy, _source_key) in enumerate(rows, start=1):
            bg = GRAY_LIGHT if index % 2 else WHITE
            for column, (cell, value) in enumerate(zip(table.rows[index].cells, [
                name, f"{hours:,.0f}", f"{power:.2f}", f"{energy:,.0f}",
            ])):
                _set_cell_bg(cell, bg)
                _cell_text(cell, value, bold=False, size_pt=8.4,
                           align=WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        for cell, value in zip(table.rows[-1].cells, ["TOTAL", f"{represented_hours:,.0f}", "—", f"{total_kwh:,.0f}"]):
            _set_cell_bg(cell, BLUE_LIGHT)
            _cell_text(cell, value, bold=True, color=BLUE_DARK, size_pt=8.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        p = doc.add_paragraph()
        p.add_run("Coste anual de referencia: ").bold = True
        p.add_run(f"{total_kwh * tariff:,.0f} EUR/año ")
        p.add_run(
            f"(tarifa {tariff:.3f} EUR/kWh; hipótesis de proyecto: "
            f"{annual_hours:,.0f} h/año totales, {night_hours:,.0f} h nocturnas"
            + (
                f" con {night_reduced_share * 100:.0f}% en régimen reducido"
                if has_night_normal and has_night_reduced else ""
            )
            + ")."
        )
        p.runs[-1].italic = True
        p.runs[-1].font.size = Pt(8.2)
        if represented_hours < annual_hours - 1:
            note = doc.add_paragraph()
            note.add_run(
                "Aviso energético: faltan escenas verificadas para cubrir todas las horas adoptadas; el consumo sólo suma las escenas disponibles."
            ).font.size = Pt(8.2)
            note.runs[0].font.color.rgb = RGBColor.from_string("B26A00")
    else:
        doc.add_paragraph(
            f"Potencia instalada: {installed_kw:.2f} kW. No se ha definido una distribución de horas por escena; no se declara consumo anual."
        )

    _section_heading(doc, "Criterios de mantenimiento y explotación", level=2)
    lum_spec = lum.get("luminaire") or {}
    mf = lum_spec.get("maintenance_factor", params.get("maintenance_factor", "—"))
    try:
        mf = f"{float(mf):.2f} (aplicado en el cálculo)"
    except (TypeError, ValueError):
        mf = str(mf)
    _kv_table(doc, [
        ("Factor de mantenimiento aplicado", mf),
        ("Reflectancia de paredes", params.get("rho_wall", params.get("wall_reflectance", "—"))),
        ("Reflectancia de techo", params.get("rho_ceiling", "—")),
        ("Revisión en recepción", "Modelo/LDT, flujo, corriente, altura, posición, tilt, escenas DALI y mediciones finales."),
    ], col_widths=(6.0, 8.0))


def _requires_detailed_schedule(params: dict, luminaire: Any) -> bool:
    """Sólo añade listado punto a punto cuando la instalación es realmente irregular."""
    lum = _as_luminaire_dict(luminaire)
    if bool((params or {}).get("include_detailed_schedule")):
        return True
    return bool(
        (params or {}).get("tilt_overrides")
        or (params or {}).get("tandem_overrides")
        or lum.get("tilt_overrides")
        or lum.get("tandem_overrides")
    )


def _build_installation_schedule(doc: Document, luminaire: Any,
                                 params: dict = None,
                                 photometric: dict = None) -> None:
    """Listado físico y su evidencia CIE 140, una referencia por luminaria."""
    lum = _as_luminaire_dict(luminaire)
    params = params or {}
    lum_params = params.get("luminaire") or params.get("lum_config") or {}
    if not isinstance(lum_params, dict):
        lum_params = {}
    road_width = float(lum.get("road_width_m", 0) or 0)
    arrangement = str(lum.get("arrangement", "central_single") or "central_single")
    wall_offset = float(lum_params.get("wall_offset_m", 0.30) or 0.30)
    if road_width > 0:
        wall_offset = min(max(0.05, wall_offset), max(0.05, road_width / 2.0 - 0.05))

    def rows_for_setpoint(setpoint_index: int) -> list[tuple[str, float]]:
        """Filas físicas con Y desde la pared izquierda, igual que CIE 140."""
        if arrangement in {"central_double", "bilateral_sym", "bilateral"}:
            return [("I", wall_offset), ("D", road_width - wall_offset)]
        if arrangement in {"bilateral_stag", "staggered"}:
            return [("I", wall_offset)] if (setpoint_index - 1) % 2 == 0 else [("D", road_width - wall_offset)]
        if arrangement in {"central_offset", "lateral_left"}:
            return [("I", wall_offset)]
        if arrangement in {"lateral_right", "unilateral"}:
            return [("D", road_width - wall_offset)]
        return [("Eje", road_width / 2.0)]

    positions = []
    for zone_order, zone in enumerate(lum.get("zones", [])):
        for sequence, item in enumerate(zone.get("setpoints", []) or [], start=1):
            index = int(item.get("idx", sequence) or sequence)
            for side, y_pos in rows_for_setpoint(index):
                positions.append({
                    "zone": zone, "item": item, "side": side, "y": y_pos,
                    "index": index, "zone_order": zone_order,
                })
    if not positions:
        return
    positions.sort(key=lambda row: (row["zone_order"], float(row["item"].get("s", 0) or 0), row["y"]))
    for reference, row in enumerate(positions, start=1):
        row["reference"] = f"L-{reference:03d}"

    # The reported d→ is the actual distance to the next luminaire in the
    # same physical row, not the zone average.  It remains meaningful for
    # bilateral layouts and for the two physical units of a tandem pair.
    by_row = {}
    for row in positions:
        key = (row["zone_order"], round(float(row["y"]), 4))
        by_row.setdefault(key, []).append(row)
    for row_group in by_row.values():
        row_group.sort(key=lambda row: float(row["item"].get("s", 0) or 0))
        for current, following in zip(row_group, row_group[1:]):
            current["d_next"] = float(following["item"].get("s", 0) or 0) - float(current["item"].get("s", 0) or 0)
        if row_group:
            row_group[-1]["d_next"] = None

    _section_heading(doc, "ANEXO C. Listado completo de luminarias y reglajes")
    note = doc.add_paragraph()
    note.add_run(
        "C.1 identifica cada luminaria física. PK es la coordenada longitudinal desde la boca A; Y se mide desde la pared izquierda; d→ es la distancia al siguiente punto de la misma fila. "
        "La lente/óptica identifica la fotometría LDT usada en CIE 140. Los valores eléctricos son los del régimen de cálculo diurno/base; las consignas de cada escena DALI se documentan en el capítulo de control."
    ).font.size = Pt(8.3)
    _section_heading(doc, "C.1. Implantación, producto y reglaje base", level=2)
    headers = ["Ref.", "Zona / capa", "Par", "PK (m)", "Y (m)", "d→ (m)", "Lente", "Tilt", "Flujo (lm)", "mA", "W", "Modelo / driver"]
    table = doc.add_table(rows=1 + len(positions), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], headers)

    def _number(value, decimals=2):
        try:
            return f"{float(value):.{decimals}f}"
        except (TypeError, ValueError):
            return "—"

    for i, entry in enumerate(positions):
        zone, item = entry["zone"], entry["item"]
        row = table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        model = str(item.get("model", zone.get("model", "—")) or "—")
        driver = str(item.get("pcb", zone.get("pcb", "")) or "")
        values = [
            entry["reference"],
            f"{zone.get('zone_name', zone.get('zone_type', '—'))} / {_humanize(zone.get('control_layer'))}",
            str(item.get("tandem", "—") or "—"),
            _number(item.get("s")), _number(entry["y"]), _number(entry.get("d_next")),
            item.get("optic", zone.get("optic", lum.get("optic", "—"))),
            f"{_number(item.get('tilt_deg', zone.get('tilt_deg', 0)), 1)}°",
            _number(item.get("flux_lm", zone.get("flux_lm")), 0),
            _number(item.get("current_mA", zone.get("current_mA")), 0),
            _number(item.get("power_w", zone.get("power_w")), 1),
            model if not driver else f"{model} / {driver}",
        ]
        for col, value in enumerate(values):
            _cell_text(row.cells[col], str(value), bold=col == 0, size_pt=6.5,
                       align=WD_ALIGN_PARAGRAPH.LEFT if col in (0, 1, 6, 11) else WD_ALIGN_PARAGRAPH.CENTER)

    fields = ((photometric or {}).get("real_profile") or {}).get("fields", [])
    fields = [field for field in fields if isinstance(field, dict)]
    if not fields:
        return
    _add_page_break(doc)
    _section_heading(doc, "C.2. Campo CIE 140 gobernante asociado a cada luminaria", level=2)
    trace_note = doc.add_paragraph()
    trace_run = trace_note.add_run(
        "Para cada referencia se informa el campo contiguo más desfavorable de su misma zona. Lavg, Lreq, U0 y Ul proceden de la verificación fotométrica CIE 140; no son luminancias puntuales bajo la luminaria. Un guion indica que esa referencia no delimita un campo CIE 140 resuelto; no se rellena con una estimación."
    )
    trace_run.italic = True
    trace_run.font.size = Pt(8.1)
    trace_run.font.color.rgb = RGBColor.from_string("666666")

    def _governing_field(entry):
        zone_name = str(entry["zone"].get("zone_name", "") or "")
        x_pos = float(entry["item"].get("s", 0) or 0)
        same_zone = [field for field in fields if str(field.get("zone_name", "") or "") == zone_name]
        adjacent = [field for field in same_zone if float(field.get("field_start", 0) or 0) - 1e-4 <= x_pos <= float(field.get("field_end", 0) or 0) + 1e-4]
        # In a layered installation, a permanent BASE luminaire may physically
        # delimit a field whose regulatory requirement is governed by CTH/CTR.
        # If its own layer has no field, associate the coincident governing
        # field rather than leaving a real CIE 140 result untraceable.
        coincident = [field for field in fields if float(field.get("field_start", 0) or 0) - 1e-4 <= x_pos <= float(field.get("field_end", 0) or 0) + 1e-4]
        candidates = adjacent or same_zone or coincident
        if not candidates:
            return None
        def _score(field):
            try:
                return float(field.get("L_ratio", float("inf")))
            except (TypeError, ValueError):
                return float("inf")
        return min(candidates, key=_score)

    trace_headers = ["Ref.", "Campo CIE 140 (m)", "Carril", "Lavg (cd/m²)", "Lreq (cd/m²)", "Lavg/Lreq", "U0", "Ul", "TI (%)"]
    trace_table = doc.add_table(rows=1 + len(positions), cols=len(trace_headers))
    trace_table.style = "Table Grid"
    _header_row(trace_table.rows[0], trace_headers)
    for i, entry in enumerate(positions):
        field = _governing_field(entry)
        row = trace_table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        if field:
            values = [
                entry["reference"],
                f"{_number(field.get('field_start'), 1)}–{_number(field.get('field_end'), 1)}",
                f"C{field.get('governing_lane_number', field.get('observer_lane_number', '—'))}",
                _number(field.get("Lavg_governing", field.get("L")), 2),
                _number(field.get("L_required"), 2),
                _number(field.get("L_ratio"), 3),
                _number(field.get("U0"), 3), _number(field.get("Ul"), 3), _number(field.get("TI"), 2),
            ]
        else:
            values = [entry["reference"], "—", "—", "—", "—", "—", "—", "—", "—"]
        for col, value in enumerate(values):
            _cell_text(row.cells[col], str(value), bold=col == 0, size_pt=6.7,
                       align=WD_ALIGN_PARAGRAPH.CENTER)


def _build_conclusions(doc: Document, result: dict, photometric: dict = None,
                       title: str = "7. Conclusiones y Verificaciones de Recepción") -> None:
    """Cierre orientado a la aprobación y a la recepción de obra."""
    status, explanation = _report_status(result, photometric)
    _section_heading(doc, title)
    p = doc.add_paragraph()
    run = p.add_run(f"Estado del diseño: {status}. ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("1A7A3C" if status == "CONFORME" else "B26A00")
    p.add_run(explanation)
    items = [
        "Verificar en obra la geometría efectiva, la posición transversal y la altura de montaje de cada fila.",
        "Comprobar que modelo, óptica, flujo, corriente y orientación/tilt coinciden con el listado de luminarias.",
        "Programar y ensayar las escenas y grupos DALI antes de la recepción de la instalación.",
        "Confirmar las condiciones de mantenimiento y las reflectancias adoptadas si difieren de las existentes en obra.",
    ]
    for item in items:
        bullet = doc.add_paragraph(style="List Bullet")
        bullet.add_run(item).font.size = Pt(9)


def _build_delivery_annex_index(doc: Document, result: dict, params: dict,
                                luminaire: Any = None) -> None:
    """Índice de documentos que componen el expediente de entrega del túnel."""
    summary = result.get("summary", {}) or {}
    tube_id = str(summary.get("tube_id") or params.get("tube_id") or "T1")
    _section_heading(doc, "10. Anexos de Entrega y Documentación Complementaria")
    intro = doc.add_paragraph()
    intro.add_run(
        "Este índice organiza el expediente entregable del tubo. Los anexos técnicos de cálculo "
        "se incorporan a continuación en este Word; los archivos independientes deben emitirse "
        "junto con él y conservar su misma revisión de proyecto."
    ).font.size = Pt(9)

    headers = ["Ref.", "Documento / archivo", "Contenido y finalidad", "Situación de entrega"]
    rows = [
        (
            "A",
            f"Libro Excel de cálculo — Tubo {tube_id}",
            "Resultados, zonas, verificación, consumos, control DALI y listado físico completo de luminarias con coordenadas, potencia, corriente, tilt y grupos.",
            "Generado por SALVI con la exportación Excel; entregar como archivo independiente.",
        ),
        (
            "B",
            "Documento normativo y metodología de cálculo",
            "Documento de apoyo imprimible desde Ayuda › Normativa. Describe el sistema CIE 88 / CIE 140 / CIE 144 e IEC 62386 y el alcance de cada comprobación.",
            "Emitir en el idioma de entrega y archivar junto al informe.",
        ),
        (
            "C",
            "Planos de implantación CAD / PDF",
            "Planta acotada, secciones, posiciones de luminarias, circuitos, grupos DALI, portal A/B, referencias de replanteo y cuadro de leyenda.",
            "Documento de proyecto externo: adjuntar DWG y PDF validados antes de obra.",
        ),
        (
            "D",
            "Fichas de fabricante y fotometrías LDT",
            "Ficha APHEX del modelo/driver instalado, óptica, declaración de prestaciones, curva polar y fichero LDT empleado para la comprobación CIE 140.",
            "Adjuntar al expediente de suministro; las fichas resumidas figuran en la sección 6.",
        ),
        (
            "E",
            "Instrucciones de instalación y puesta en marcha",
            "Fijación, estanqueidad, altura, posición transversal, orientación/tilt, cableado, direccionamiento DALI, escenas, rampas, ensayos y mediciones de recepción.",
            "Adjuntar manual de fabricante y acta de puesta en marcha firmada.",
        ),
        (
            "F",
            "Configuración de control DALI",
            "Exportación del plan de control, direcciones/grupos, escenas de respaldo, curvas de regulación continua, parámetros del luminancímetro y copia de seguridad del controlador.",
            "Exportar desde Control y entregar al integrador / explotación.",
        ),
        (
            "G",
            "Anexos técnicos de cálculo",
            "Clasificación, distancia de parada, L20/Lth, zonas, perfil y campos críticos de verificación CIE 140.",
            "Incorporados en las páginas siguientes de este mismo Word.",
        ),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], headers)
    widths = [1.0, 3.4, 6.5, 4.1]
    for row_index, values in enumerate(rows, start=1):
        row = table.rows[row_index]
        bg = GRAY_LIGHT if row_index % 2 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        for column, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = Cm(widths[column])
            _cell_text(
                cell, value, bold=column in (0, 1), size_pt=7.7,
                align=WD_ALIGN_PARAGRAPH.CENTER if column == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(5)
    run = note.add_run(
        "Nota de trazabilidad: los planos CAD/PDF, manuales y actas son documentos externos al motor de cálculo; "
        "este índice exige su incorporación al expediente, pero no declara su existencia ni su aprobación hasta que se adjunten y revisen."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("666666")


def _append_tube_report(doc: Document, result: dict, params: dict,
                        photometric: dict = None, luminaire: Any = None) -> None:
    """Informe principal seguido de anexos, siempre en el mismo documento Word."""
    _build_executive_summary(doc, result, photometric, luminaire)
    _build_design_inputs(doc, result, params, luminaire)

    _add_page_break(doc)
    _build_methodology(doc, result, photometric, params=params, luminaire=luminaire)
    _build_longitudinal_profile(doc, result, photometric)
    _build_compliance_summary(doc, result, photometric)

    _add_page_break(doc)
    _build_installation_summary(doc, luminaire)
    _build_installation_visuals(doc, result, params, luminaire)
    _build_report_video(doc, params)
    if _as_luminaire_dict(luminaire).get("zones"):
        _add_page_break(doc)
        _build_luminaire_catalog(doc, luminaire)
    _build_control(doc, result, title="7. Control DALI y Escenas de Operación")
    _build_operation_summary(doc, luminaire)
    _build_operation_chart(doc, luminaire)
    _build_energy_and_maintenance(doc, params, luminaire, title="8. Potencia, Energía y Mantenimiento")
    _build_conclusions(doc, result, photometric,
                       title="9. Conclusiones y Verificaciones de Recepción")

    _add_page_break(doc)
    _build_delivery_annex_index(doc, result, params, luminaire)
    _add_page_break(doc)
    _section_heading(doc, "ANEXOS TÉCNICOS DE CÁLCULO")
    intro = doc.add_paragraph()
    intro.add_run(
        "Los anexos conservan la trazabilidad necesaria para auditoría, dirección facultativa y puesta en marcha."
    ).italic = True
    intro.runs[0].font.size = Pt(9)

    _build_classification(doc, result, title="ANEXO A. Clasificación y Distancia de Parada")
    _build_luminances(doc, result, title="ANEXO A.2 Datos L20, Lth e Interior")
    _build_zones(doc, result, title="ANEXO A.3 Zonas y Perfil Normativo")
    _build_quality(doc, result, photometric=photometric,
                   title="ANEXO B. Verificación Fotométrica CIE 140:2019")
    _build_cie140_field_annex(doc, photometric, luminaire=luminaire)
    if _requires_detailed_schedule(params, luminaire):
        _add_page_break(doc)
        _build_installation_schedule(doc, luminaire)


def _build_project_summary(doc: Document, tubes_data: list) -> None:
    """Resumen de proyecto antes de las fichas individuales de cada tubo."""
    _section_heading(doc, "Resumen Comparativo de Tubos")
    headers = ["Tubo", "Longitud", "Lth A / B", "L interior", "CIE 140", "Estado"]
    table = doc.add_table(rows=1 + len(tubes_data), cols=len(headers))
    table.style = "Table Grid"
    _header_row(table.rows[0], headers)
    for i, item in enumerate(tubes_data):
        result = item["result"]
        summary = result.get("summary", {})
        lth = result.get("lth", {})
        status, _ = _report_status(result, item.get("photometric"))
        row = table.rows[i + 1]
        bg = GRAY_LIGHT if i % 2 == 0 else WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        values = [
            summary.get("tube_id", f"T{i + 1}"),
            f"{summary.get('length_m', '—')} m",
            f"{summary.get('Lth', '—')} / {lth.get('Lth_b', '—')}",
            f"{summary.get('Lin', '—')} cd/m²",
            "Disponible" if (item.get("photometric") or {}).get("available") else "Pendiente",
            status,
        ]
        for col, value in enumerate(values):
            _cell_text(row.cells[col], value, bold=col in (0, 5), size_pt=8.5,
                       align=WD_ALIGN_PARAGRAPH.CENTER)


def _build_quality(doc: Document, result: dict, photometric: dict = None,
                   title: str = "5. Criterios de Calidad CIE 88:2004 / CIE 140:2019") -> None:
    validation = result.get('validation', {})
    ph       = photometric or {}
    ph_avail = ph.get('available', False)

    _section_heading(doc, title)

    if ph_avail:
        # ── CIE 140:2019 — tabla por zona ─────────────────────────────────
        p_inf = doc.add_paragraph()
        r_inf = p_inf.add_run(
            f"CIE 140:2019  ·  Tabla {ph.get('rtable','--')}  ·  "
            f"Optica {ph.get('optic','--')}  ·  H = {ph.get('H_m','--')} m"
        )
        r_inf.italic = True
        r_inf.font.size = Pt(8.5)
        r_inf.font.color.rgb = RGBColor.from_string("444444")

        zones  = ph.get('zones', {})
        col_w  = [2.2, 2.5, 1.7, 1.7, 1.7, 2.0]
        tbl    = doc.add_table(rows=1 + len(zones), cols=6)
        tbl.style = 'Table Grid'
        _header_row(tbl.rows[0],
                    ["Zona", "L_avg / L_req", "U0 >=0.40", "Ul >=0.60", "TI <=15%", "Estado"])
        for j, w in enumerate(col_w):
            tbl.rows[0].cells[j].width = Cm(w)

        for i, (zk, zv) in enumerate(zones.items()):
            row = tbl.rows[i + 1]
            bg  = GRAY_LIGHT if i % 2 == 0 else WHITE
            for c in row.cells:
                _set_cell_bg(c, bg)
            for j, w in enumerate(col_w):
                row.cells[j].width = Cm(w)
            _cell_text(row.cells[0], zk, bold=True, size_pt=9)
            if 'error' in zv:
                _cell_text(row.cells[1], str(zv['error']), size_pt=8, color="C0392B")
            else:
                chk = zv.get('checks', {})
                _cell_text(row.cells[1],
                           f"{zv['L_avg']} / {zv['L_req']}",
                           size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER,
                           color="1A7A3C" if chk.get('L_avg') else "C0392B")
                _cell_text(row.cells[2], str(zv['U0']), size_pt=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           color="1A7A3C" if chk.get('U0') else "C0392B")
                _cell_text(row.cells[3], str(zv['Ul']), size_pt=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           color="1A7A3C" if chk.get('Ul') else "C0392B")
                _cell_text(row.cells[4], str(zv['TI']), size_pt=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           color="1A7A3C" if chk.get('TI') else "C0392B")
                st = "CUMPLE" if zv['compliant'] else "REVISAR"
                _cell_text(row.cells[5], st, bold=True, size_pt=9,
                           align=WD_ALIGN_PARAGRAPH.CENTER,
                           color="1A7A3C" if zv['compliant'] else "C0392B")

        lane_rows = [
            (zone_name, lane)
            for zone_name, zone_data in zones.items()
            for lane in zone_data.get('by_lane', [])
        ]
        if lane_rows:
            doc.add_paragraph()
            _section_heading(
                doc,
                "Verificación por carril y sentido",
                level=2,
            )
            lane_tbl = doc.add_table(
                rows=1 + len(lane_rows),
                cols=8,
            )
            lane_tbl.style = 'Table Grid'
            _header_row(
                lane_tbl.rows[0],
                [
                    "Zona", "Carril", "Sentido", "L / Lreq",
                    "U0 carril*", "Ul", "TI", "Estado",
                ],
            )
            for i, (zone_name, lane) in enumerate(lane_rows):
                row = lane_tbl.rows[i + 1]
                bg = GRAY_LIGHT if i % 2 == 0 else WHITE
                for cell in row.cells:
                    _set_cell_bg(cell, bg)
                compliant = bool(lane.get('diagnostic_compliant'))
                direction = (
                    "A → B"
                    if float(lane.get('direction', 1)) >= 0
                    else "B → A"
                )
                values = [
                    zone_name,
                    str(lane.get('lane_number', '—')),
                    direction,
                    (
                        f"{lane.get('L_avg', 0)} / "
                        f"{lane.get('L_required', 0)}"
                    ),
                    str(lane.get('U0', 0)),
                    str(lane.get('Ul', 0)),
                    str(lane.get('TI', 0)),
                    "CUMPLE" if compliant else "REVISAR",
                ]
                for column, value in enumerate(values):
                    _cell_text(
                        row.cells[column],
                        value,
                        bold=column in (0, 7),
                        size_pt=7.5,
                        align=(
                            WD_ALIGN_PARAGRAPH.LEFT
                            if column == 0
                            else WD_ALIGN_PARAGRAPH.CENTER
                        ),
                        color=(
                            "1A7A3C" if compliant else "C0392B"
                        ) if column == 7 else "222222",
                    )
            note = doc.add_paragraph(
                "* U0 por carril es un diagnóstico adicional. "
                "La conformidad normativa de U0 se calcula sobre toda la "
                "calzada para cada observador; Ul se verifica sobre el eje "
                "de cada carril. En el resumen gobierna el peor resultado."
            )
            note.runs[0].font.size = Pt(7.5)
            note.runs[0].font.color.rgb = RGBColor.from_string("666666")

        doc.add_paragraph()
        overall       = ph.get('overall_compliant', False)
        criteria_extra = [
            ("CIE 140:2019 global",          "todas las zonas",
             "CUMPLE" if overall else "REVISAR"),
            ("Perfil longitudinal normativo", "CIE 88:2004",
             "CUMPLE" if validation.get('valid') else "REVISAR"),
            ("Ratios de transicion",         "<= 3:1",
             "CUMPLE" if validation.get('valid') else "REVISAR"),
        ]
    else:
        criteria_extra = [
            ("Uniformidad general Uo",           ">= 0.40",       "Requiere LDT"),
            ("Uniformidad longitudinal Ul",       ">= 0.60",       "Requiere LDT"),
            ("Deslumbramiento TI",               "<= 15 %",       "Requiere LDT"),
            ("Ratio luminancia pared/calzada",    ">= 0.60",       "Requiere LDT"),
            ("Perfil longitudinal normativo",     "CIE 88:2004",
             "CUMPLE" if validation.get('valid') else "REVISAR"),
            ("Ratios de transicion",             "<= 3:1",
             "CUMPLE" if validation.get('valid') else "REVISAR"),
        ]

    table = doc.add_table(rows=1 + len(criteria_extra), cols=3)
    table.style = 'Table Grid'
    _header_row(table.rows[0], ["Criterio", "Requisito", "Estado"])
    for j, w in enumerate([8.0, 2.8, 2.8]):
        table.rows[0].cells[j].width = Cm(w)

    for i, (crit, req, state) in enumerate(criteria_extra):
        row  = table.rows[i + 1]
        bg   = GRAY_LIGHT if i % 2 == 0 else WHITE
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        _set_cell_bg(row.cells[2], bg)

        _cell_text(row.cells[0], crit, size_pt=9)
        _cell_text(row.cells[1], req, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        if "CUMPLE" in state:
            st_color = "1A7A3C"
        elif "REVISAR" in state:
            st_color = "C0392B"
        else:
            st_color = "666666"
        _cell_text(row.cells[2], state, bold=("CUMPLE" in state or "REVISAR" in state),
                   color=st_color, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        for j, w in enumerate([8.0, 2.8, 2.8]):
            row.cells[j].width = Cm(w)

    # Warnings
    warnings = result.get('warnings', [])
    if warnings:
        doc.add_paragraph()
        _section_heading(doc, "Advertencias del cálculo", level=2)
        for w in warnings:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(w)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("C0392B")

    # Validation errors
    val_errors = (validation or {}).get('errors', [])
    if val_errors:
        for e in val_errors:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"ERROR: {e}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("C0392B")


def _build_footer(doc: Document) -> None:
    """Pie de página con número de página y referencia."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SALVI Tunnel Engine (STE) — CIE 88:2004 — ")
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string("888888")

        # Campo de número de página
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run2 = p.add_run()
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = RGBColor.from_string("888888")
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)


# ══════════════════════════════════════════════════════════════════
# FUNCIÓN PRINCIPAL
# ══════════════════════════════════════════════════════════════════

def generate_report(result: dict, params: dict = None, photometric: dict = None,
                    luminaire: Any = None) -> bytes:
    """
    Genera el informe técnico Word a partir del resultado del motor.

    Args:
        result: dict devuelto por run_tunnel_calculation()
        params: dict con los parámetros originales del formulario (opcional)
        photometric: verificación CIE 140 y perfil longitudinal real (opcional)
        luminaire: diseño APHEX con zonas y posiciones de montaje (opcional)

    Returns:
        bytes del .docx listo para enviar como descarga
    """
    if params is None:
        params = {}

    doc = Document()
    _set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)

    # Tipografía corporativa del informe.
    _configure_word_typography(doc)

    # Pie de página
    _build_footer(doc)

    # ── PORTADA ──────────────────────────────────────────────────
    _build_cover(doc, result, params)

    doc.add_paragraph()  # Espacio antes del salto
    _add_page_break(doc)

    _append_tube_report(doc, result, params, photometric=photometric,
                        luminaire=luminaire)

    # Serializar a bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_combined_report(tubes_data: list, project_name: str = "") -> bytes:
    """
    Genera un informe Word combinado con múltiples tubos.
    """
    if not tubes_data:
        raise ValueError("Se necesita al menos un tubo para el informe combinado")

    if len(tubes_data) == 1:
        t = tubes_data[0]
        return generate_report(
            t["result"],
            t.get("params", {}),
            photometric=t.get("photometric"),
            luminaire=t.get("luminaire"),
        )

    doc = Document()
    _set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)

    _configure_word_typography(doc)

    _build_footer(doc)

    first_result = tubes_data[0]["result"]
    first_params  = tubes_data[0].get("params", {})
    if project_name:
        first_params = dict(first_params)
        first_params["project_name"] = project_name

    _build_cover(doc, first_result, first_params)

    doc.add_paragraph()
    _add_page_break(doc)

    _build_project_summary(doc, tubes_data)
    _add_page_break(doc)

    for i, t in enumerate(tubes_data):
        result = t["result"]
        params = t.get("params", {})
        tid    = result.get("summary", {}).get("tube_id", f"T{i+1}")

        if i > 0:
            _add_page_break(doc)

        p_sec = doc.add_paragraph()
        p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sec.add_run(f"-- TUBO {tid} --")
        run.bold = True
        _set_word_font(run.font, TITLE_FONT)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(BLUE_DARK)
        doc.add_paragraph()

        _append_tube_report(
            doc,
            result,
            params,
            photometric=t.get("photometric"),
            luminaire=t.get("luminaire"),
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
