"""Informe técnico V2 de SALVI Tunnel Engine.

La V2 conserva el generador histórico ``report.py`` sin modificarlo. Su foco
es que un revisor pueda reconstruir el camino de cálculo: datos de partida,
regla normativa, valor calculado, eventual decisión de proyecto y evidencia
fotométrica final.
"""

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from .report import (
    BLUE_DARK, BLUE_LIGHT, GRAY_LIGHT, WHITE,
    _add_page_break, _as_luminaire_dict, _build_cie140_field_annex,
    _build_compliance_summary, _build_conclusions, _build_control,
    _build_cover, _build_energy_and_maintenance, _build_footer,
    _build_installation_schedule, _build_installation_summary,
    _build_installation_visuals, _build_longitudinal_profile,
    _build_luminaire_catalog, _build_operation_chart,
    _build_operation_summary, _build_quality, _cell_text,
    _configure_word_typography, _header_row, _humanize, _kv_table,
    _build_methodology,
    _report_status, _requires_detailed_schedule, _section_heading,
    _set_cell_bg, _set_page_margins, _set_word_font,
)


_STANDARD_REFERENCES = [
    (
        "CIE 88:2004",
        "Guide for the Lighting of Road Tunnels and Underpasses.",
        "Clasificación, campo de adaptación, zonas de acceso/umbral/transición/interior y perfil de luminancias.",
    ),
    (
        "CIE 140:2019",
        "Road Lighting Calculations.",
        "Verificación fotométrica de la instalación: retícula de luminancias, U0, Ul y TI a partir de la LDT declarada.",
    ),
    (
        "CIE 144:2001",
        "Road Surface and Lighting.",
        "Tabla R adoptada para relacionar la iluminancia y la luminancia de la calzada.",
    ),
    (
        "Orden Circular 36/2015, Tomo II",
        "Recomendaciones para la iluminación de túneles.",
        "Aplicación nacional: clase de túnel y relación k = Lth/L20 en función de la distancia de parada cuando se selecciona este método.",
    ),
]


def _paragraph(doc: Document, text: str, *, italic: bool = False,
               color: str | None = None, size: float = 9.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _fmt(value: Any, digits: int = 2, unit: str = "") -> str:
    if value in (None, ""):
        return "—"
    try:
        value = float(value)
        return f"{value:.{digits}f}{unit}"
    except (TypeError, ValueError):
        return f"{value}{unit}"


def _portal_rows(result: dict, params: dict) -> list[list[str]]:
    """Una fila por boca que conserva el valor calculado y el de proyecto."""
    summary = result.get("summary", {})
    lth = result.get("lth", {})
    bidirectional = str(params.get("traffic_direction", "one_way")) == "two_way"
    overrides = result.get("project_overrides", {}).get("items", []) or []
    override_names = {
        str(item.get("parameter", item.get("label", ""))) for item in overrides
    }

    portals = [("A", "Lth portal A", summary.get("Lth"), lth.get("Lth_auto"),
                lth.get("L20"), summary.get("SD_m"), params.get("portal_orientation", "—"))]
    if bidirectional:
        portals.append(("B", "Lth portal B", lth.get("Lth_b"), lth.get("Lth_b_auto"),
                        lth.get("L20_b"), lth.get("SD_b_m"), lth.get("orientation_b", "—")))

    rows = []
    for name, override_name, lth_design, lth_reference, l20, dp, orientation in portals:
        is_override = override_name in override_names
        try:
            effective_k = float(lth_reference) / float(l20) if float(l20) > 0 else None
        except (TypeError, ValueError):
            effective_k = None
        rows.append([
            name,
            str(orientation),
            _fmt(l20, 0),
            _fmt(dp, 1),
            _fmt(effective_k, 4),
            _fmt(lth_reference, 1),
            _fmt(lth_design, 1),
            "Valor de proyecto" if is_override else "Valor calculado",
        ])
    return rows


def _build_v2_front_matter(doc: Document, result: dict, params: dict,
                           photometric: dict, luminaire: Any) -> None:
    """Primera página de decisión, pensada para revisión de ingeniería."""
    summary = result.get("summary", {})
    lth = result.get("lth", {})
    status, explanation = _report_status(result, photometric)
    status_color = "1A7A3C" if status == "CONFORME" else (
        "B26A00" if status.startswith("PENDIENTE") else "C0392B"
    )

    _section_heading(doc, "1. Dictamen técnico y alcance")
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    _set_cell_bg(banner.cell(0, 0), "E8F5E9" if status == "CONFORME" else "FFF4E5")
    _cell_text(
        banner.cell(0, 0), f"{status} - {explanation}", bold=True,
        color=status_color, size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _paragraph(
        doc,
        "Este documento es una memoria de cálculo y especificación de implantación. "
        "La conformidad final exige contrastar en obra la geometría, la luminaria instalada, "
        "el fichero fotométrico, los reglajes, el control y las mediciones de recepción.",
        italic=True, color="666666", size=8.5,
    )
    _kv_table(doc, [
        ("Proyecto / tubo", f"{summary.get('project', params.get('project_name', '—'))} / {summary.get('tube_id', '—')}"),
        ("Longitud y circulación", f"{_fmt(summary.get('length_m'), 1, ' m')} / {_humanize(params.get('traffic_direction', 'one_way'))}"),
        ("Condición de diseño", f"v={_fmt(summary.get('speed_kmh'), 0, ' km/h')} · DP A={_fmt(summary.get('SD_m'), 1, ' m')}"),
        ("Luminancia interior", f"Lin={_fmt(summary.get('Lin'), 2, ' cd/m²')} · noche reducida={_fmt(summary.get('L_night'), 2, ' cd/m²')}"),
        ("Criterio de Lth", f"{lth.get('standard', '—')} · método {lth.get('method', '—')}"),
    ])

    lum = _as_luminaire_dict(luminaire)
    totals = lum.get("totals", {}) or {}
    if totals:
        _paragraph(
            doc,
            f"Instalación calculada: {totals.get('n_luminaires', '—')} luminarias físicas; "
            f"potencia instalada {totals.get('power_kw', '—')} kW. La configuración por tramo se declara en la sección 6.",
            size=8.8,
        )


def _build_v2_basis(doc: Document, result: dict, params: dict) -> None:
    summary = result.get("summary", {})
    speed = result.get("speed", {})
    interior = result.get("interior", {})
    l20 = result.get("l20", {})

    _section_heading(doc, "2. Base de cálculo trazable")
    _section_heading(doc, "Datos geométricos, tráfico y pavimento", level=2)
    _kv_table(doc, [
        ("Calzada", f"{_fmt(params.get('road_width_m', params.get('width_m')), 2, ' m')} · {params.get('num_lanes', '—')} carriles de {_fmt(params.get('lane_width_m'), 2, ' m')}"),
        ("Sección", f"{_humanize(params.get('tunnel_shape', '—'))} · altura {_fmt(params.get('height_m'), 2, ' m')}"),
        ("Acera / arcén", f"izq. {_fmt(params.get('sidewalk_left_m', 0), 2, ' m')} / der. {_fmt(params.get('sidewalk_right_m', 0), 2, ' m')}"),
        ("Tráfico de proyecto", f"{params.get('traffic_veh_h', '—')} veh/h · {interior.get('traffic_per_lane_veh_h', '—')} veh/h/carril"),
        ("Pavimento y tabla R", f"{params.get('road_surface', '—')} · Q0={params.get('road_q0', params.get('q0', '—'))}"),
        ("Reflectancias declaradas", f"paredes ρ={params.get('rho_wall', params.get('wall_reflectance', '—'))} · techo ρ={params.get('rho_ceiling', params.get('ceiling_reflectance', '—'))}"),
    ])
    _section_heading(doc, "Distancia de parada", level=2)
    _kv_table(doc, [
        ("Velocidad de cálculo", _fmt(speed.get('v_kmh', summary.get('speed_kmh')), 0, ' km/h')),
        ("Tiempo de reacción", _fmt(speed.get('reaction_time_s'), 2, ' s')),
        ("Coeficiente de fricción", _fmt(speed.get('friction_coefficient'), 3)),
        ("Distancia de reacción", _fmt(speed.get('d_reaction_m'), 1, ' m')),
        ("Distancia de frenado", _fmt(speed.get('d_braking_m'), 1, ' m')),
        ("DP adoptada (portal A)", f"{_fmt(summary.get('SD_m'), 1, ' m')} · origen: {speed.get('SD_source', '—')}"),
    ])
    _paragraph(
        doc,
        "La distancia de parada (DP) delimita la distancia desde la que debe detectarse el obstáculo. "
        "En circulación bidireccional se calcula y declara una DP independiente para la boca B cuando la pendiente efectiva cambia.",
        size=8.8,
    )
    _section_heading(doc, "Campo de adaptación L20", level=2)
    _kv_table(doc, [
        ("Método", f"{l20.get('method', '—')} · confianza {l20.get('confidence', '—')}"),
        ("Entorno / orientación A", f"{_humanize(params.get('environment_type', '—'))} / {params.get('portal_orientation', '—')}"),
        ("Condición de cielo", _humanize(params.get('sky_condition', '—'))),
        ("Resultado", _fmt(l20.get('L20', summary.get('L20')), 0, ' cd/m²')),
        ("Observación", l20.get('note', '—')),
    ])
    _paragraph(
        doc,
        "L20 representa la luminancia media del campo de visión de 20° del conductor antes de entrar. "
        "Cuando se utiliza un modelo o una tabla de entorno, es una estimación de diseño: debe sustituirse o validarse con una evaluación específica del portal si el emplazamiento presenta cielo visible, pantallas, vegetación o geometría singulares.",
        italic=True, color="666666", size=8.5,
    )


def _build_v2_lth_trace(doc: Document, result: dict, params: dict) -> None:
    summary = result.get("summary", {})
    lth = result.get("lth", {})
    method = str(lth.get("method", "k_factor"))
    overrides = result.get("project_overrides", {}) or {}

    _section_heading(doc, "3.6. Determinación de Lth por boca")
    _paragraph(
        doc,
        "Lth es la luminancia de diseño de la zona umbral. Su función es proporcionar, durante la adaptación visual desde el exterior, la luminancia mínima que permita la detección del obstáculo de referencia desde la distancia de parada. El resultado no se infiere de la potencia de la luminaria: se fija antes de diseñar la implantación.",
    )
    if method == "lseq":
        formula = "Método Lseq seleccionado: Lth = ceil(Lseq × Cobs / qc). qc y Cobs participan explícitamente; qc debe proceder de un dato fotométrico o de proyecto validado."
    else:
        formula = "Método k seleccionado: Lth = ceil(k × L20). k se toma de la tabla normativa seleccionada o de un valor de proyecto expresamente declarado."
    _paragraph(doc, formula, italic=True, color=BLUE_DARK, size=9.2)

    _kv_table(doc, [
        ("Norma / regla activa", lth.get("standard", "—")),
        ("Método de cálculo", method),
        ("Clase de túnel", f"adoptada {lth.get('tunnel_class', '—')} · automática {lth.get('calculated_tunnel_class', '—')} · origen {lth.get('tunnel_class_source', '—')}"),
        ("Origen de k", lth.get("k_source", "—")),
        ("k aplicado", _fmt(lth.get("k_factor"), 4)),
        ("qc / Cobs", f"qc={_fmt(lth.get('qc'), 3)} · Cobs={_fmt(lth.get('C_obs'), 3)} · qc interviene: {'sí' if lth.get('qc_used') else 'no'}"),
        ("Lseq", f"{_fmt(summary.get('Lseq'), 0, ' cd/m²')} · origen {lth.get('Lseq_source', '—')}"),
    ])

    rows = _portal_rows(result, params)
    table = doc.add_table(rows=1 + len(rows), cols=8)
    table.style = "Table Grid"
    _header_row(table.rows[0], ["Boca", "Orient.", "L20\n(cd/m²)", "DP\n(m)", "k efectivo", "Lth referencia\n(cd/m²)", "Lth de diseño\n(cd/m²)", "Trazabilidad"])
    widths = [1.1, 1.3, 1.8, 1.4, 1.5, 2.1, 2.1, 2.7]
    for i, values in enumerate(rows, start=1):
        bg = GRAY_LIGHT if i % 2 else WHITE
        for col, value in enumerate(values):
            _set_cell_bg(table.rows[i].cells[col], bg)
            _cell_text(table.rows[i].cells[col], value, bold=col in (0, 5), size_pt=8.0,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            table.rows[i].cells[col].width = Cm(widths[col])

    if overrides.get("has_overrides"):
        _section_heading(doc, "Decisiones de proyecto que sustituyen la referencia", level=2)
        _paragraph(doc, overrides.get("note", ""), italic=True, color="9A4B00", size=8.5)
        items = overrides.get("items", []) or []
        table = doc.add_table(rows=1 + len(items), cols=5)
        table.style = "Table Grid"
        _header_row(table.rows[0], ["Parámetro", "Referencia", "Proyecto", "Diferencia", "Unidad"])
        for index, item in enumerate(items, start=1):
            bg = "FFF4E5" if index % 2 else WHITE
            values = [item.get("parameter", item.get("label", "—")), item.get("cie_value", "—"), item.get("project_value", "—"), item.get("difference", "—"), item.get("unit", "—")]
            for col, value in enumerate(values):
                _set_cell_bg(table.rows[index].cells[col], bg)
                _cell_text(table.rows[index].cells[col], value, bold=col == 0, size_pt=8.2,
                           align=WD_ALIGN_PARAGRAPH.CENTER if col > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    if lth.get("note"):
        _paragraph(doc, f"Nota del motor: {lth['note']}", italic=True, color="666666", size=8.5)


def _build_v2_norms(doc: Document) -> None:
    _section_heading(doc, "3.7. Referencias normativas y límites de validez")
    table = doc.add_table(rows=1 + len(_STANDARD_REFERENCES), cols=3)
    table.style = "Table Grid"
    _header_row(table.rows[0], ["Referencia", "Documento", "Aplicación en esta memoria"])
    for index, values in enumerate(_STANDARD_REFERENCES, start=1):
        bg = GRAY_LIGHT if index % 2 else WHITE
        for col, value in enumerate(values):
            _set_cell_bg(table.rows[index].cells[col], bg)
            _cell_text(table.rows[index].cells[col], value, bold=col == 0, size_pt=8.2,
                       align=WD_ALIGN_PARAGRAPH.LEFT)
    _paragraph(
        doc,
        "La CIE 88 establece el diseño de luminancias y las zonas; la CIE 140 verifica la instalación con la geometría, la LDT y los observadores definidos. La disponibilidad de una LDT no valida por sí sola el producto instalado: antes de recepción deben comprobarse fabricante, referencia, óptica, corriente, flujo, factor de mantenimiento y posición de montaje.",
        italic=True, color="666666", size=8.5,
    )


def _append_v2_tube(doc: Document, result: dict, params: dict,
                    photometric: dict = None, luminaire: Any = None) -> None:
    _build_v2_front_matter(doc, result, params, photometric, luminaire)
    _build_v2_basis(doc, result, params)
    _add_page_break(doc)
    # Se conserva íntegramente el capítulo metodológico de V1. La V2 añade
    # después la cadena explícita L20 -> k -> Lth por boca.
    _build_methodology(doc, result, photometric, params=params, luminaire=luminaire)
    _add_page_break(doc)
    _build_v2_lth_trace(doc, result, params)
    _build_v2_norms(doc)
    _build_longitudinal_profile(doc, result, photometric)
    _build_compliance_summary(doc, result, photometric)

    _add_page_break(doc)
    _build_installation_summary(doc, luminaire)
    _build_installation_visuals(doc, result, params, luminaire)
    if _as_luminaire_dict(luminaire).get("zones"):
        _build_luminaire_catalog(doc, luminaire)
    _build_control(doc, result, title="7. Control, escenas y trazabilidad DALI")
    _build_operation_summary(doc, luminaire)
    _build_operation_chart(doc, luminaire)
    _build_energy_and_maintenance(doc, params, luminaire, title="8. Potencia, energía y mantenimiento")
    _build_conclusions(doc, result, photometric, title="9. Conclusiones y requisitos de recepción")

    _add_page_break(doc)
    _section_heading(doc, "ANEXOS TÉCNICOS AUDITABLES")
    _paragraph(doc, "Los anexos documentan las evidencias de cálculo. Los valores de resultados no sustituyen los planos de obra, el esquema eléctrico ni el protocolo de medición de recepción.", italic=True, size=8.7)
    _build_quality(doc, result, photometric=photometric, title="ANEXO A. Verificación fotométrica CIE 140:2019")
    _build_cie140_field_annex(doc, photometric, luminaire=luminaire)
    if _as_luminaire_dict(luminaire).get("zones"):
        _add_page_break(doc)
        _build_installation_schedule(doc, luminaire, params=params,
                                     photometric=photometric)


def _new_document() -> Document:
    doc = Document()
    _set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)
    _configure_word_typography(doc)
    _build_footer(doc)
    return doc


def generate_report_v2(result: dict, params: dict = None,
                       photometric: dict = None, luminaire: Any = None) -> bytes:
    """Genera la nueva memoria V2 sin alterar el informe histórico."""
    params = params or {}
    doc = _new_document()
    _build_cover(doc, result, params)
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("MEMORIA DE CÁLCULO E IMPLANTACIÓN - V2")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    _add_page_break(doc)
    _append_v2_tube(doc, result, params, photometric, luminaire)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_combined_report_v2(tubes_data: list, project_name: str = "") -> bytes:
    """Genera una memoria V2 para uno o varios tubos, conservando la trazabilidad individual."""
    if not tubes_data:
        raise ValueError("Se necesita al menos un tubo para el informe combinado")
    if len(tubes_data) == 1:
        item = tubes_data[0]
        return generate_report_v2(item["result"], item.get("params", {}), item.get("photometric"), item.get("luminaire"))

    doc = _new_document()
    first = tubes_data[0]
    cover_params = dict(first.get("params", {}))
    if project_name:
        cover_params["project_name"] = project_name
    _build_cover(doc, first["result"], cover_params)
    _add_page_break(doc)
    for index, item in enumerate(tubes_data):
        if index:
            _add_page_break(doc)
        tube_id = item.get("result", {}).get("summary", {}).get("tube_id", f"T{index + 1}")
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"TUBO {tube_id} - MEMORIA V2")
        run.bold = True
        _set_word_font(run.font, "Exposure[-50]")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor.from_string(BLUE_DARK)
        _append_v2_tube(doc, item["result"], item.get("params", {}), item.get("photometric"), item.get("luminaire"))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
